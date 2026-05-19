#!/usr/bin/env python3
"""
Generate a Word document with all user-facing copy from bloodsrus.com,
organized for non-technical review by Dr. Suparno Chakrabarti and Dr. Mahak Agarwal.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LABEL_COLOR = RGBColor(0x55, 0x3C, 0x9A)  # muted purple/dark
HINT_COLOR = RGBColor(0x70, 0x70, 0x70)   # gray
PAGE_HDR_COLOR = RGBColor(0x7A, 0x15, 0x30)  # brand maroon

doc = Document()

# Set overall base font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)


def add_page_break():
    doc.add_page_break()


def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = PAGE_HDR_COLOR
    run.font.name = 'Calibri'


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x2E, 0x4E)


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2F, 0x3E, 0x5E)


def add_label(label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[{label}]")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = LABEL_COLOR
    run.font.name = 'Consolas'


def add_hint(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = HINT_COLOR


def add_para(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_multi_para(text):
    # Split text on \n\n to make multiple paragraphs
    for chunk in text.split("\n\n"):
        chunk = chunk.strip()
        if chunk:
            add_para(chunk)


def add_block(label, hint, content):
    """Standard editable block: label, context hint, then content paragraphs."""
    add_label(label)
    add_hint(hint)
    if isinstance(content, list):
        for c in content:
            add_para(c)
    else:
        add_multi_para(content)


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.size = Pt(11)


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.size = Pt(11)


# ============================================================
# PAGE 1 — INSTRUCTIONS
# ============================================================
add_h1("How to Use This Document")

add_para(
    "This document contains every piece of text currently published on the Bloods R Us website "
    "(bloodsrus.com). It is organized page-by-page, in the same order a visitor sees them, "
    "so you can read through it like a book and revise any copy you would like to change."
)

add_para(
    "Edit the text in the white boxes directly — either in Microsoft Word (using Track Changes: "
    "Review → Track Changes) or in Google Docs (Edit menu → \"Track changes\"). "
    "Please do not change the gray section labels like [HOME_HERO_HEADLINE] — those tell us "
    "where each piece of text lives on the site. When you're done, save and send this file back."
)

add_h3("How to turn on Track Changes in Google Docs")
add_bullets([
    "Upload this .docx to Google Drive and open it with Google Docs.",
    "In the top menu, click Edit → \"Track changes\" (or switch the pencil icon on the top right from \"Editing\" to \"Suggesting\").",
    "Edit away — every change appears as a coloured suggestion that we can accept on our side.",
])

add_h3("How to turn on Track Changes in Microsoft Word")
add_bullets([
    "Open this file in Word.",
    "Click the Review tab → Track Changes.",
    "Edit the white text below each label. Changes will appear coloured; save and return the file.",
])

add_para(
    "If you'd rather just type notes, you can also write in-line comments (Insert → Comment) "
    "next to any block. We will read every note and apply the changes to the live website."
)

add_page_break()


# ============================================================
# PAGE 2 — NAVIGATION & GLOBAL
# ============================================================
add_h1("Navigation & Global Elements")
add_hint("Text that appears in the top menu, footer, and other site-wide elements, on every page.")

add_h2("Top Navigation Menu")

add_block("NAV_LINK_HOME", "Top menu item.", "Home")
add_block("NAV_LINK_ABOUT", "Top menu item.", "About")
add_block("NAV_LINK_CONDITIONS", "Top menu item (opens a dropdown).", "Conditions ▾")
add_block("NAV_LINK_BMT_CENTRE", "Top menu item.", "BMT Centre")
add_block("NAV_LINK_FOR_PATIENTS", "Top menu item.", "For Patients")
add_block("NAV_LINK_FOR_DOCTORS", "Top menu item.", "For Doctors")
add_block("NAV_LINK_CONTACT", "Top menu item.", "Contact")
add_block("NAV_CTA_BOOK_APPOINTMENT", "Blue button in top-right of menu.", "Book Appointment")

add_h3("Conditions Dropdown — Column Headings")
add_block("NAV_DROPDOWN_HEADING_GENETIC", "Heading inside the Conditions dropdown menu.", "Genetic Blood Disorders")
add_block("NAV_DROPDOWN_HEADING_CANCERS", "Heading inside the Conditions dropdown menu.", "Blood Cancers")
add_block("NAV_DROPDOWN_HEADING_LYMPHOMAS", "Heading inside the Conditions dropdown menu.", "Lymphomas & MDS")
add_block("NAV_DROPDOWN_HEADING_MARROW_IMMUNE", "Heading inside the Conditions dropdown menu.", "Bone Marrow & Immune")
add_block("NAV_DROPDOWN_VIEW_ALL", "Link at the bottom of the Conditions dropdown.", "View all conditions →")

add_h3("Mobile Menu Section Headings")
add_block("MOBILE_MENU_NAV", "Section heading inside the mobile hamburger menu.", "Navigation")
add_block("MOBILE_MENU_INFORMATION", "Section heading inside the mobile hamburger menu.", "Information")
add_block("MOBILE_MENU_ALL_CONDITIONS_LINK", "Link inside the mobile menu to the conditions hub.", "All Conditions")
add_block("MOBILE_MENU_CONTACT_LINK", "Link inside the mobile menu.", "Contact & Location")

add_h2("Accessibility / Small UI Labels")
add_block("SKIP_TO_MAIN", "Hidden link for screen-reader users to skip to content.", "Skip to main content")
add_block("SCROLL_HINT", "Small text at the very bottom of the home hero section.", "Scroll to explore")
add_block("FORM_BUTTON_DEFAULT", "Default submit button label on the contact form.", "Send Enquiry")
add_block("FORM_BUTTON_AFTER_SUBMIT", "Button label after the form is submitted.", "Enquiry Sent")
add_block("SEARCH_PLACEHOLDER_HOME", "Placeholder text inside the search box on the home page.", "Search conditions, treatments, FAQs...")
add_block("SEARCH_PLACEHOLDER_HUB", "Placeholder text inside the search box on the conditions hub page.", "Search conditions...")
add_block("SEARCH_NO_RESULTS", "Message shown in search when no matches are found (the quoted word is replaced with the user's query).", 'No results found for "..."')

add_h2("Footer")

add_block("FOOTER_BRAND_SUB", "Short tagline under the logos in the footer.", "Dr. Suparno Chakrabarti's Centre of Excellence for Blood Disorders. Pioneers of Haploidentical BMT.")
add_block("FOOTER_ADDRESS", "Small address line below the phone number in the footer.", "Action Cancer Hospital, A-4, Paschim Vihar, New Delhi")

add_h3("Footer — Column Titles")
add_block("FOOTER_COL_CONDITIONS", "Column heading in the footer.", "Conditions")
add_block("FOOTER_COL_INFORMATION", "Column heading in the footer.", "Information")
add_block("FOOTER_COL_OPENING_HOURS", "Column heading in the footer.", "Opening Hours")
add_block("FOOTER_COL_CONTACT", "Column heading in the footer.", "Contact")

add_h3("Footer — Links")
add_block("FOOTER_LINK_ABOUT_TEAM", "Link in the footer 'Information' column.", "About & Team")
add_block("FOOTER_LINK_BMT_KC", "Link in the footer 'Information' column.", "BMT Knowledge Centre")
add_block("FOOTER_LINK_PHYSICIANS", "Link in the footer 'Information' column.", "For Physicians")
add_block("FOOTER_LINK_PATIENT_GUIDE", "Link in the footer 'Information' column.", "Patient Guide")
add_block("FOOTER_LINK_DONOR_GUIDE", "Link in the footer 'Information' column.", "Donor Guide")
add_block("FOOTER_LINK_RESEARCH", "Link in the footer 'Information' column.", "Research Activities")

add_h3("Footer — Hours & Contact")
add_block("FOOTER_HOURS_DAYS", "Hours line.", "Mon – Sat")
add_block("FOOTER_HOURS_MORNING", "Hours line.", "10:00 AM – 1:00 PM")
add_block("FOOTER_HOURS_EVENING", "Hours line.", "6:00 PM – 8:00 PM")
add_block("FOOTER_APPOINTMENT_LABEL", "Secondary number shown in footer.", "+91-9911800616 (Appointment)")
add_block("FOOTER_WHATSAPP_LABEL", "WhatsApp number shown in footer.", "+91-9871264073 (WhatsApp)")

add_block("FOOTER_COPYRIGHT", "Copyright line at the bottom of every page.", "© 2024 Bloods <em>R</em> Us — Centre of Excellence for Blood Disorders. Action Cancer Hospital, New Delhi.")

add_block("WHATSAPP_FAB_ARIA", "Tooltip on the floating WhatsApp button.", "Chat on WhatsApp")

add_page_break()


# ============================================================
# PAGE: HOME
# ============================================================
add_h1("Home Page")

add_h2("Hero Section (top of the page)")
add_block("HOME_HERO_OVERLINE", "Small line above the main headline.", "Action Cancer Hospital, New Delhi")
add_block("HOME_HERO_HEADLINE", "The largest text on the home page, first thing visitors see. 'Blood Disorders' appears in italics.", "Centre of Excellence for Blood Disorders")
add_block("HOME_HERO_SUBTITLE", "One-line introduction right below the headline.", "Internationally acclaimed haematologists and BMT specialists offering world-class care. Pioneers of Haploidentical Bone Marrow Transplantation in India.")

add_label("HOME_HERO_BULLETS")
add_hint("Six-point bulleted list shown in the hero.")
add_bullets([
    "Centre of excellence for Haploidentical BMT",
    "Successfully done more than 250 Haploidentical BMT",
    "75% long-term survival in high risk blood cancers",
    "Innovative immunotherapies to cure blood cancers",
    "More than 125 peer-reviewed publications",
    "The only Indian centre with international standing",
])

add_block("HOME_HERO_CTA_PRIMARY", "Orange button in the hero.", "Book Appointment")
add_block("HOME_HERO_CTA_PHONE", "Phone number button in the hero.", "+91-8287078906")

add_h2("Statistics Bar (colored band below the hero)")
add_block("HOME_STAT_1_NUMBER", "Big number.", "250+")
add_block("HOME_STAT_1_LABEL", "Label under the number.", "Haploidentical BMTs")
add_block("HOME_STAT_2_NUMBER", "Big number.", "75%")
add_block("HOME_STAT_2_LABEL", "Label under the number.", "Long-term Survival")
add_block("HOME_STAT_3_NUMBER", "Big number.", "125+")
add_block("HOME_STAT_3_LABEL", "Label under the number.", "Peer-reviewed Publications")
add_block("HOME_STAT_4_NUMBER", "Text instead of a number.", "Only")
add_block("HOME_STAT_4_LABEL", "Label under the Only.", "Indian Centre with International Standing")

add_h2("About Bloods <em>R</em> Us (brief section)")
add_block("HOME_ABOUT_OVERLINE", "Small line above the section title.", "About Bloods <em>R</em> Us")
add_block("HOME_ABOUT_TITLE", "Section heading.", "Internationally Acclaimed Excellence in Haematology & BMT")
add_block("HOME_ABOUT_PARA_1", "First paragraph of the About section on home.",
    "Welcome to the official website of Bloods <em>R</em> Us Centre of Excellence for Blood Disorders. Our Internationally and Nationally acclaimed team of Haemato-oncologists and Haematologists has a great profile and extensive experience in BMT.")
add_block("HOME_ABOUT_PARA_2", "Second paragraph of the About section on home.",
    "Dr Suparno Chakrabarti and Dr Mahak Agarwal, along with a dedicated team of Haematopathologists, Paediatric Oncologists, Trained Transplant Nurses, Medical Oncologists, Radiation Oncologists, Surgical Oncologists, Microbiologist (Infection control specialist), Research Scientist, Technicians, Clinical Coordinators, Pharmacists, Dieticians, Physiotherapist and counsellors — ensure that each patient's journey from diagnosis, treatment, and long-term follow-up is integrated and seamless.")
add_block("HOME_ABOUT_CTA", "Button under the About paragraphs.", "Meet the Team")

add_h3("'Our Team Includes' side card")
add_block("HOME_TEAM_CARD_OVERLINE", "Small label at the top of the side card.", "Our Team Includes")
add_label("HOME_TEAM_CARD_LIST")
add_hint("Seven-item list inside the side card.")
add_bullets([
    "Dedicated co-ordinators for BMT and non-BMT patients",
    "Dedicated team of highly skilled nurses",
    "Comprehensively trained dedicated housekeeping staff",
    "Team of trained technicians and scientists to run highly specialised BMT lab",
    "Haematopathologists, Paediatric Oncologists",
    "Medical, Radiation & Surgical Oncologists",
    "Research Scientists, Pharmacists, Dieticians & Counsellors",
])

add_h2("What Makes Us Different (four cards)")
add_block("HOME_DIFF_OVERLINE", "Small line above the section title.", "Why Bloods <em>R</em> Us")
add_block("HOME_DIFF_TITLE", "Section heading.", "What Makes Us Different")

add_block("HOME_DIFF_CARD_1_TITLE", "First card title.", "Pioneers of Haploidentical BMT")
add_block("HOME_DIFF_CARD_1_BODY", "First card body text.",
    "First in India to establish a comprehensive Haploidentical BMT programme. Over 250 successful procedures with outcomes matching the best centres worldwide.")

add_block("HOME_DIFF_CARD_2_TITLE", "Second card title.", "International Training & Standing")
add_block("HOME_DIFF_CARD_2_BODY", "Second card body text.",
    "Our specialists trained at Birmingham, London, Italy, and Seattle's Fred Hutchinson Centre. The only Indian centre with international standing in Haploidentical BMT.")

add_block("HOME_DIFF_CARD_3_TITLE", "Third card title.", "Research-Driven Care")
add_block("HOME_DIFF_CARD_3_BODY", "Third card body text.",
    "125+ peer-reviewed publications. First in the world to identify specific risks in Haploidentical BMT. Every treatment protocol is backed by rigorous research.")

add_block("HOME_DIFF_CARD_4_TITLE", "Fourth card title.", "Integrated Multidisciplinary Team")
add_block("HOME_DIFF_CARD_4_BODY", "Fourth card body text.",
    "Haematopathologists, Paediatric Oncologists, Transplant Nurses, Research Scientists, and Counsellors — ensuring seamless care from diagnosis to long-term follow-up.")

add_h2("Comprehensive Blood Disorder Care (four category cards)")
add_block("HOME_TREAT_OVERLINE", "Small line above the section title.", "Conditions We Treat")
add_block("HOME_TREAT_TITLE", "Section heading.", "Comprehensive Blood Disorder Care")
add_block("HOME_TREAT_BODY", "Description paragraph under the section heading.",
    "From inherited disorders to aggressive blood cancers, our team delivers evidence-based, internationally benchmarked care for patients of all ages.")

add_block("HOME_CAT1_TITLE", "First category card title.", "Genetic Blood Disorders")
add_block("HOME_CAT1_DESC", "First category card description.",
    "Inherited conditions affecting haemoglobin and blood cells. BMT offers up to 90% cure rates.")
add_label("HOME_CAT1_LIST")
add_hint("Conditions listed inside the first category card.")
add_bullets(["Thalassemia", "Sickle Cell Anemia"])
add_block("HOME_CAT1_LINK", "Link text at the bottom of each category card.", "Explore conditions →")

add_block("HOME_CAT2_TITLE", "Second category card title.", "Blood Cancers & Lymphomas")
add_block("HOME_CAT2_DESC", "Second category card description.",
    "Leukemias, lymphomas, and myelodysplastic syndromes. 75% long-term survival with our approach.")
add_label("HOME_CAT2_LIST")
add_hint("Conditions listed inside the second category card.")
add_bullets(["ALL", "AML", "CML", "Hodgkin's", "NHL", "MDS"])

add_block("HOME_CAT3_TITLE", "Third category card title.", "Bone Marrow Disorders")
add_block("HOME_CAT3_DESC", "Third category card description.",
    "When the bone marrow fails to produce healthy blood cells. BMT cures 70–90% of patients.")
add_label("HOME_CAT3_LIST")
add_hint("Conditions listed inside the third category card.")
add_bullets(["Aplastic Anemia", "Bone Marrow Failure Syndromes"])

add_block("HOME_CAT4_TITLE", "Fourth category card title.", "Immune & Metabolic Disorders")
add_block("HOME_CAT4_DESC", "Fourth category card description.",
    "Autoimmune conditions, immunodeficiencies, and inherited metabolic diseases treatable with BMT.")
add_label("HOME_CAT4_LIST")
add_hint("Conditions listed inside the fourth category card.")
add_bullets(["Autoimmune Diseases", "Primary Immunodeficiency", "Metabolic Disorders"])

add_h2("Patient Testimonials (home page)")
add_block("HOME_TESTI_OVERLINE", "Small line above the section title.", "Global Impact")
add_block("HOME_TESTI_TITLE", "Section heading.", "Lives Transformed")
add_block("HOME_TESTI_BODY", "Paragraph under the heading.",
    "Patients from across India and internationally trust Bloods <em>R</em> Us for life-changing treatment.")

add_block("HOME_TESTI_1_QUOTE", "First testimonial quote.",
    '"I have taken treatment of Acute Myeloid Leukemia (AML) for my daughter. My daughter had a Successful Bone Marrow Transplant."')
add_block("HOME_TESTI_1_AUTHOR", "First testimonial author.", "Aimira Isamidinov")
add_block("HOME_TESTI_1_ROLE", "First testimonial role.", "Patient's Mother")

add_block("HOME_TESTI_2_QUOTE", "Second testimonial quote.",
    '"The care and dedication shown by the entire Bloods <em>R</em> Us team is truly remarkable. From diagnosis to treatment and follow-up, every step was handled with professionalism and compassion."')
add_block("HOME_TESTI_2_AUTHOR", "Second testimonial author.", "Nitin Garg")
add_block("HOME_TESTI_2_ROLE", "Second testimonial role.", "Patient")

add_block("HOME_TESTI_3_QUOTE", "Third testimonial quote.",
    '"Success Story of Fatima Yahaya Zango from Nigeria. Treated for Sickle Cell Disease."')
add_block("HOME_TESTI_3_AUTHOR", "Third testimonial author.", "Fatima Yahaya Zango")
add_block("HOME_TESTI_3_ROLE", "Third testimonial role.", "Patient, Nigeria — Sickle Cell Disease")

add_block("HOME_TESTI_MORE_LINK", "Link under the three testimonials.", "Read all patient stories →")

add_h2("CTA Bar (colored strip at bottom of home page)")
add_block("HOME_CTA_TITLE", "Large text in the bottom CTA strip.", "Ready to speak with our specialists?")
add_block("HOME_CTA_SUBTITLE", "Small line under the CTA title.",
    "Appointments available Mon–Sat, 10 AM–1 PM & 6–8 PM. Teleconsults available.")
add_block("HOME_CTA_BUTTON_BOOK", "Orange button.", "Book Appointment")
add_block("HOME_CTA_BUTTON_PHONE", "Phone button.", "+91-8287078906")
add_block("HOME_CTA_BUTTON_WHATSAPP", "Green WhatsApp button.", "WhatsApp")

add_page_break()


# ============================================================
# PAGE: ABOUT
# ============================================================
add_h1("About Page")

add_h2("Page Hero")
add_block("ABOUT_HERO_OVERLINE", "Small label at the top of the page hero.", "Bloods <em>R</em> Us")
add_block("ABOUT_HERO_TITLE", "Main page title.", "About Our Team & Foundation")
add_block("ABOUT_HERO_SUBTITLE", "Subtitle under the page title.",
    "Internationally and nationally acclaimed haematologists dedicated to excellence in blood disorder care and research.")

add_h2("Meet the Team — Section Heading")
add_block("ABOUT_TEAM_OVERLINE", "Small label above section title.", "Our Specialists")
add_block("ABOUT_TEAM_TITLE", "Section title.", "Meet the Team")

add_h3("Dr. Suparno Chakrabarti — Profile Card")
add_block("ABOUT_SUPARNO_NAME", "Doctor name on profile card.", "Dr. Suparno Chakrabarti")
add_block("ABOUT_SUPARNO_ROLE", "Role/affiliations line under the name.",
    "Principal Director, Action Institute for Blood Diseases, Transplantation & Cellular Therapy (AIBTCT) · Action Cancer Hospital · Hon Professor, Amity University · Adjunct Professor, Jamia Hamdard University · Director & Senior Scientist, Manashi Chakrabarti Foundation")
add_block("ABOUT_SUPARNO_CREDENTIALS", "Degrees / credentials line.",
    "MD (Internal Medicine, PGIMER, Chandigarh) • Doctor of Medicine (Hematopoietic Cell Transplantation, University of Birmingham, UK) • FRCPath (Haematology, Royal College of Pathologists, London) • CCT in Haematology (UK)")
add_block("ABOUT_SUPARNO_SPECIALIZATION", "Specialization paragraph. Begins with 'Specialization:' in bold on the website.",
    "Specialization: Blood and Marrow Transplantation, Haploidentical BMT, T Cell Co-stimulation Blockade, Adaptive NK Cell Biology in Transplantation, NK Cell-Based Immunotherapy, Viral Immunology Post-BMT, Cellular Therapy, Hemato-Oncology, Mesenchymal Cell Therapy for GVHD, Refractory Lymphoma & Myeloma, Post-Transplant Immune Reconstitution.")
add_block("ABOUT_SUPARNO_TRAINING", "Training paragraph. Begins with 'Training:' in bold.",
    "Training: 13 years in the UK as research fellow then consultant in BMT. Played a substantial role in developing Campath-1H based T cell depletion and reduced intensity conditioning.")
add_block("ABOUT_SUPARNO_SOCIETY_ROLES", "Academic & Society Roles paragraph.",
    "Academic & Society Roles: Peer reviewer — Blood, JAMA, Lancet, Haematologica, BJH • CIBMTR Expert Panel (Post-Transplant Infections) • EBMT Working Party (Infectious Diseases) • Editorial board member, multiple journals • Ex-member: ASH, BSH, BSBMT, ASBMT")

add_h3("Dr. Mahak Agarwal — Profile Card")
add_block("ABOUT_MAHAK_NAME", "Doctor name on profile card.", "Dr. Mahak Agarwal")
add_block("ABOUT_MAHAK_ROLE", "Role/affiliations line under the name.",
    "Associate Consultant, Action Institute for Blood Diseases, Transplantation and Cellular Therapy (AIBTraCT) · Action Cancer Hospital")
add_block("ABOUT_MAHAK_CREDENTIALS", "Degrees / credentials line.",
    "MBBS • MD • PGDMLS (Symbiosis Centre for Health Care, Pune) • PGDHHM (Symbiosis Centre for Health Care, Pune) • Fellowship in Clinical Hematology & BMT (Dharamshila Cancer Hospital, New Delhi) • Specialized Training in GMP-Based Cellular Therapy • Clinical Observer, ACTREC, Tata Memorial Centre, Mumbai")
add_block("ABOUT_MAHAK_SPECIALIZATION", "Specialization paragraph.",
    "Specialization: Blood and Marrow Transplantation, Anemia Treatment, Leukemia & Lymphoma, Myeloma, Sickle Cell Disease, Thalassemia, Bone Marrow Transplant, NK Cell-Based Cellular Therapy, CAR-T Cell Therapy, Bispecific Antibody Therapy, Extracorporeal Photopheresis & GVHD Management.")

add_h3("Full Team Composition Box")
add_block("ABOUT_TEAM_COMPOSITION_HEADING", "Heading above the full team paragraph.", "Full Team Composition")
add_block("ABOUT_TEAM_COMPOSITION_TEXT", "Long paragraph listing every role in the team.",
    "Dr Suparno Chakrabarti • Dr Mahak Agarwal • Haematopathologists • Paediatric Oncologists • Trained Transplant Nurses • Medical Oncologists • Radiation Oncologists • Surgical Oncologists • Microbiologist (Infection Control Specialist) • Research Scientist • Technicians • Clinical Coordinators • Pharmacists • Dieticians • Physiotherapist • Counsellors")

add_h2("Manashi Chakrabarti Foundation — In Memoriam")
add_block("ABOUT_MCF_OVERLINE", "Small label above section title.", "In Memoriam")
add_block("ABOUT_MCF_TITLE", "Section title.", "Manashi Chakrabarti Foundation")
add_block("ABOUT_MCF_SUBTITLE", "Subtitle paragraph under the section title.",
    "A tribute to the life, spirit, and enduring legacy of Manashi Chakrabarti — whose memory inspires our mission of research and care for children with life-threatening blood disorders.")

add_block("ABOUT_MCF_BIO_PARA_1", "Paragraph 1 of Manashi's biography.",
    "Born as the first child to Ganendranath Chakrabarti and Uma Devi in 1939, Manashi was very special to her parents. Her father, an eminent pathologist, noted some rare talents in the young girl – a mind that could perceive the abstract and express it through her paintings and sketches. As is the story of many such families of her time, Manashi's talents were suffocated by the existence of a middle class family that found painting to be an unnecessary pastime.")
add_block("ABOUT_MCF_BIO_PARA_2", "Paragraph 2 of Manashi's biography.",
    "At the age of 20, Manashi got married to Bimalangshu. Bimalangshu hailed from a family in Rajshahi, a part of East Bengal that became East Pakistan in 1947. Bimalangshu, a bright medical student at that time, joined the army and became a distinguished soldier and anaesthetist. Manashi adapted to living in a large joint family, but, the promise and talents were submerged in her chores. She raised two children, Mousumi and Suparno. In 1974, she lost her younger brother at the age of 21. This changed her life forever. She started to seek the truth about this world and the hereafter.")
add_block("ABOUT_MCF_BIO_PARA_3", "Paragraph 3.",
    "Ganendranath had retired as a frustrated academic pathologist, getting little recognition for his brilliance. This pained Manashi. She felt that her children must fulfil the academic promise that she and her father were not allowed to express. Suparno went abroad to earn further expertise in medicine, promising to come back to her side and fulfil the family's mission of research in medicine.")
add_block("ABOUT_MCF_BIO_PARA_4", "Paragraph 4.",
    "On 13th June 2005, Bimalangshu was admitted to the hospital intensive care with a chest infection. Manashi was by his side for the next few days. On 17th June, she suffered a massive heart attack, which she had smilingly defeated when only a few days old, to grace this world with all her talents.")
add_block("ABOUT_MCF_BIO_PARA_5", "Paragraph 5 (closing).",
    "Death could not have taken Manashi's legacy away from us. It will live on with this charity founded by her children to fulfil the family's mission of research in medicine and the care of children with life-threatening blood disorders.")

add_block("ABOUT_MCF_POEM", "Poem shown in a highlighted side block.",
    "'It's time to say good-bye\n"
    "But, please don't cry.\n"
    "When I have left for a different shore,\n"
    "Just gently shut the door.\n"
    "Yet let the memories stay,\n"
    "Bright as the sunshine may.\n"
    "Then, every drop of rain you'll find,\n"
    "Leaves a rainbow behind.'")

add_h2("Research Activities (four cards)")
add_block("ABOUT_RESEARCH_OVERLINE", "Small label.", "Manashi Chakrabarti Foundation")
add_block("ABOUT_RESEARCH_TITLE", "Section title.", "Research Activities")

add_h3("Research Card 1 — Children with Blood Disorders")
add_block("ABOUT_RESEARCH1_OVERLINE", "Card label.", "1. Children with Blood Disorders")
add_block("ABOUT_RESEARCH1_TITLE", "Card heading.", "Supporting the cause of children suffering from blood disorders")
add_block("ABOUT_RESEARCH1_PARA_1", "Paragraph 1.",
    "Blood cancers account for half of childhood malignancies. They are broad of two types, myeloid and lymphoid. Acute Lymphoblastic Leukemia or ALL account for 90% of blood cancers between the ages of 1-16 years. Diligent research and exhaustive clinical trials have now made it one of the most curable cancers. In the western world, 80% of children diagnosed with ALL get cured with chemotherapy alone. Of the other 20% which relapse, 50-60% are cured with further treatment including a bone marrow transplantation.")
add_block("ABOUT_RESEARCH1_PARA_2", "Paragraph 2.",
    "In developing countries including India, the majority of children with ALL fail to access proper healthcare facilities and even if they do so, a large number of the default treatment due to logistic and financial constraints.")
add_block("ABOUT_RESEARCH1_PARA_3", "Paragraph 3.",
    "Our endeavor is to provide awareness, access, and support to all such children who could lead a happy and healthy life and productively contribute to the development of the society. Thus, the journey does not end with successful treatment. That is the beginning of a long and productive life through proper guidance and support for educational and psychological rehabilitation.")
add_block("ABOUT_RESEARCH1_PARA_4", "Paragraph 4.",
    "The success in treatment of cancer depends on the understanding and amalgamation of individual biology and the environment we live in. The factors predisposing to childhood cancers are poorly understood. In a country where air and water pollution is on the rise and many carcinogenic substances are used in daily life without regulation, we need to know the cause of blood cancers in our children.")
add_block("ABOUT_RESEARCH1_PARA_5", "Paragraph 5.",
    "Disease biology differs from geographical and ethnic variations. Little do we know if the biology of childhood leukemia in India is different from that in the west. Uncompromised research on each of these areas is the need of the hour. Our organization is striving to gather the infrastructure and human resources to start answering these questions. We invite all interested researchers, collaborators, philanthropists to join us in this endeavor.")

add_h3("Research Card 2 — BMT Research")
add_block("ABOUT_RESEARCH2_OVERLINE", "Card label.", "2. BMT Research")
add_block("ABOUT_RESEARCH2_TITLE", "Card heading.", "Research on Bone Marrow Transplantation (BMT) in Children")
add_block("ABOUT_RESEARCH2_PARA_1", "Paragraph 1.",
    "BMT from a donor or Allogeneic BMT is often the only curative treatment for advanced blood cancers. If a matched donor is not available in the family, the patient can go for a transplant from an unrelated donor.")
add_block("ABOUT_RESEARCH2_PARA_2", "Paragraph 2.",
    "We inherit half of our HLA genes from each parent and pass it on likewise to our children. Thus, HLA or tissue type is 50% matched between the children and their parents. This is called a Haploidentical match.")
add_block("ABOUT_RESEARCH2_PARA_3", "Paragraph 3.",
    "We were the first in the world to point out that this approach for children was wrought with high risks of rejection and Graft Versus Host Disease (GVHD). Through years of diligent clinical research, we have now developed the most effective way of transplanting such children from a parent or a haploidentical donor. This discovery has changed the lives of many such children. Our organization has provided expertise and logistic support for such research activities over the last decade.")

add_h3("Research Card 3 — Cancer Survivorship")
add_block("ABOUT_RESEARCH3_OVERLINE", "Card label.", "3. Cancer Survivorship")
add_block("ABOUT_RESEARCH3_TITLE", "Card heading.", "Surviving after Blood Cancer")
add_block("ABOUT_RESEARCH3_PARA_1", "Paragraph 1.",
    "Following the ordeal of going through treatment for blood cancer, the biggest challenge lies in rehabilitation and long-term surveillance. Growth and mental development of the survivors of childhood cancer is of paramount importance.")
add_block("ABOUT_RESEARCH3_PARA_2", "Paragraph 2.",
    "In addition, cancer drugs can also have an effect on the heart and the lungs and need regular monitoring. Society at large should take the responsibility of rehabilitating the survivors of childhood cancer, so that they can take up a lead role in society in future.")

add_h3("Research Card 4 — Thalassemia")
add_block("ABOUT_RESEARCH4_OVERLINE", "Card label.", "4. Thalassemia")
add_block("ABOUT_RESEARCH4_TITLE", "Card heading.", "A life free of Thalassemia")
add_block("ABOUT_RESEARCH4_PARA_1", "Paragraph 1.",
    "β-Thalassemia Major is the commonest genetic disorder in India. About 10,000 children with thalassemia are born every year. Despite the improvement in supportive care, the long-term outcome of children with transfusion-dependent β thalassemia in developing countries is disappointing.")
add_block("ABOUT_RESEARCH4_PARA_2", "Paragraph 2.",
    "Recent data from WHO confirms that about 12% of children born with transfusion-dependent β thalassemia are actually transfused, and less than 5% receive adequate iron chelation. BMT from a matched sibling donor was established as a curative treatment for this condition in the early eighties. However, only 10-20% of thalassemia sufferers find a matched family donor.")
add_block("ABOUT_RESEARCH4_PARA_3", "Paragraph 3.",
    "Our endeavor is to develop facilities for transfusion and chelation for children under the optimum guidance of pediatric hematologists. We also provide expertise in curative treatment for this condition.")

add_page_break()


# ============================================================
# PAGE: CONDITIONS HUB
# ============================================================
add_h1("Conditions Hub Page")
add_hint("The landing page visitors see when they click 'Conditions' in the top menu. Lists all conditions grouped into four categories.")

add_h2("Conditions Hub — Page Header")
add_block("CONDITIONS_HERO_OVERLINE", "Small label above the title.", "Clinical Reference")
add_block("CONDITIONS_HERO_TITLE", "Page title.", "Blood Disorder Information")
add_block("CONDITIONS_HERO_SUBTITLE", "Subtitle paragraph.",
    "Comprehensive clinical information authored by Dr. Suparno Chakrabarti and Dr. Mahak Agarwal.")

add_block("CONDITIONS_HUB_OVERLINE", "Label on the hub landing page.", "All Conditions")
add_block("CONDITIONS_HUB_TITLE", "Heading on the hub landing page.", "Find Your Condition")
add_block("CONDITIONS_HUB_BODY", "Paragraph under the heading.",
    "Select a category to learn about specific conditions and how we treat them.")

add_h2("Category 1 — Genetic Blood Disorders")
add_block("HUB_CAT_GENETIC_TITLE", "Category section heading.", "Genetic Blood Disorders")

add_block("HUB_THAL_TAG", "Tag label on condition card.", "Genetic Disorder")
add_block("HUB_THAL_TITLE", "Condition card title.", "Thalassemia")
add_block("HUB_THAL_DESC", "Condition card description.",
    "Inherited blood disorder with low haemoglobin. BMT offers up to 90% cure rate in eligible patients.")

add_block("HUB_SICKLE_TAG", "Tag label on condition card.", "Genetic Disorder")
add_block("HUB_SICKLE_TITLE", "Condition card title.", "Sickle Cell Anemia")
add_block("HUB_SICKLE_DESC", "Condition card description.",
    "Common inherited RBC disorder. Haploidentical BMT now offers curative potential for most patients.")

add_h2("Category 2 — Blood Cancers")
add_block("HUB_CAT_CANCERS_TITLE", "Category section heading.", "Blood Cancers")

add_block("HUB_ALL_TAG", "Tag label on condition card.", "Blood Cancer")
add_block("HUB_ALL_TITLE", "Condition card title.", "Acute Lymphoblastic Leukemia")
add_block("HUB_ALL_DESC", "Condition card description.",
    "Commonest childhood cancer. 90% of standard-risk children cured with chemotherapy.")

add_block("HUB_AML_TAG", "Tag label on condition card.", "Blood Cancer")
add_block("HUB_AML_TITLE", "Condition card title.", "Acute Myeloid Leukemia")
add_block("HUB_AML_DESC", "Condition card description.",
    "BMT reduces relapse risk by 80% compared to chemotherapy alone.")

add_block("HUB_CML_TAG", "Tag label on condition card.", "Blood Cancer")
add_block("HUB_CML_TITLE", "Condition card title.", "Chronic Myeloid Leukemia")
add_block("HUB_CML_DESC", "Condition card description.",
    "Targeted therapy with TKIs. BMT for resistant or accelerated phase disease.")

add_h2("Category 3 — Lymphomas & MDS")
add_block("HUB_CAT_LYMPHOMAS_TITLE", "Category section heading.", "Lymphomas & MDS")

add_block("HUB_HL_TAG", "Tag label on condition card.", "Lymphoma")
add_block("HUB_HL_TITLE", "Condition card title.", "Hodgkin's Lymphoma")
add_block("HUB_HL_DESC", "Condition card description.",
    "Highly treatable cancer with >95% cure rate in limited disease. BMT for relapsed cases.")

add_block("HUB_NHL_TAG", "Tag label on condition card.", "Lymphoma")
add_block("HUB_NHL_TITLE", "Condition card title.", "Non-Hodgkin's Lymphoma")
add_block("HUB_NHL_DESC", "Condition card description.",
    "Diverse group of lymphoid cancers. 70% cure rate with Haploidentical BMT in relapsed/high-risk cases.")

add_block("HUB_MDS_TAG", "Tag label on condition card.", "Bone Marrow Disorder")
add_block("HUB_MDS_TITLE", "Condition card title.", "Myelodysplastic Syndromes")
add_block("HUB_MDS_DESC", "Condition card description.",
    "Allogeneic BMT is the only curative treatment. Best results when done early.")

add_h2("Category 4 — Bone Marrow & Immune Disorders")
add_block("HUB_CAT_BMID_TITLE", "Category section heading.", "Bone Marrow & Immune Disorders")

add_block("HUB_AA_TAG", "Tag label on condition card.", "Bone Marrow Failure")
add_block("HUB_AA_TITLE", "Condition card title.", "Aplastic Anemia")
add_block("HUB_AA_DESC", "Condition card description.",
    "Bone marrow stops producing blood cells. BMT cures 90% of patients when done early.")

add_block("HUB_BMF_TAG", "Tag label on condition card.", "Bone Marrow Failure")
add_block("HUB_BMF_TITLE", "Condition card title.", "Inherited Bone Marrow Failure")
add_block("HUB_BMF_DESC", "Condition card description.",
    "Fanconi Anemia, DKC, SDS and others. 70–80% cure rates with Allogeneic BMT.")

add_block("HUB_AUTO_TAG", "Tag label on condition card.", "Immune Disorder")
add_block("HUB_AUTO_TITLE", "Condition card title.", "Autoimmune Diseases")
add_block("HUB_AUTO_DESC", "Condition card description.",
    "BMT can 'reboot' the immune system. Most promising for Scleroderma and Multiple Sclerosis.")

add_block("HUB_PID_TAG", "Tag label on condition card.", "Pediatric")
add_block("HUB_PID_TITLE", "Condition card title.", "Primary Immunodeficiency")
add_block("HUB_PID_DESC", "Condition card description.",
    "Nearly 100% cure rate with T cell depleted Haploidentical BMT in our experience.")

add_block("HUB_IMD_TAG", "Tag label on condition card.", "Pediatric")
add_block("HUB_IMD_TITLE", "Condition card title.", "Inherited Metabolic Disorders")
add_block("HUB_IMD_DESC", "Condition card description.",
    "BMT from healthy donor populates body with enzyme-producing cells, improving organ function.")

add_page_break()


# ============================================================
# CONDITION DETAIL PAGES - Helper
# ============================================================
def condition_cta(prefix, title, subtitle, show_phone=True):
    add_h3("Bottom CTA Bar")
    add_block(f"{prefix}_CTA_TITLE", "Heading of the CTA strip at the bottom of the condition page.", title)
    add_block(f"{prefix}_CTA_SUBTITLE", "Subtitle under the CTA title.", subtitle)
    add_block(f"{prefix}_CTA_BUTTON", "Orange button.", "Book Consultation")
    if show_phone:
        add_block(f"{prefix}_CTA_PHONE", "Phone button.", "+91-8287078906")


# ============================================================
# CONDITION: THALASSEMIA
# ============================================================
add_h1("Condition — Thalassemia")
add_hint("This is a sub-page under Conditions. It has two tabs: Pediatric and Adult.")

add_block("THAL_TAB_LABEL_PED", "Tab button text.", "In Children (Pediatric)")
add_block("THAL_TAB_LABEL_ADULT", "Tab button text.", "In Adults")

# ----- Pediatric tab -----
add_h2("Pediatric Tab")
add_block("THAL_PED_OVERLINE", "Overline text on the pediatric tab.", "Pediatric BMT — Dr. Suparno Chakrabarti")
add_block("THAL_PED_TITLE", "Pediatric tab title.", "Thalassemia in Children")

add_h3("FAQ — What is Thalassemia?")
add_block("THAL_PED_FAQ1_Q", "Accordion (FAQ) question.", "What is Thalassemia and Types of Thalassemia?")
add_block("THAL_PED_FAQ1_A1", "Answer paragraph 1.",
    "Thalassemia is an inherited blood disorder characterised by low haemoglobin (Haemoglobin is an iron rich protein present within the Red Blood Cells of the body. Haemoglobin is responsible for carrying oxygen to all parts of the body). Normal haemoglobin binds iron with alfa and beta globin (protein) chains. Due to genetic mutations in both the genes inherited from the parents, either alfa or beta chains are absent or reduced. This results in precipitation of other globin chains inside the red blood cells, leading to deformity of the red cells. These deformed red cells get destroyed either within the bone marrow (spongy material within the bones) or spleen resulting in anemia (Pallor). This is called Thalassemia Major.")
add_block("THAL_PED_FAQ1_A2", "Answer paragraph 2.",
    "When an individual inherits only one abnormal gene, the manifestations are very mild. This is called Thalassemia Minor or Thalassemia Trait.")
add_block("THAL_PED_FAQ1_A3", "Answer paragraph 3.",
    "The commonest form of Thalassemia in India is Beta-Thalassemia, where Beta-globin chains are deficient. When instead of absolute absence of one type of globin chain, there is reduction only; the condition is called Thalassemia Intermedia.")
add_block("THAL_PED_FAQ1_A4", "Answer paragraph 4.",
    "Other Thalassemia like Syndromes such as HbE, HbD etc. when co-inherited with Beta-thalassemia gene give rise to Thalassemia Intermedia.")

add_h3("FAQ — Symptoms")
add_block("THAL_PED_FAQ2_Q", "Accordion question.", "Symptoms")
add_block("THAL_PED_FAQ2_H1", "Sub-heading inside answer.", "Symptoms of Thalassemia Major")
add_block("THAL_PED_FAQ2_A1", "Answer paragraph.",
    "The symptoms appear between 3-7 years age. The child starts becoming listless, irritable and pale. The parents often visit the doctor several times before this is detected. The spleen enlarges in size and the eyes turn yellow from jaundice. The child is unable to grow and play like other children of his age.")
add_block("THAL_PED_FAQ2_H2", "Sub-heading.", "Symptoms of Thalassemia Intermedia")
add_block("THAL_PED_FAQ2_A2", "Answer paragraph.",
    "This depends on the severity of the affection of globin gene production. In more severe form, the symptoms are same as Thalassemia major, but they start later in life. Often they don't need regular blood transfusion.")
add_block("THAL_PED_FAQ2_H3", "Sub-heading.", "Symptoms of Thalassemia Trait")
add_block("THAL_PED_FAQ2_A3", "Answer paragraph.",
    "This does not produce any symptom and is often detected on routine blood tests. The red cell morphology is abnormal but the Hb is usually near normal.")

add_h3("FAQ — Diagnosis")
add_block("THAL_PED_FAQ3_Q", "Accordion question.", "Diagnosis")
add_block("THAL_PED_FAQ3_H1", "Sub-heading.", "Complete Blood Count")
add_block("THAL_PED_FAQ3_A1", "Answer paragraph.",
    "An astute haematologist easily picks up the diagnosis on routine blood tests.")
add_block("THAL_PED_FAQ3_H2", "Sub-heading.", "Hemoglobin Electrophoresis")
add_block("THAL_PED_FAQ3_A2", "Answer paragraph.",
    "This is needed to confirm the diagnosis. HbF is the major component with very little of normal haemoglobin (HbA).")
add_block("THAL_PED_FAQ3_H3", "Sub-heading.", "Genetic Tests")
add_block("THAL_PED_FAQ3_A3", "Answer paragraph.",
    "Genetic tests are done to test for the mutations of Thalassemia. This is needed in the following situations:")
add_label("THAL_PED_FAQ3_LIST")
add_hint("Bullets under the Genetic Tests heading.")
add_bullets([
    "When the child is already transfused and the diagnosis cannot be confirmed by electrophoresis.",
    "When a couple wants to confirm the status of the foetus i.e. antenatal testing.",
])

add_h3("FAQ — Treatment of Thalassemia Major")
add_block("THAL_PED_FAQ4_Q", "Accordion question.", "Treatment of Thalassemia Major")
add_block("THAL_PED_FAQ4_INTRO", "Intro paragraph.", "There are two approaches to the treatment of Thalassemia major:")
add_block("THAL_PED_FAQ4_H1", "Sub-heading.", "1. Transfusion and Chelation (Removal of iron from body)")
add_block("THAL_PED_FAQ4_A1a", "Answer paragraph 1.",
    "This involves regular red cell transfusion to maintain the Hb above 12 gm% and not let it drop below 9.5 gm%. This enables normal growth and development.")
add_block("THAL_PED_FAQ4_A1b", "Answer paragraph 2.",
    "However, such regular transfusion loads our body with excess iron which affects the heart, liver and other organs. There are medicines, both Intravenous and Oral, to remove excess iron from our body. Effective iron chelation helps in maintaining organ function in regularly transfused patients.")
add_block("THAL_PED_FAQ4_SE_INTRO", "Intro line before side-effects bullets.", "Other side effects of transfusion include:")
add_label("THAL_PED_FAQ4_SE_LIST")
add_hint("Bullets of transfusion side effects.")
add_bullets([
    "Transmission of viral infections",
    "Transfusion reactions",
    "Development of antibodies against red cells.",
])
add_block("THAL_PED_FAQ4_A1c", "Answer paragraph 3.",
    "During chelation, the effect on eyes and ears has to be monitored. Odd infections with bacteria like Yersinia can also occur.")
add_block("THAL_PED_FAQ4_H2", "Sub-heading.", "2. Bone Marrow Transplantation (BMT)")
add_block("THAL_PED_FAQ4_A2a", "Answer paragraph 1.",
    "ALLOGENEIC Bone Marrow Transplantation (BMT) is the only curative treatment for this condition. This works by replacement of diseased stem cells with healthy donor stem cells. Once healthy stem cells start producing normal red cells with normal Hb, the patient is cured.")
add_block("THAL_PED_FAQ4_A2b", "Answer paragraph 2.",
    "This doesn't change the genetic composition of the patient, it only changes the genetic composition of the bone marrow.")

add_h3("FAQ — Role of BMT (Pediatric)")
add_block("THAL_PED_FAQ5_Q", "Accordion question.", "Role of BMT — Best Age, Donor, Outcomes")
add_block("THAL_PED_FAQ5_H1", "Sub-heading.", "What is the best age for patients with Thalassemia to undergo Bone Marrow Transplantation?")
add_block("THAL_PED_FAQ5_A1", "Answer.",
    "Best age is between 2 to 5 years. However, successful BMT has been carried out in older children and young adults.")
add_block("THAL_PED_FAQ5_H2", "Sub-heading.", "Who can be the donor?")
add_label("THAL_PED_FAQ5_DONOR_LIST")
add_hint("Donor bullets.")
add_bullets([
    "Matched Family donor is the best option.",
    "A person who has Thalassemia trait can be a donor for a patient with Thalassemia major.",
    "Results with Matched Unrelated Donor are improving.",
    "Recent studies have shown that HAPLOIDENTICAL (Half Matched) FAMILY DONOR might be a viable option for patients with Thalassemia major.",
])
add_block("THAL_PED_FAQ5_H3", "Sub-heading.", "How do we predict the success or failure of BMT in a particular patient?")
add_block("THAL_PED_FAQ5_A3", "Answer.",
    "Prof. Lucarelli from Pessaro, Italy had proposed a scoring system for patients undergoing BMT for Thalassemia. This takes into account the effectiveness of transfusion, chelation and degree of iron overload. The patients are classified as Class I, Class II, Class III.")
add_block("THAL_PED_FAQ5_H4", "Sub-heading.", "Outcome based on Pessaro Classification: Thalassemia free Survival")
add_label("THAL_PED_FAQ5_OUTCOME_LIST")
add_hint("Outcome bullets.")
add_bullets(["Class I: 90%", "Class II: 80%", "Class III: 60%"])
add_block("THAL_PED_FAQ5_H5", "Sub-heading.", "Do patients with Thalassemia Intermedia need BMT?")
add_block("THAL_PED_FAQ5_A5", "Intro line.",
    "Following categories of patients need Bone Marrow Transplantation:")
add_label("THAL_PED_FAQ5_INTERMEDIA_LIST")
add_hint("Bullets under Thalassemia Intermedia question.")
add_bullets([
    "Those with severe symptoms",
    "Those who are transfusion dependent",
    "Those who choose to be cured knowing the risks and benefits",
])

# ----- Adult tab -----
add_h2("Adult Tab")
add_block("THAL_AD_OVERLINE", "Overline text on the adult tab.", "Adult BMT — Dr. Suparno Chakrabarti")
add_block("THAL_AD_TITLE", "Adult tab title.", "Thalassemia and Related Disorders in Adults")

add_h3("FAQ — What is Thalassemia? (Adult)")
add_block("THAL_AD_FAQ1_Q", "Accordion question.", "What is Thalassemia and Types of Thalassemia?")
add_block("THAL_AD_FAQ1_A1", "Same definition as pediatric tab.",
    "Thalassemia is an inherited blood disorder characterised by low haemoglobin (Haemoglobin is an iron rich protein present within the Red Blood Cells of the body. Haemoglobin is responsible for carrying oxygen to all parts of the body). Normal haemoglobin binds iron with alfa and beta globin (protein) chains. Due to genetic mutations in both the genes inherited from the parents, either alfa or beta chains are absent or reduced. This results in precipitation of other globin chains inside the red blood cells, leading to deformity of the red cells. These deformed red cells get destroyed either within the bone marrow (spongy material within the bones) or spleen resulting in anemia (Pallor). This is called Thalassemia Major.")
add_block("THAL_AD_FAQ1_A2", "Paragraph.",
    "When an individual inherits only one abnormal gene, the manifestations are very mild. This is called Thalassemia Minor or Thalassemia Trait.")
add_block("THAL_AD_FAQ1_A3", "Paragraph.",
    "The commonest form of Thalassemia in India is Beta-Thalassemia, where Beta-globin chains are deficient. When instead of absolute absence of one type of globin chain, there is reduction only; the condition is called Thalassemia Intermedia.")
add_block("THAL_AD_FAQ1_A4", "Paragraph.",
    "Other Thalassemia like Syndromes such as HbE, HbD etc. when co-inherited with Beta-thalassemia gene give rise to Thalassemia Intermedia.")

add_h3("FAQ — Symptoms (Adult)")
add_block("THAL_AD_FAQ2_Q", "Accordion question.", "Symptoms")
add_block("THAL_AD_FAQ2_H1", "Sub-heading.", "Symptoms of Thalassemia Major")
add_block("THAL_AD_FAQ2_A1", "Answer (adult version — note 'Thalassemia Facies').",
    "The symptoms appear between 3-7 years age. The child starts becoming listless, irritable and pale. The parents often visit the doctor several times before this is detected. The spleen enlarges in size and the eyes turn yellow from jaundice. The child is unable to grow and the bones of the skull and face thicken. This is often called Thalassemia Facies.")
add_block("THAL_AD_FAQ2_H2", "Sub-heading.", "Symptoms of Thalassemia Intermedia")
add_block("THAL_AD_FAQ2_A2", "Answer (adult version).",
    "This depends on the severity of the affection of globin gene production. In more severe form, the appearance is like Thalassemia Major and the patient needs red cell transfusion every 1-3 months. In less severe forms, there might be only mild anemia and the patient does not need regular transfusion.")
add_block("THAL_AD_FAQ2_H3", "Sub-heading.", "Symptoms of Thalassemia Trait")
add_block("THAL_AD_FAQ2_A3", "Answer (adult version).",
    "This does not produce any symptom and is often detected on routine blood tests. The red cell morphology is like those with iron deficiency anemia.")

add_h3("FAQ — Diagnosis (Adult)")
add_block("THAL_AD_FAQ3_Q", "Accordion question.", "Diagnosis")
add_block("THAL_AD_FAQ3_H1", "Sub-heading.", "Complete Blood Count")
add_block("THAL_AD_FAQ3_A1", "Answer.", "An astute haematologist easily picks up the diagnosis on routine blood tests.")
add_block("THAL_AD_FAQ3_H2", "Sub-heading.", "Hemoglobin Electrophoresis")
add_block("THAL_AD_FAQ3_A2", "Answer.",
    "This is needed to confirm the diagnosis. HbF is the major component with very little of normal haemoglobin i.e. HbA.")
add_block("THAL_AD_FAQ3_H3", "Sub-heading.", "Genetic Tests")
add_block("THAL_AD_FAQ3_A3", "Intro line.",
    "Genetic tests are done to test for the mutations of Thalassemia. This is needed in the following situations:")
add_label("THAL_AD_FAQ3_LIST")
add_hint("Bullets under Genetic Tests.")
add_bullets([
    "When the child is already transfused and the diagnosis cannot be confirmed by electrophoresis.",
    "When a couple wants to confirm the status of the foetus i.e. antenatal testing.",
])

add_h3("FAQ — Treatment of Thalassemia Major (Adult)")
add_block("THAL_AD_FAQ4_Q", "Accordion question.", "Treatment of Thalassemia Major")
add_block("THAL_AD_FAQ4_H1", "Sub-heading.", "1. Transfusion and Chelation (Removal of iron from body)")
add_block("THAL_AD_FAQ4_A1a", "Answer paragraph 1.",
    "This involves regular red cell transfusion to maintain the Hb above 12 gm% and not let it drop below 9 gm%.")
add_block("THAL_AD_FAQ4_A1b", "Answer paragraph 2.",
    "However, such regular transfusion, loads our body with excess iron which affects the heart, liver and other organs. There are medicines, both Intravenous and Oral to remove excess iron from our body. Effective iron chelation helps in maintaining organ function in regularly transfused patients.")
add_block("THAL_AD_FAQ4_SE_INTRO", "Intro line.", "Other side effects of transfusion include:")
add_label("THAL_AD_FAQ4_SE_LIST")
add_hint("Bullets of transfusion side effects.")
add_bullets([
    "Transmission of viral infections",
    "Transfusion reactions",
    "Development of antibodies against red cells.",
])
add_block("THAL_AD_FAQ4_A1c", "Answer paragraph.",
    "During chelation, the effect on eyes and ears has to be monitored. Odd infections with bacteria like Yersinia can also occur.")
add_block("THAL_AD_FAQ4_H2", "Sub-heading.", "2. Bone Marrow Transplantation (BMT)")
add_block("THAL_AD_FAQ4_A2a", "Answer paragraph 1.",
    "ALLOGENEIC Bone Marrow Transplantation (BMT) is the only curative treatment for this condition. This works by replacement of diseased stem cells by normal stem cells from a donor.")
add_block("THAL_AD_FAQ4_A2b", "Answer paragraph 2.",
    "This doesn't change the genetic composition of the patient, it only changes the genetic composition of the blood and immune system.")

add_h3("FAQ — Role of BMT (Adult)")
add_block("THAL_AD_FAQ5_Q", "Accordion question.", "Role of BMT — Best Age, Donor, Outcomes")
add_block("THAL_AD_FAQ5_H1", "Sub-heading.", "What is the best age for patients with Thalassemia to undergo Bone Marrow Transplantation?")
add_block("THAL_AD_FAQ5_A1", "Answer.",
    "Best age is between 2 to 5 years. However, successful BMT has been carried out in older children and young adults.")
add_block("THAL_AD_FAQ5_H2", "Sub-heading.", "Who can be the donor?")
add_label("THAL_AD_FAQ5_DONOR_LIST")
add_hint("Donor bullets.")
add_bullets([
    "Matched Family donor is the best option.",
    "A person who has Thalassemia trait can be a donor for a patient with Thalassemia major.",
    "Results with Matched Unrelated Donor are improving.",
    "Recent studies have shown that HAPLOIDENTICAL (Half Matched) FAMILY DONOR might be a viable option for those who do not have a matched donor.",
])
add_block("THAL_AD_FAQ5_H4", "Sub-heading.", "Outcome based on Pessaro Classification: Thalassemia free Survival")
add_label("THAL_AD_FAQ5_OUTCOME_LIST")
add_hint("Outcome bullets.")
add_bullets(["Class I: 90%", "Class II: 80%", "Class III: 60%"])
add_block("THAL_AD_FAQ5_H5", "Sub-heading.", "Do patients with Thalassemia Intermedia need BMT?")
add_label("THAL_AD_FAQ5_INTERMEDIA_LIST")
add_hint("Bullets answering the question.")
add_bullets([
    "Those with severe symptoms",
    "Those who are transfusion dependent",
    "Those who choose to be cured knowing the risks and benefits",
])

condition_cta("THAL",
    "When to see us for Thalassemia",
    "If your child shows pallor, fatigue, or jaundice — consult our specialists without delay.")

add_page_break()


# ============================================================
# CONDITION: SICKLE CELL
# ============================================================
add_h1("Condition — Sickle Cell Anemia")
add_hint("Sub-page under Conditions. Has Pediatric and Adult tabs.")

add_h2("Pediatric Tab")
add_block("SC_PED_OVERLINE", "Overline.", "Pediatric BMT — Dr. Mahak Agarwal")
add_block("SC_PED_TITLE", "Tab title.", "Sickle Cell Anemia in Children")

add_h3("FAQ — What Is Sickle Cell Anemia?")
add_block("SC_PED_FAQ1_Q", "Accordion question.", "What Is Sickle Cell Anemia?")
add_block("SC_PED_FAQ1_A1", "Paragraph 1.",
    'Sickle Cell Anaemia is the most common form of sickle cell disease (SCD). SCD is a serious disorder in which the body makes sickle-shaped red blood cells. "Sickle-shaped" means that the red blood cells are shaped like a crescent.')
add_block("SC_PED_FAQ1_A2", "Paragraph 2.",
    "Normal red blood cells are disc-shaped and look like doughnuts without holes in the centre. They move easily through the blood vessels. Red blood cells contain an iron-rich protein called haemoglobin. This protein carries oxygen from the lungs to the rest of the body.")
add_block("SC_PED_FAQ1_A3", "Paragraph 3.",
    "Sickle cells contain abnormal haemoglobin called sickle hemoglobin or hemoglobin S. Sickle hemoglobin causes the cells to develop a sickle, or crescent shape.")
add_block("SC_PED_FAQ1_A4", "Paragraph 4.",
    "Sickle cells are stiff and sticky. They tend to block blood flow in the blood vessels of the limbs and organs. Blocked blood flow can cause pain and organ damage. It can also raise the risk for infection.")

add_h3("FAQ — Causes & Who Is at Risk?")
add_block("SC_PED_FAQ2_Q", "Accordion question.", "Causes & Who Is at Risk?")
add_block("SC_PED_FAQ2_A1", "Paragraph 1.",
    "When both parents have a normal gene and an abnormal gene, each child has a 25 percent chance of inheriting two normal genes; a 50 percent chance of inheriting one normal gene and one abnormal gene; and a 25 percent chance of inheriting two abnormal genes.")
add_block("SC_PED_FAQ2_A2", "Paragraph 2.",
    "People who inherit a sickle hemoglobin gene from one parent and a normal gene from the other parent have sickle cell trait. Their bodies make both sickle hemoglobin and normal hemoglobin. People who have sickle cell trait usually have few, if any, symptoms and lead normal lives. However, they can pass the sickle hemoglobin gene to their children.")
add_block("SC_PED_FAQ2_A3", "Paragraph 3.",
    "People who inherit both sickle hemoglobin genes from both parents develop Sickle Cell Disease. Their bodies make mostly sickle hemoglobin.")
add_block("SC_PED_FAQ2_A4", "Paragraph 4.",
    "Sickle cell anemia is most common in people whose families come from Africa, South or Central America (especially Panama), Caribbean islands, Mediterranean countries (such as Turkey, Greece, and Italy), India, and Saudi Arabia.")

add_h3("FAQ — Signs & Symptoms")
add_block("SC_PED_FAQ3_Q", "Accordion question.", "Signs & Symptoms")
add_block("SC_PED_FAQ3_A1", "Intro paragraph.",
    "The signs and symptoms of sickle cell anemia vary. Some people have mild symptoms. Others have very severe symptoms and are often hospitalized for treatment. Sickle cell anemia is present at birth, but many infants don't show any signs until after 4 months of age.")
add_block("SC_PED_FAQ3_A2", "Intro to bullets.",
    "The most common signs and symptoms are linked to anemia and pain. Symptoms may include:")
add_label("SC_PED_FAQ3_LIST")
add_hint("Symptom bullets.")
add_bullets([
    "Severe pain in arms, back, abdomen and legs",
    "Fever",
    "Strokes — one-sided weakness or paralysis in arms, face or legs; or finding difficulty in talking, walking or seeing",
    "Shortness of breath or difficulty in breathing",
    "Chest pain",
    "Severe infections",
    "Blood flow blockage in liver or spleen",
    "Anemia",
    "Delayed growth",
    "Hand & foot syndrome i.e. swelling in the feet and hands",
    "Pale skin",
    "Vision problems",
    "Kidney failure",
    "Pulmonary Hypertension",
])

add_h3("FAQ — Diagnosis")
add_block("SC_PED_FAQ4_Q", "Accordion question.", "Diagnosis")
add_block("SC_PED_FAQ4_H1", "Sub-heading.", "Complete Blood Count and PERIPHERAL Blood Smear")
add_block("SC_PED_FAQ4_A1", "Paragraph.",
    "The characteristic finding of SCD is circulating sickle shaped red cells. Routine examination of the blood count reveals anemia. Reticulocytosis i.e. increase in young red blood cells is almost always seen.")
add_block("SC_PED_FAQ4_H2", "Sub-heading.", "Hemoglobin Electrophoresis")
add_block("SC_PED_FAQ4_A2", "Paragraph.",
    "This is the cornerstone of diagnosis and is used to detect Hemoglobin S. HbA2 and often HbF are increased in SCD.")
add_block("SC_PED_FAQ4_H3", "Sub-heading.", "Prenatal Diagnosis")
add_block("SC_PED_FAQ4_A3", "Paragraph.",
    "May also be detected by examining the DNA from the chorionic villus biopsy or from the cells of amniotic fluid.")

add_h3("FAQ — Treatment")
add_block("SC_PED_FAQ5_Q", "Accordion question.", "Treatment")
add_block("SC_PED_FAQ5_H1", "Sub-heading.", "Supportive Treatment")
add_label("SC_PED_FAQ5_LIST1")
add_hint("Supportive treatment bullets.")
add_bullets([
    "Folic acid",
    "Penicillin prophylaxis",
    "Pneumococcal vaccine.",
    "All the infections should be identified and should be treated early.",
    "Treatment of painful crisis or other complications.",
])
add_block("SC_PED_FAQ5_H2", "Sub-heading.", "Specific Treatment")
add_label("SC_PED_FAQ5_LIST2")
add_hint("Specific treatment bullets (bold term at start of each line).")
add_bullets([
    "Hydroxyurea: Data have shown that it decreases the painful crisis episodes. This drug decreases Neutrophil concentration and stickiness of red cells and increases fetal hemoglobin (HbF).",
    "Red cell Transfusion: Used frequently in Sickle Cell Disease to increase Hb concentration and decrease the concentration of sickle red cells.",
    "Exchange red cell transfusions: This is done to reduce the concentration of sickle cells without increasing the concentration of red cells too much.",
])

add_h3("FAQ — Role of BMT in Sickle Cell Disease")
add_block("SC_PED_FAQ6_Q", "Accordion question.", "Role of BMT in Sickle Cell Disease")
add_block("SC_PED_FAQ6_A1", "Paragraph.",
    "ALLOGENEIC BMT was first carried out in 1980 in a patient with Acute Leukemia who also had Sickle Cell Disease. Since then several centers in the Europe and North America have successfully performed BMT for Sickle Cell Disease.")
add_block("SC_PED_FAQ6_A2", "Paragraph.",
    "This shows that 90% of the most severely affected patients with Sickle Cell Disease are cured with BMT if they have a matched family donor.")
add_block("SC_PED_FAQ6_A3", "Paragraph.",
    "However, very few patients with Sickle Cell Disease have a Matched Family Donor. The Unrelated Donor Transplants are often complicated by severe Graft Versus Host Disease (GVHD).")
add_block("SC_PED_FAQ6_A4", "Paragraph.",
    "BMT in Sickle Cell Disease has produced the best results one can expect in any severe inherited disease.")
add_block("SC_PED_FAQ6_A5", "Paragraph.",
    "Recent studies have shown that Haploidentical BMT can cure the majority of patients with Sickle Cell Disease with nearly 100% chance of finding a donor.")
add_block("SC_PED_FAQ6_H1", "Sub-heading.", "How is Conditioning for BMT done in Sickle Cell Disease?")
add_block("SC_PED_FAQ6_A6", "Paragraph.",
    "Primarily Reduced Intensity conditioning is used in adults and conventional high dose chemotherapy or radiation is the preferred option in children.")
add_block("SC_PED_FAQ6_H2", "Sub-heading.", "Who can be a donor for BMT?")
add_block("SC_PED_FAQ6_A7", "Paragraph.",
    "Although we prefer a matched family donor, a Half matched (Haploidentical) family donor or an unrelated donor provide excellent survival.")
add_block("SC_PED_FAQ6_HIGHLIGHT", "Highlighted box (all caps emphasis).",
    "WE OFFER HAPLOIDENTICAL BMT FOR PATIENTS WITH SICKLE CELL DISEASE LACKING A MATCHED FAMILY DONOR.")

# Adult tab
add_h2("Adult Tab")
add_block("SC_AD_OVERLINE", "Overline.", "Adult BMT — Dr. Suparno Chakrabarti")
add_block("SC_AD_TITLE", "Tab title.", "Sickle Cell Anemia in Adults")

add_h3("FAQ — What Is Sickle Cell Anemia? (Adult)")
add_block("SC_AD_FAQ1_Q", "Accordion question.", "What Is Sickle Cell Anemia?")
add_block("SC_AD_FAQ1_A1", "Paragraph 1.",
    'Sickle Cell Anaemia is the most common form of sickle cell disease (SCD). SCD is a serious disorder in which the body makes sickle-shaped red blood cells. "Sickle-shaped" means that the red blood cells are shaped like a crescent.')
add_block("SC_AD_FAQ1_A2", "Paragraph 2.",
    "Normal red blood cells are disc-shaped and look like doughnuts without holes in the centre. They move easily through the blood vessels. Red blood cells contain an iron-rich protein called haemoglobin. This protein carries oxygen from the lungs to the rest of the body.")
add_block("SC_AD_FAQ1_A3", "Paragraph 3.",
    "Sickle cells contain abnormal haemoglobin called sickle hemoglobin or hemoglobin S. Sickle hemoglobin causes the cells to develop a sickle, or crescent shape. Sickle cells are stiff and sticky. They tend to block blood flow in the blood vessels of the limbs and organs. Blocked blood flow can cause pain and organ damage. It can also raise the risk for infection.")

add_h3("FAQ — Causes, Risk Factors & Symptoms (Adult)")
add_block("SC_AD_FAQ2_Q", "Accordion question.", "Causes, Risk Factors & Symptoms")
add_block("SC_AD_FAQ2_A1", "Paragraph.",
    "When both parents have a normal gene and an abnormal gene, each child has a 25 percent chance of inheriting two normal genes; a 50 percent chance of inheriting one normal gene and one abnormal gene; and a 25 percent chance of inheriting two abnormal genes.")
add_block("SC_AD_FAQ2_A2", "Paragraph.",
    "People who inherit a sickle hemoglobin gene from one parent and a normal gene from the other parent have sickle cell trait. Their bodies make both sickle hemoglobin and normal hemoglobin. People who have sickle cell trait usually have few, if any, symptoms and lead normal lives. However, they can pass the sickle hemoglobin gene to their children.")
add_block("SC_AD_FAQ2_A3", "Paragraph.",
    "People who inherit both sickle hemoglobin genes from both parents, develop Sickle Cell Disease. Their bodies make only sickle hemoglobin.")
add_block("SC_AD_FAQ2_A4", "Paragraph.",
    "Sickle cell anemia is most common in people whose families come from Africa, South or Central America (especially Panama), Caribbean islands, Mediterranean countries (such as Turkey, Greece, and Italy), India, and Saudi Arabia.")
add_block("SC_AD_FAQ2_A5", "Symptoms summary (runs as single paragraph).",
    "Symptoms include: Severe pain in arms, back, abdomen and legs; Fever; Strokes; Shortness of breath or difficulty in breathing; Chest pain; Severe infections; Blood flow blockage in liver or spleen; Anemia; Delayed growth; Hand & foot syndrome; Pale skin; Vision problems; Kidney failure; Pulmonary Hypertension.")

add_h3("FAQ — Diagnosis (Adult)")
add_block("SC_AD_FAQ3_Q", "Accordion question.", "Diagnosis")
add_block("SC_AD_FAQ3_H1", "Sub-heading.", "Complete Blood Count and PERIPHERAL Blood Smear")
add_block("SC_AD_FAQ3_A1", "Paragraph.",
    "The characteristic finding of SCD is circulating sickle shaped red cells. Routine examination of the blood under the microscope may confirm the presence of sickle cells. Reticulocytosis i.e. increase in young red blood cells is almost always seen.")
add_block("SC_AD_FAQ3_H2", "Sub-heading.", "Hemoglobin Electrophoresis")
add_block("SC_AD_FAQ3_A2", "Paragraph.",
    "This is the cornerstone of diagnosis and is used to detect Hemoglobin S. HbA2 and often HbF are increased.")
add_block("SC_AD_FAQ3_H3", "Sub-heading.", "Prenatal Diagnosis")
add_block("SC_AD_FAQ3_A3", "Paragraph.",
    "This may also be detected by examining the DNA from the chorionic villus biopsy or from the cells of amniotic fluid.")

add_h3("FAQ — Treatment & Role of BMT (Adult)")
add_block("SC_AD_FAQ4_Q", "Accordion question.", "Treatment & Role of BMT")
add_block("SC_AD_FAQ4_H1", "Sub-heading.", "Supportive Treatment")
add_label("SC_AD_FAQ4_LIST1")
add_hint("Supportive treatment bullets.")
add_bullets([
    "Folic acid",
    "Penicillin prophylaxis",
    "Pneumococcal vaccine.",
    "All the infections should be identified and should be treated early.",
    "Treatment of painful crisis or other complications.",
])
add_block("SC_AD_FAQ4_H2", "Sub-heading.", "Specific Treatment")
add_label("SC_AD_FAQ4_LIST2")
add_hint("Specific treatment bullets.")
add_bullets([
    "Hydroxyurea: Data have shown that it decreases the painful crisis episodes. This drug decreases Neutrophil count and increases Hemoglobin F concentration.",
    "Red cell Transfusion: Used frequently in Sickle Cell Disease to increase Hb concentration and decrease the concentration of sickle cells.",
    "Exchange red cell transfusions: This is done to reduce the concentration of sickle cells without increasing the concentration of red blood cells.",
])
add_block("SC_AD_FAQ4_H3", "Sub-heading.", "Blood and Marrow Transplant (BMT) for Sickle Cell Disease")
add_block("SC_AD_FAQ4_A1", "Paragraph.",
    "ALLOGENEIC BMT was first carried out in 1980 in a patient with Acute Leukemia who also had Sickle Cell Anemia. The BMT cured both.")
add_block("SC_AD_FAQ4_A2", "Paragraph.",
    "This shows that 90% of the most severely affected patients with Sickle Cell Disease are cured with BMT.")
add_block("SC_AD_FAQ4_A3", "Paragraph.",
    "However, very few patients with Sickle Cell Disease have a Matched Family Donor. The Unrelated Donor search for patients of African or Indian origin is very difficult.")
add_block("SC_AD_FAQ4_A4", "Paragraph.",
    "BMT in Sickle Cell Disease has produced the best results one can expect in any severe inherited disorder. Recent studies have shown that Haploidentical BMT can cure the majority of patients with Sickle Cell Disease.")
add_block("SC_AD_FAQ4_H4", "Sub-heading.", "How is Conditioning for BMT done in Sickle Cell Disease?")
add_block("SC_AD_FAQ4_A5", "Paragraph.",
    "Primarily Reduced Intensity conditioning is used in adults and conventional high dose chemotherapy is used in children.")
add_block("SC_AD_FAQ4_H5", "Sub-heading.", "Who can be a donor for BMT?")
add_block("SC_AD_FAQ4_A6", "Paragraph.",
    "Although we prefer a matched family donor, a Half matched (Haploidentical) family donor or an unrelated donor can also be used.")
add_block("SC_AD_FAQ4_HIGHLIGHT", "Highlighted box.",
    "WE OFFER HAPLOIDENTICAL BMT FOR PATIENTS WITH SICKLE CELL DISEASE LACKING A MATCHED FAMILY DONOR.")

condition_cta("SC",
    "When to see us for Sickle Cell Disease",
    "Frequent pain crises, strokes, or organ damage require specialist BMT evaluation.")

add_page_break()


# ============================================================
# CONDITION: APLASTIC ANEMIA
# ============================================================
add_h1("Condition — Aplastic Anemia")
add_hint("Sub-page under Conditions. Has Pediatric and Adult tabs.")

add_h2("Pediatric Tab")
add_block("AA_PED_OVERLINE", "Overline.", "Pediatric BMT — Dr. Suparno Chakrabarti")
add_block("AA_PED_TITLE", "Tab title.", "Aplastic Anemia in Children")

add_h3("FAQ — What is Aplastic Anemia?")
add_block("AA_PED_FAQ1_Q", "Accordion question.", "What is Aplastic Anemia?")
add_block("AA_PED_FAQ1_A1", "Paragraph 1.",
    "The blood cells have a limited lifespan. Red blood cells survive 3 months, white cells survive for 24 hours and platelets survive for 5 days. The bone marrow is an amazing organ which produces billions of blood cells every day to keep up with the requirements of our body. When the bone marrow fails, the cell production slows down and often stops. This is called aplastic anemia.")
add_block("AA_PED_FAQ1_A2", "Paragraph 2.",
    "The incidence of Aplastic anemia in Asian countries is 5 times higher than in Europe or USA.")

add_h3("FAQ — Causes")
add_block("AA_PED_FAQ2_Q", "Accordion question.", "Causes")
add_block("AA_PED_FAQ2_INTRO", "Intro.",
    "Aplastic anemia results from damage to the mother cells or stem cells which are no longer able to produce normal blood cells. This can happen due to the following causes:")
add_label("AA_PED_FAQ2_LIST")
add_hint("Causes bullets.")
add_bullets([
    "Chemicals and toxins such as benzene and pesticides",
    "Infections such as hepatitis",
    "Drugs such as arsenic, gold, painkillers and some antibiotics",
    "Inherited conditions such as Fanconi's Anemia",
    "Most commonly, the cause is not known and thought to be due to an immune mediated attack by the body on its own stem cells.",
])

add_h3("FAQ — Diagnosis & Classification")
add_block("AA_PED_FAQ3_Q", "Accordion question.", "Diagnosis & Classification")
add_block("AA_PED_FAQ3_A1", "Paragraph 1.",
    "The diagnosis of aplastic anemia needs awareness of the condition and a strong index of suspicion. The patient can present with fever due to infection, bleeding and bruising due to low platelet count and fatigue due to anemia.")
add_block("AA_PED_FAQ3_A2", "Paragraph 2.",
    "A routine blood test shall show at least two of the three cell lines of red cells, white cells and platelets being low. A marrow biopsy is mandatory to confirm the diagnosis and rule out other causes of marrow failure like leukemia.")
add_block("AA_PED_FAQ3_A3", "Paragraph 3.",
    "It is important to look for a causative factor and stop any drug or exposure to chemical that might have caused it.")
add_block("AA_PED_FAQ3_A4", "Paragraph 4.",
    "The investigations include blood tests to look for infections, autoimmune diseases like lupus and chromosomal breakage study to rule out Fanconi's anemia.")
add_block("AA_PED_FAQ3_H1", "Sub-heading.", "Classification")
add_label("AA_PED_FAQ3_LIST")
add_hint("Classification bullets (bold term at start).")
add_bullets([
    "Severe: neutrophil count less than 500 cells/microliter (normal 1500-6000/microliter)",
    "Very Severe: neutrophil count less than 200 cells/microliter",
    "Nonsevere: none of the above",
])

add_h3("FAQ — Treatment & Role of BMT")
add_block("AA_PED_FAQ4_Q", "Accordion question.", "Treatment & Role of BMT")
add_block("AA_PED_FAQ4_H1", "Sub-heading.", "Supportive Therapy")
add_block("AA_PED_FAQ4_A1", "Paragraph.",
    "We cannot survive without functioning blood cells. Immediate treatment involves transfusion of red cells and platelets and treatment of infections. However, the patient must be referred to a center where curative treatment can be given as soon as possible. Repeated blood transfusions reduce the success rate of BMT.")
add_block("AA_PED_FAQ4_H2", "Sub-heading.", "Immunosuppressive Treatment")
add_block("AA_PED_FAQ4_A2", "Paragraph.",
    "Antithymocyte globulin, a drug produced from horse or rabbit by injecting human lymphocytes in them, is the main drug. It suppresses the immune system and allows the marrow to recover. This drug is given with another oral medicine called Cyclosporine. About 60-70% of the patients improve with this treatment. However, it takes 3-6 months for the marrow to recover and many patients may relapse or develop other marrow disorders in the future.")
add_block("AA_PED_FAQ4_H3", "Sub-heading.", "Blood and Marrow Transplantation (BMT)")
add_block("AA_PED_FAQ4_A3", "Paragraph.",
    "BMT cures 90% of the patients with aplastic anemia if carried out early and from a matched family donor. BMT is the treatment of choice for young patients with a matched sibling donor.")
add_block("AA_PED_FAQ4_A4", "Paragraph.",
    "For patients without a matched family donor, HAPLOIDENTICAL DONOR TRANSPLANTATION is now giving results as good as a matched family donor. At BLOODS-R-US, we offer the best survival for patients with aplastic anemia through BMT.")

add_h2("Adult Tab")
add_block("AA_AD_OVERLINE", "Overline.", "Adult BMT — Dr. Suparno Chakrabarti")
add_block("AA_AD_TITLE", "Tab title.", "Aplastic Anemia in Adults")

add_h3("FAQ — What is Aplastic Anemia? (Adult)")
add_block("AA_AD_FAQ1_Q", "Accordion question.", "What is Aplastic Anemia?")
add_block("AA_AD_FAQ1_A1", "Paragraph 1.",
    "The blood cells have a limited lifespan. Red blood cells survive 3 months, white cells survive for 24 hours and platelets survive for 5 days. The bone marrow is an amazing organ which produces billions of blood cells every day to keep up with the requirements of our body. When the bone marrow fails, the cell production slows down and often stops. This is called aplastic anemia.")
add_block("AA_AD_FAQ1_A2", "Paragraph 2.",
    "The incidence of Aplastic anemia in Asian countries is 5 times higher than in Europe or USA.")

add_h3("FAQ — Causes, Diagnosis & Classification (Adult)")
add_block("AA_AD_FAQ2_Q", "Accordion question.", "Causes, Diagnosis & Classification")
add_block("AA_AD_FAQ2_A1", "Paragraph (causes combined).",
    "Aplastic anemia results from damage to the mother cells or stem cells which are no longer able to produce normal blood cells. Causes include: Chemicals and toxins such as benzene and pesticides; Infections such as hepatitis; Drugs such as arsenic, gold, painkillers and some antibiotics; Inherited conditions such as Fanconi's Anemia. Most commonly, the cause is not known and thought to be due to an immune mediated attack by the body on its own stem cells.")
add_block("AA_AD_FAQ2_H1", "Sub-heading.", "Diagnosis")
add_label("AA_AD_FAQ2_LIST1")
add_hint("Diagnosis bullets.")
add_bullets([
    "The diagnosis of aplastic anemia needs awareness of the condition and a strong index of suspicion. The patient can present with fever, weakness or bleeding.",
    "A routine blood test shall show at least two of the three cell lines of red cells, white cells and platelets being low.",
    "A bone marrow examination is essential for confirmation of the diagnosis. It shows absence of normal blood cells and replacement with fat cells.",
    "It is important to look for a causative factor and stop any drug or exposure to chemical that might have caused the condition.",
    "The investigations include blood tests to look for infections, autoimmune diseases like lupus and chromosomal abnormalities.",
])
add_block("AA_AD_FAQ2_H2", "Sub-heading.", "Classification")
add_label("AA_AD_FAQ2_LIST2")
add_hint("Classification bullets.")
add_bullets([
    "Severe: neutrophil count less than 500 cells/microliter (normal 1500-6000/microliter)",
    "Very Severe: neutrophil count less than 200 cells/microliter",
    "Nonsevere: none of the above",
])

add_h3("FAQ — Treatment & Role of BMT (Adult)")
add_block("AA_AD_FAQ3_Q", "Accordion question.", "Treatment & Role of BMT")
add_block("AA_AD_FAQ3_H1", "Sub-heading.", "Supportive Therapy")
add_block("AA_AD_FAQ3_A1", "Paragraph.",
    "We cannot survive without functioning blood cells. Immediate treatment involves transfusion of red cells and platelets. Infections have to be treated with high-end antibiotics.")
add_block("AA_AD_FAQ3_H2", "Sub-heading.", "Immunosuppressive Treatment")
add_block("AA_AD_FAQ3_A2", "Paragraph.",
    "Antithymocyte globulin, a drug produced from horse or rabbit by injecting human lymphocytes in them, is the mainstay of treatment. This drug is given for 4 to 5 days and it takes 3 to 4 months for blood counts to recover. This treatment produces response in 60 to 70% of the patients.")
add_block("AA_AD_FAQ3_H3", "Sub-heading.", "Blood and Marrow Transplantation (BMT)")
add_block("AA_AD_FAQ3_A3", "Paragraph.",
    "BMT cures 90% of the patients with aplastic anemia if carried out early and from a matched family donor. Other options like Haploidentical BMT or Unrelated donor BMT are also showing promising results.")

condition_cta("AA",
    "When to see us for Aplastic Anemia",
    "Prompt evaluation is critical — repeated transfusions reduce the success rate of BMT.")

add_page_break()


# ============================================================
# CONDITION: ALL (LEUKEMIA)
# ============================================================
add_h1("Condition — Acute Lymphoblastic Leukemia (ALL)")
add_hint("Sub-page under Conditions. Has Pediatric and Adult tabs.")

add_h2("Pediatric Tab")
add_block("ALL_PED_OVERLINE", "Overline.", "Pediatric BMT — Dr. Suparno Chakrabarti")
add_block("ALL_PED_TITLE", "Tab title.", "Acute Lymphoblastic Leukemia (ALL) in Children")

add_h3("FAQ — What is ALL in Children?")
add_block("ALL_PED_FAQ1_Q", "Accordion question.", "What is ALL in Children?")
add_block("ALL_PED_FAQ1_A1", "Paragraph 1.",
    "Acute Lymphoblastic leukemia (ALL), is a cancer that starts from white blood cells called lymphocytes in the bone marrow (the soft inner part of the bones, where new blood cells are generated).")
add_block("ALL_PED_FAQ1_A2", "Paragraph 2.",
    "The term \"acute\" means that the leukemia can progress quickly, and if not treated, would probably be fatal within a few months. 'Lymphoblastic' means it develops from early (immature) forms of lymphocytes, a type of white blood cell.")
add_block("ALL_PED_FAQ1_A3", "Paragraph 3.",
    "Acute Lymphoblastic Leukemia is the commonest childhood cancer. This arises from abnormalities in the precursors of normal lymphocytes which are otherwise a normal part of our immune system necessary to fight infections. These abnormal cells rapidly fill up the bone marrow space suppressing the normal cells.")
add_block("ALL_PED_FAQ1_A4", "Paragraph 4.",
    "In most cases, the disease is characterised by certain abnormalities in the chromosomes which promote abnormal growth of certain lymphoid cells rather than maturation to normal lymphocytes.")

add_h3("FAQ — Causes & Symptoms")
add_block("ALL_PED_FAQ2_Q", "Accordion question.", "Causes & Symptoms")
add_block("ALL_PED_FAQ2_A1", "Paragraph.",
    "This is not known in most cases. However, exposure to radiation, exposure to high voltage electric wires and chemical exposures are few known risk factors.")
add_block("ALL_PED_FAQ2_H1", "Sub-heading.", "Symptoms")
add_label("ALL_PED_FAQ2_LIST")
add_hint("Symptom bullets.")
add_bullets([
    "General weakness", "Fatigue", "High temperature (fever)", "Weight loss", "Frequent infections",
    "Bruising easily or with no obvious cause", "Bleeding from the gums or nose",
    "A fine rash of dark red spots (called purpura)", "Blood in urine or stools",
    "Pain in the bones or joints", "Breathlessness", "Swollen lymph glands",
    "Enlarged liver or spleen",
])

add_h3("FAQ — Diagnosis & Classification")
add_block("ALL_PED_FAQ3_Q", "Accordion question.", "Diagnosis & Classification")
add_block("ALL_PED_FAQ3_H1", "Sub-heading.", "Complete Blood Count")
add_block("ALL_PED_FAQ3_A1", "Paragraph.",
    "The characteristic finding of Acute Lymphoblastic Leukemia is high white blood cell count with low haemoglobin and platelets. Routine examination of the blood under the microscope may confirm the presence of abnormal cells (blasts).")
add_block("ALL_PED_FAQ3_H2", "Sub-heading.", "Bone Marrow Examination")
add_block("ALL_PED_FAQ3_A2", "Paragraph.",
    "This is necessary for confirmation of the diagnosis. There must be more than 30% of abnormal cells (blasts) in the marrow for it to be called Acute Lymphoblastic Leukemia.")
add_block("ALL_PED_FAQ3_H3", "Sub-heading.", "Flow Cytometry")
add_block("ALL_PED_FAQ3_A3", "Paragraph.",
    "ALL has two broad subtypes- B cell Acute Lymphoblastic Leukemia and T cell Acute Lymphoblastic Leukemia. There are further subtypes in the individual categories. This has got implications on the choice of treatment and the likely outcome.")
add_block("ALL_PED_FAQ3_H4", "Sub-heading.", "Cytogenetics")
add_block("ALL_PED_FAQ3_A4", "Paragraph.",
    "This is testing for abnormalities in the chromosomes. Certain genes are brought next to each other by a process called 'translocation' leading to uninterrupted growth and proliferation of white blood cell precursors in the bone marrow (e.g., Philadelphia Chromosome). Both are normal genes, but when they interact they produce an abnormal protein which drives the cancer. The absolute increase (Hyperdiploid) or decrease (Hypodiploid) in the number of chromosomes in the cancer cells are also features of ALL. Certain changes in sequence of the genes called mutations are also seen which have an impact on outcome of treatment.")
add_block("ALL_PED_FAQ3_H5", "Sub-heading.", "Classification")
add_block("ALL_PED_FAQ3_A5", "Paragraph.",
    "Unlike cancers of solid organs, Acute Lymphoblastic Leukemia and other blood cancers are not staged by the extent of involvement. Certain features related to the patient and the disease help in classifying the patients as: Standard-Risk or Good-Risk, Intermediate-Risk and High-Risk.")
add_label("ALL_PED_FAQ3_LIST")
add_hint("Classification factor bullets.")
add_bullets([
    "Age: if the patient is 2–10 years old, it is considered favourable",
    "Cytogenetics: Certain chromosomal abnormalities such as increase in number of chromosomes (Hyperploidy) and translocation are features of a favourable outcome.",
    "Response to Treatment: Response to steroids after 7 days is probably a very good indicator of how the disease is going to respond to further treatment.",
])
add_block("ALL_PED_FAQ3_H6", "Sub-heading.", "Minimal Residual Disease (MRD)")
add_block("ALL_PED_FAQ3_A6", "Paragraph.",
    "Minimal Residual Disease (MRD) means patient still has a minimum disease and has not fully recovered. Response to treatment is conventionally checked with a bone marrow morphology which establishes COMPLETE REMISSION (CR). However, CR means that the number of leukemia cells in the body are less than 10⁹ or 1000 million cells. At this level, they cannot be seen under the microscope.")
add_block("ALL_PED_FAQ3_A7", "Paragraph.",
    "MRD is the detection of cells below 1%, which can range from 1 in 1000 to 1 in a million cells. This enables us to check for response to treatment with a much higher sensitivity than possible under the microscope.")

add_h3("FAQ — Treatment")
add_block("ALL_PED_FAQ4_Q", "Accordion question.", "Treatment")
add_block("ALL_PED_FAQ4_H1", "Sub-heading.", "Treatment Phases")
add_label("ALL_PED_FAQ4_LIST1")
add_hint("Treatment phase bullets.")
add_bullets([
    "Induction Therapy: The first phase of treatment to kill most of the leukemia cells in the blood and bone marrow and restore normal blood cell production.",
    "Consolidation Therapy: Also known as post-remission therapy. Aimed at destroying any remaining leukemia cells in the body.",
    "Maintenance Therapy: Helps in preventing leukemia cells from regrowing. Treatments used in this stage are often given at lower doses over a long period of time.",
    "Preventive treatment to the spinal cord: Patients suffering from acute lymphocytic leukemia may also receive treatment to kill leukemia cells located in the central nervous system. In this treatment, chemotherapy drugs are injected directly into the fluid that surrounds the spinal cord.",
    "Radiation Therapy: Radiation therapy uses high-powered beams, such as X-rays, to kill cancer cells. If the cancer cells have spread to the central nervous system, your doctor may recommend radiation therapy.",
])
add_block("ALL_PED_FAQ4_NOTE", "Paragraph.",
    "Depending on patient's evaluation, the phases of treatment for acute lymphocytic leukemia can span up to 2-3 years.")
add_block("ALL_PED_FAQ4_H2", "Sub-heading.", "Role of Chemotherapy")
add_label("ALL_PED_FAQ4_LIST2")
add_hint("Chemotherapy outcome bullets.")
add_bullets([
    "90% of the children with Standard Risk Acute Lymphoblastic Leukemia are cured with Chemotherapy alone.",
    "60-70% of children with Intermediate Risk Acute Lymphoblastic Leukemia are cured with Chemotherapy alone.",
    "40% of children with High Risk Acute Lymphoblastic Leukemia are cured with Chemotherapy alone.",
])
add_block("ALL_PED_FAQ4_A1", "Paragraph.",
    "ALL has a tendency to spread to the brain. Hence, the treatment involves giving intrathecal chemotherapy periodically through a Lumbar Puncture. Intravenous chemotherapy is also given over 6 months, followed by oral drugs for 18-30 months. Boys are treated for a longer duration than girls.")
add_block("ALL_PED_FAQ4_H3", "Sub-heading.", "Adolescents and Young Adults (AYA)")
add_block("ALL_PED_FAQ4_A2", "Paragraph.",
    "Patients between ages 16 and 25 are called Adolescents and Young Adults (AYA). The results in this age group are slightly inferior than in younger children.")

add_h3("FAQ — Role of BMT in ALL")
add_block("ALL_PED_FAQ5_Q", "Accordion question.", "Role of BMT in ALL")
add_block("ALL_PED_FAQ5_H1", "Sub-heading.", "When is BMT needed?")
add_label("ALL_PED_FAQ5_LIST1")
add_hint("Indications bullets.")
add_bullets([
    "Most children with Standard or Intermediate Risk Acute Lymphoblastic Leukemia are cured with chemotherapy and do not need a BMT.",
    "Most children with High Risk Acute Lymphoblastic Leukemia might need a BMT.",
    "Most Children with a positive MRD after intensive chemotherapy would benefit from a BMT.",
])
add_block("ALL_PED_FAQ5_H2", "Sub-heading.", "Conditioning")
add_block("ALL_PED_FAQ5_A1", "Paragraph.", "Total Body Radiation is an important component of conditioning for BMT.")
add_block("ALL_PED_FAQ5_H3", "Sub-heading.", "Who can be a donor?")
add_block("ALL_PED_FAQ5_A2", "Paragraph.",
    "Although we prefer a matched family donor, a Half matched (Haploidentical) family donor or an unrelated donor can also be used.")
add_block("ALL_PED_FAQ5_H4", "Sub-heading.", "Results of BMT")
add_label("ALL_PED_FAQ5_LIST2")
add_hint("Outcomes bullets.")
add_bullets([
    "BMT reduces the risk of relapse i.e. recurrence of disease by 50% compared to chemotherapy.",
    "If the child is MRD negative before BMT, the risk of relapse is negligible.",
    "BMT is the only cure for patients who relapse early after chemotherapy.",
])

add_h2("Adult Tab")
add_block("ALL_AD_OVERLINE", "Overline.", "Adult BMT — Dr. Suparno Chakrabarti")
add_block("ALL_AD_TITLE", "Tab title.", "Acute Lymphoblastic Leukemia (ALL) in Adults")

add_h3("FAQ — What is ALL? Causes & Symptoms (Adult)")
add_block("ALL_AD_FAQ1_Q", "Accordion question.", "What is ALL? Causes & Symptoms")
add_block("ALL_AD_FAQ1_A1", "Paragraph 1.",
    "Acute Lymphoblastic leukemia (ALL), is a cancer that starts from white blood cells called lymphocytes in the bone marrow (the soft inner part of the bones, where new blood cells are generated). The term \"acute\" means that the leukemia can progress quickly, and if not treated, would probably be fatal within a few months. 'Lymphoblastic' means it develops from early (immature) forms of lymphocytes, a type of white blood cell.")
add_block("ALL_AD_FAQ1_A2", "Paragraph 2.",
    "Acute Lymphoblastic Leukemia is the commonest childhood cancer. This arises from abnormalities in the precursors of normal lymphocytes which are otherwise a normal part of our immune system necessary to fight infections. These abnormal cells rapidly fill up the bone marrow space suppressing the normal cells. In most cases, the disease is characterised by certain abnormalities in the chromosomes which promote abnormal growth of certain lymphoid cells rather than maturation to normal lymphocytes.")
add_block("ALL_AD_FAQ1_A3", "Paragraph 3.",
    "This is not known in most cases. However, exposure to radiation, exposure to high voltage electric wires and chemical exposures are few known risk factors.")
add_block("ALL_AD_FAQ1_H1", "Sub-heading.", "Symptoms")
add_block("ALL_AD_FAQ1_LIST", "Symptoms listed as one paragraph on the site.",
    "General weakness; Fatigue; High temperature (fever); Weight loss; Frequent infections; Bruising easily or with no obvious cause; Bleeding from the gums or nose; A fine rash of dark red spots (called purpura); Blood in urine or stools; Pain in the bones or joints; Breathlessness; Swollen lymph glands; Enlarged liver or spleen")

add_h3("FAQ — Diagnosis, Classification & MRD (Adult)")
add_block("ALL_AD_FAQ2_Q", "Accordion question.", "Diagnosis, Classification & MRD")
add_block("ALL_AD_FAQ2_A1", "Paragraph.",
    "Complete Blood Count: The characteristic finding of Acute Lymphoblastic Leukemia is high white blood cell count with low haemoglobin and platelets.")
add_block("ALL_AD_FAQ2_A2", "Paragraph.",
    "Bone Marrow Examination: There must be more than 30% of abnormal cells (blasts) in the marrow.")
add_block("ALL_AD_FAQ2_A3", "Paragraph.",
    "Flow Cytometry: ALL has two broad subtypes- B cell and T cell ALL.")
add_block("ALL_AD_FAQ2_A4", "Paragraph.",
    "Cytogenetics: Testing for chromosomal abnormalities, including the Philadelphia Chromosome (translocation between chromosomes 9 and 22). The absolute increase (Hyperdiploid) or decrease (Hypodiploid) in the number of chromosomes are also features of ALL.")
add_block("ALL_AD_FAQ2_A5", "Paragraph.",
    "Classification: Standard-Risk or Good-Risk, Intermediate-Risk and High-Risk, based on age, cytogenetics, and response to treatment.")
add_block("ALL_AD_FAQ2_A6", "Paragraph.",
    "Minimal Residual Disease (MRD): Detection of cells below 1%, ranging from 1 in 1000 to 1 in a million cells — enables checking for response to treatment with much higher sensitivity.")

add_h3("FAQ — How is ALL in Adults Different? (Adult)")
add_block("ALL_AD_FAQ3_Q", "Accordion question.", "How is ALL in Adults Different? Treatment & BMT")
add_block("ALL_AD_FAQ3_H1", "Sub-heading.", "How is ALL in Adults Different from that in Children?")
add_label("ALL_AD_FAQ3_LIST1")
add_hint("Difference bullets.")
add_bullets([
    "ALL is less common in adults",
    "Certain bad CYTOGENETICS such as Philadelphia Chromosome (translocation between chromosomes 9 and 22) are more common in adults.",
    "Adults do not tolerate intensive chemotherapy as well as children",
    "Only 30% of adults are cured by chemotherapy alone.",
])
add_block("ALL_AD_FAQ3_H2", "Sub-heading.", "Treatment")
add_block("ALL_AD_FAQ3_A1", "Paragraph.",
    "Treatment falls into phases: Induction, Consolidation, Maintenance, CNS prophylaxis, and Radiation if needed. Can span 2-3 years.")
add_block("ALL_AD_FAQ3_A2", "Paragraph.",
    "ALL has a tendency to spread to the brain. Hence, the treatment involves giving intrathecal chemotherapy periodically through a Lumbar Puncture. Intravenous chemotherapy is also given over 6 months, followed by oral drugs for 18-30 months. Boys are treated for a longer duration than girls.")
add_block("ALL_AD_FAQ3_H3", "Sub-heading.", "Bone Marrow Transplantation (BMT) for Adults with ALL")
add_label("ALL_AD_FAQ3_LIST2")
add_hint("BMT bullets.")
add_bullets([
    "Most adults with ALL benefit from a BMT at the first remission",
    "Most adults beyond the age of 40 years cannot tolerate high doses of radiotherapy or chemotherapy and they are often treated with Reduced Intensity Conditioning BMT.",
    "Once the disease recurs, the chance of achieving a second remission is much lower than in children.",
])
add_block("ALL_AD_FAQ3_A3", "Paragraph.", "Total Body Radiation is an important component of conditioning for BMT.")
add_block("ALL_AD_FAQ3_A4", "Paragraph.",
    "BMT reduces the risk of relapse by 50% compared to chemotherapy. BMT is the only cure for patients who relapse early after chemotherapy.")

condition_cta("ALL",
    "When to see us for ALL",
    "High-risk ALL and MRD-positive cases benefit most from early BMT evaluation.")

add_page_break()


# ============================================================
# CONDITION: AML
# ============================================================
add_h1("Condition — Acute Myeloid Leukemia (AML)")
add_hint("Sub-page under Conditions. Single-tab page (no Pediatric / Adult split).")

add_block("AML_OVERLINE", "Overline.", "Pediatric BMT — Dr. Mahak Agarwal")
add_block("AML_TITLE", "Page title.", "Acute Myeloid Leukemia (AML)")

add_h3("FAQ — What is Acute Myeloid Leukemia?")
add_block("AML_FAQ1_Q", "Accordion question.", "What is Acute Myeloid Leukemia?")
add_block("AML_FAQ1_A", "Answer.",
    "Acute Myeloid Leukemia (AML) is a disorder of the process that normally produces neutrophils, a type of white blood cell. These abnormal cells rapidly fill up the bone marrow space, suppressing the normal cells. AML may sometimes be called Acute Myelogenous Leukemia, Acute Myelocytic Leukemia, and or Acute Nonlymphocytic Leukemia. Unlike chronic leukemia, acute leukemia develops quickly and generally needs immediate treatment. Acute Myeloid Leukemia is the most common type of acute leukemia which affects adults and chances of getting Acute Myeloid Leukemia increases with age.")

add_h3("FAQ — Causes & Risk Factors")
add_block("AML_FAQ2_Q", "Accordion question.", "Causes & Risk Factors")
add_block("AML_FAQ2_INTRO", "Intro paragraph.",
    "Acute Myeloid Leukemia is caused by damage to the DNA of developing cells in bone marrow. Though exact cause is not known, several risk factors have been identified:")
add_label("AML_FAQ2_LIST")
add_hint("Risk factor bullets.")
add_bullets([
    "Exposure to the chemical benzene",
    "Being male",
    "Smoking, especially after age of 60",
    "Treatment with chemotherapy or radiation therapy in the past",
    "Treatment for childhood acute lymphoblastic leukemia (ALL) in the past",
    "Being exposed to radiation from an atomic bomb",
    "History of a blood disorder such as myelodysplastic syndrome",
    "Weakened immune system due to an organ transplant",
    "Certain genetic disorders, including Down's syndrome and Fanconi's anaemia",
])

add_h3("FAQ — Symptoms")
add_block("AML_FAQ3_Q", "Accordion question.", "Symptoms")
add_block("AML_FAQ3_A", "Paragraph.",
    "The symptoms are usually similar to ALL. In certain types of AML, like Acute Monoblastic Leukemia, or Acute Myelomonocytic Leukemia, one can have symptoms of gum swelling or skin nodules.")
add_block("AML_FAQ3_LIST", "Symptoms listed (single item on site).",
    "Fever; Fatigue; Weight loss or loss of appetite; Shortness of breath; Anaemia; Easy bruising; Severe Bleeding; Petechiae (flat, pin-head sized spots under the skin caused by bleeding); Bone and joint pain; Persistent or frequent infections; Abnormal swelling of gums around teeth")

add_h3("FAQ — Diagnosis & Classification")
add_block("AML_FAQ4_Q", "Accordion question.", "Diagnosis & Classification")
add_block("AML_FAQ4_H1", "Sub-heading.", "Complete Blood Count")
add_block("AML_FAQ4_A1", "Paragraph.",
    "The characteristic finding of AML is high white blood cell count with low haemoglobin and platelet count.")
add_block("AML_FAQ4_H2", "Sub-heading.", "Peripheral Blood Smear")
add_block("AML_FAQ4_A2", "Paragraph.",
    "Presumptive diagnosis of AML can be made via examination of the peripheral blood smear when there are more than 20% blasts.")
add_block("AML_FAQ4_H3", "Sub-heading.", "Bone Marrow Examination")
add_block("AML_FAQ4_A3", "Paragraph.",
    "This is necessary for confirmation of the diagnosis. There must be more than 20% of abnormal cells (blasts) in the marrow for it to be called Acute Myeloid Leukemia. However, in certain cases of Acute Myeloid Leukemia, presence of the characteristic abnormality in chromosome is enough to call it Acute Myeloid Leukemia even with fewer than 20% blasts in the marrow.")
add_block("AML_FAQ4_H4", "Sub-heading.", "Flow Cytometry")
add_block("AML_FAQ4_A4", "Paragraph.",
    "AML has seven subtypes (AML-M1 ---- AML-M7). There are further subtypes in the individual categories which are confirmed via flow cytometry.")
add_block("AML_FAQ4_H5", "Sub-heading.", "Cytogenetics")
add_block("AML_FAQ4_A5", "Paragraph.",
    "This is testing for abnormalities in the chromosomes. Certain genes are brought next to each other by a process called 'translocation' leading to uninterrupted growth and proliferation of white blood cell precursors in the bone marrow. Missing pair of a chromosome in positions 5, 7, 8 is also a feature of AML. This is called Monosomy. Certain changes in sequence of the genes called mutations are also seen which have an impact on outcome of treatment.")
add_block("AML_FAQ4_H6", "Sub-heading.", "Classification")
add_block("AML_FAQ4_A6", "Paragraph.",
    "Unlike cancers of solid organs, ALL and other blood cancers are not staged by the extent of involvement of different organs. They are classified into: Standard-Risk or Good-Risk, Intermediate-Risk and High-Risk. This is predominantly based on the abnormality of chromosomes or genes.")
add_block("AML_FAQ4_H7", "Sub-heading.", "Minimal Residual Disease (MRD)")
add_block("AML_FAQ4_A7", "Paragraph 1.",
    "Minimal Residual Disease (MRD) means patient still has a minimum disease and has not fully recovered. Response to treatment is conventionally checked with a bone marrow morphology which establishes COMPLETE REMISSION (CR). However, CR means that the number of leukemia cells in the body are less than 10⁹ or 1000 million cells. At this level, they cannot be seen under the microscope.")
add_block("AML_FAQ4_A8", "Paragraph 2.",
    "MRD is the detection of cells below 1%, which can range from 1 in 1000 to 1 in a million cells. This enables us to check for response to treatment with a much higher sensitivity than possible under the microscope.")

add_h3("FAQ — Treatment")
add_block("AML_FAQ5_Q", "Accordion question.", "Treatment")
add_block("AML_FAQ5_H1", "Sub-heading.", "Chemotherapy")
add_block("AML_FAQ5_A1", "Paragraph.",
    "Chemotherapy for AML is divided in two phases, INDUCTION and CONSOLIDATION. About 60-70% of patients achieve a complete remission. The chances of a sustained remission or cure are as follows:")
add_label("AML_FAQ5_LIST1")
add_hint("Cure rate bullets.")
add_bullets(["Good Risk: 50-70%", "Intermediate Risk: 30%", "High Risk: Less than 10%"])
add_block("AML_FAQ5_H2", "Sub-heading.", "What is different about treatment of APML or AML-M3?")
add_block("AML_FAQ5_A2", "Paragraph.",
    "This is a unique condition which responds to high doses of a form of Vitamin A called ALL-TRANS RETINOIC ACID (ATRA) and ARSENIC TRIOXIDE (ATO). This condition carries a 90% cure rate.")
add_block("AML_FAQ5_H3", "Sub-heading.", "How is AML in Children Different from that in Adults?")
add_label("AML_FAQ5_LIST2")
add_hint("Difference bullets.")
add_bullets([
    "AML is less common in children; particularly APML",
    "Results of Chemotherapy are 10% better in children.",
    "Children less commonly have BAD chromosomes.",
])

add_h3("FAQ — Role of BMT in AML")
add_block("AML_FAQ6_Q", "Accordion question.", "Role of BMT in AML")
add_block("AML_FAQ6_H1", "Sub-heading.", "When is BMT needed for AML?")
add_block("AML_FAQ6_A1", "Paragraph.",
    "Chemotherapy is curative in only one-third of patients with AML. All patients with AML apart from the ones with good risk genetics or APML should undergo BMT. BMT is best done when the disease is in first CR. Very few patients who RELAPSE or do not respond to treatment are cured.")
add_block("AML_FAQ6_H2", "Sub-heading.", "Conditioning")
add_block("AML_FAQ6_A2", "Paragraph.", "High to Moderate dose of chemotherapy is generally used in conditioning for BMT.")
add_block("AML_FAQ6_H3", "Sub-heading.", "Who can be a donor?")
add_block("AML_FAQ6_A3", "Paragraph.",
    "Although we prefer a matched family donor, a Half matched (Haploidentical) family donor or an unrelated donor can also be used. However, a HAPLOIDENTICAL DONOR who has Natural Killer Cell mismatch with the patient provides the best cure for AML via Graft versus Leukemia (GVL) effect.")
add_block("AML_FAQ6_H4", "Sub-heading.", "Results of BMT in AML")
add_label("AML_FAQ6_LIST")
add_hint("Results bullets.")
add_bullets([
    "BMT reduces the risk of relapse i.e., recurrence of disease by 80% compared to chemotherapy",
    "If the patient is MRD negative before BMT, the risk of relapse is low",
    "BMT is the only cure for patients, who relapse after chemotherapy",
])

condition_cta("AML",
    "When to see us for AML",
    "Most AML patients (except good-risk/APML) should be evaluated for BMT at first remission.")

add_page_break()


# ============================================================
# CONDITION: CML
# ============================================================
add_h1("Condition — Chronic Myeloid Leukemia (CML)")
add_hint("Sub-page under Conditions. Has Pediatric and Adult tabs.")

add_h2("Pediatric Tab")
add_block("CML_PED_OVERLINE", "Overline.", "Pediatric BMT — Dr. Mahak Agarwal")
add_block("CML_PED_TITLE", "Tab title.", "Chronic Myeloid Leukemia (CML) in Children")

add_h3("FAQ — What is CML in Children?")
add_block("CML_PED_FAQ1_Q", "Accordion question.", "What is CML in Children? Natural History & Causes")
add_block("CML_PED_FAQ1_A1", "Paragraph.",
    "A slowly progressing disease in which too many white blood cells (not lymphocytes) are made in the bone marrow. Chronic Myeloid Leukemia was the first disease where the molecular abnormality leading to cancer was deciphered. A unique translocation between chromosomes 9 and 22 leads to the formation of BCR-ABL fusion protein which results in uncontrolled proliferation of white blood cells.")
add_block("CML_PED_FAQ1_H1", "Sub-heading.", "Phases of Chronic Myeloid Leukemia")
add_label("CML_PED_FAQ1_LIST1")
add_hint("Phases bullets.")
add_bullets([
    "Chronic phase: elevated blood counts with enlarged spleen. Patients are asymptomatic or mildly symptomatic.",
    "Accelerated Phase: The disease starts getting out of hand. The spleen gets bigger despite treatment. The blasts in the blood and marrow start to increase (10 to 19%).",
    "Blast Crisis: This is when the disease develops to a full blown leukaemia with marrow blasts more than 20%.",
])
add_block("CML_PED_FAQ1_H2", "Sub-heading.", "What Causes Chronic Myeloid Leukaemia?")
add_block("CML_PED_FAQ1_A2", "Paragraph.",
    "This is largely unknown as in most cases of leukaemia. Risk factors: More common in males; More common in the elderly with a median age at diagnosis of 65 years; Exposure to ionising radiation; Low Immunity; Inflammatory bowel conditions, such as ulcerative colitis or Crohn's disease; Using pesticides at work.")
add_block("CML_PED_FAQ1_H3", "Sub-heading.", "How is CML in Children different from that in Adults?")
add_label("CML_PED_FAQ1_LIST2")
add_hint("Difference bullets.")
add_bullets([
    "CML is less common in children",
    "Long term outcome of TKI therapy is not known in children. We are not talking of 10 or 20 years of treatment but for 50 to 60 years or more.",
    "BMT should still be considered in children as first line therapy if the right donor is identified.",
])

add_h3("FAQ — Symptoms, Diagnosis & Treatment")
add_block("CML_PED_FAQ2_Q", "Accordion question.", "Symptoms, Diagnosis & Treatment")
add_block("CML_PED_FAQ2_A1", "Paragraph.",
    "In most patients there are no symptoms, 25% of the patients are detected when the disease has progressed to accelerated or blast phase. Common symptoms include: Frequent Infections; Weight loss; Tiredness and looking pale; Swollen lymph glands; Abnormal bruising or bleeding; Abdominal discomfort due to an enlarged spleen; Poor appetite; Sweating at night; Headaches; Bone pain.")
add_block("CML_PED_FAQ2_H1", "Sub-heading.", "Diagnosis")
add_block("CML_PED_FAQ2_A2", "Paragraph.",
    "Complete Blood Count: high white blood cell count with low haemoglobin and high platelet count. Bone Marrow Examination: necessary for confirmation; in chronic phase, blast count is less than 5%. Flow Cytometry: to confirm type of transformation. Cytogenetics: essential for confirmation of diagnosis. Karyotyping: establishes classic translocation between chromosomes 9 and 22. PCR: a test done on the DNA and establishes the amount of BCR/ABL product.")
add_block("CML_PED_FAQ2_H2", "Sub-heading.", "Treatment")
add_block("CML_PED_FAQ2_A3", "Paragraph.",
    "Targeted therapy; Chemotherapy; Biologic therapy; High-dose chemotherapy with Bone Marrow Transplant.")

add_h3("FAQ — Role of BMT in CML")
add_block("CML_PED_FAQ3_Q", "Accordion question.", "Role of BMT in CML")
add_block("CML_PED_FAQ3_A1", "Paragraph.",
    "BMT was the treatment of choice for Chronic Myeloid Leukemia until Tyrosine Kinase Inhibitors (TKI) were developed.")
add_block("CML_PED_FAQ3_H1", "Sub-heading.", "When is BMT needed?")
add_label("CML_PED_FAQ3_LIST1")
add_hint("Indication bullets.")
add_bullets([
    "When patients stop responding to Tyrosine Kinase Inhibitors (TKI)",
    "When Chronic Myeloid Leukemia is in Accelerated Phase or Blast Crisis.",
])
add_block("CML_PED_FAQ3_H2", "Sub-heading.", "Conditioning")
add_block("CML_PED_FAQ3_A2", "Paragraph.",
    "High to Moderate dose of chemotherapy is generally used in conditioning for BMT in younger patients.")
add_block("CML_PED_FAQ3_H3", "Sub-heading.", "Who can be a donor?")
add_block("CML_PED_FAQ3_A3", "Paragraph.",
    "Although, a matched family donor is preferred, a Half matched (Haploidentical) family donor or an unrelated donor can also be used. However a Haploidentical (Half Matched) Donor who has Natural Killer Cell mismatch with the patient is an ideal donor for CML as it reduces the risk of relapse.")
add_block("CML_PED_FAQ3_H4", "Sub-heading.", "Results of BMT in CML")
add_block("CML_PED_FAQ3_A4", "Paragraph.", "BMT remains the only curative treatment for Chronic Myeloid Leukemia.")
add_label("CML_PED_FAQ3_LIST2")
add_hint("Cure rate bullets.")
add_bullets([
    "Chronic Phase: 90% are cured",
    "Accelerated Phase: 40-60% are cured",
    "Blast Crisis: 20-40% are cured",
])

add_h2("Adult Tab")
add_block("CML_AD_OVERLINE", "Overline.", "Adult BMT — Dr. Suparno Chakrabarti")
add_block("CML_AD_TITLE", "Tab title.", "Chronic Myeloid Leukemia (CML) in Adults")

add_h3("FAQ — What is CML in Adults?")
add_block("CML_AD_FAQ1_Q", "Accordion question.", "What is CML in Adults? Natural History & Causes")
add_block("CML_AD_FAQ1_A1", "Paragraph 1.",
    "A slowly progressing disease in which too many white blood cells (not lymphocytes) are made in the bone marrow which is also called Chronic Granulocytic Leukemia, Chronic Myelogenous Leukemia or Chronic Myeloid Leukemia. Chronic Myeloid Leukemia is more common in adults than in children. This arises from abnormality in Blood Stem Cells resulting in proliferation of both mature and immature white blood cells as well as platelets.")
add_block("CML_AD_FAQ1_A2", "Paragraph 2.",
    "Chronic Myeloid Leukemia was the first disease where the molecular abnormality leading to cancer was deciphered. A unique translocation between chromosomes 9 and 22 results in formation of a protein which stimulates the enzyme Tyrosine Kinase leading to uncontrolled proliferation of white blood cells. This abnormality was called the PHILADELPHIA CHROMOSOME. The molecular basis for this was later established as the translocation of BCR gene to the ABL gene.")
add_block("CML_AD_FAQ1_H1", "Sub-heading.", "Phases of Chronic Myeloid Leukemia")
add_label("CML_AD_FAQ1_LIST")
add_hint("Phases bullets.")
add_bullets([
    "Chronic phase: elevated blood counts with enlarged spleen. Patients are asymptomatic or mildly symptomatic.",
    "Accelerated Phase: The disease starts getting out of hand. The spleen gets bigger despite treatment. The blasts in the blood and marrow start to increase (10 to 19%).",
    "Blast Crisis: This is when the disease develops to a full blown leukaemia with marrow blasts more than 20%.",
])
add_block("CML_AD_FAQ1_A3", "Risk factors paragraph.",
    "Risk factors: More common in males; More common in the elderly with a median age at diagnosis of 65 years; Exposure to ionising radiation; Low Immunity; Inflammatory bowel conditions; Using pesticides at work.")

add_h3("FAQ — Symptoms, Diagnosis, Treatment & BMT (Adult)")
add_block("CML_AD_FAQ2_Q", "Accordion question.", "Symptoms, Diagnosis, Treatment & BMT")
add_block("CML_AD_FAQ2_A1", "Paragraph.",
    "In most patients there are no symptoms, 25% of the patients are detected when the disease has progressed to accelerated or blast phase. Symptoms include: Frequent Infections; Weight loss; Tiredness; Swollen lymph glands; Abnormal bruising or bleeding; Abdominal discomfort; Poor appetite; Night sweats; Headaches; Bone pain.")
add_block("CML_AD_FAQ2_A2", "Paragraph.",
    "Diagnosis: Complete Blood Count (high WBC, low Hb, high platelets); Bone Marrow Examination; Flow Cytometry; Cytogenetics; Karyotyping (establishes translocation between chromosomes 9 and 22); PCR (BCR/ABL level).")
add_block("CML_AD_FAQ2_A3", "Paragraph.",
    "Treatment: Targeted therapy; Chemotherapy; Biologic therapy; High-dose chemotherapy with BMT.")
add_block("CML_AD_FAQ2_H1", "Sub-heading.", "BMT for Chronic Myeloid Leukemia")
add_block("CML_AD_FAQ2_A4", "Paragraph.",
    "BMT was the treatment of choice for CML until TKI were developed. When is BMT needed: when patients stop responding to TKI; when CML is in Accelerated Phase or Blast Crisis.")
add_block("CML_AD_FAQ2_A5", "Paragraph.",
    "Conditioning: High to Moderate dose of chemotherapy in younger patients. A Haploidentical (Half Matched) Donor who has Natural Killer Cell mismatch with the patient is ideal. Results: Chronic Phase 90% cured; Accelerated Phase 40-60% cured; Blast Crisis 20-40% cured.")

condition_cta("CML",
    "When to see us for CML",
    "TKI resistance or accelerated/blast phase warrants urgent BMT evaluation.")

add_page_break()


# ============================================================
# CONDITION: HODGKIN'S LYMPHOMA
# ============================================================
add_h1("Condition — Hodgkin's Lymphoma")
add_hint("Sub-page under Conditions. Has Pediatric and Adult tabs.")

add_h2("Pediatric Tab")
add_block("HL_PED_OVERLINE", "Overline.", "Pediatric BMT — Dr. Mahak Agarwal")
add_block("HL_PED_TITLE", "Tab title.", "Hodgkin's Lymphoma in Children")

add_h3("FAQ — What is Hodgkin's Lymphoma?")
add_block("HL_PED_FAQ1_Q", "Accordion question.", "What is Hodgkin's Lymphoma? Types & Causes")
add_block("HL_PED_FAQ1_A1", "Paragraph.",
    "Hodgkin's lymphoma is a distinct cancer of lymphoid tissue derived from germinal B cells of lymph glands, defined by the presence of very characteristic cells called REED-STERNBERG cells or its variants with characteristic immunophenotype in a background of different types of cells.")
add_block("HL_PED_FAQ1_H1", "Sub-heading.", "Types of Hodgkin's Lymphoma")
add_label("HL_PED_FAQ1_LIST1")
add_hint("Types bullets.")
add_bullets([
    "Nodular sclerosing HL",
    "Mixed-cellularity subtype",
    "Lymphocyte-rich or Lymphocytic predominance",
    "Lymphocyte depleted",
    "Nodular lymphocyte predominant Hodgkin's lymphoma (NLPHL)",
])
add_block("HL_PED_FAQ1_H2", "Sub-heading.", "What Causes Hodgkin's Lymphoma?")
add_block("HL_PED_FAQ1_A2", "Paragraph.",
    "The exact cause of Hodgkin disease is not known. Most Hodgkin's lymphomas occur when an infection-fighting cell called a B cell develops a mutation in its DNA. The mutation tells the cells to divide rapidly and to continue living when a healthy cell would die.")
add_block("HL_PED_FAQ1_H3", "Sub-heading.", "Risk Factors")
add_label("HL_PED_FAQ1_LIST2")
add_hint("Risk factor bullets.")
add_bullets([
    "Age: Generally Age group 15-35 and 55 years and above are diagnosed with Hodgkin's lymphoma.",
    "Sex: Mostly Male are suspected to be on higher risk of Hodgkin's lymphoma.",
    "Previously exposed to the Epstein-Barr virus (EBV), a common virus that causes glandular fever.",
    "Medical condition that weaken immune system, such as HIV.",
    "Previous treatment with chemotherapy or radiotherapy for Non-Hodgkin's Lymphoma.",
])

add_h3("FAQ — Symptoms")
add_block("HL_PED_FAQ2_Q", "Accordion question.", "Symptoms")
add_label("HL_PED_FAQ2_LIST1")
add_hint("Main symptoms bullets.")
add_bullets([
    "Painless swelling of the lymph nodes in neck, armpits, groin (swollen glands)",
    "Fatigue",
    "Fever and chills that come and go (38°C / 100.4 F)",
    "Unexplained Itching all over the body",
    "Loss of appetite",
    "Soaking night sweats",
    "Unexplained Weight loss",
])
add_block("HL_PED_FAQ2_H1", "Sub-heading.", "Other symptoms that may occur")
add_label("HL_PED_FAQ2_LIST2")
add_hint("Other symptoms bullets.")
add_bullets([
    "Coughing, chest pains, or breathing problems if there are swollen lymph nodes in the chest",
    "Excessive sweating",
    "Pain or feeling of fullness below the ribs due to swollen spleen or liver",
    "Pain in lymph nodes after drinking alcohol",
    "Skin blushing or flushing",
])

add_h3("FAQ — Diagnosis & Staging")
add_block("HL_PED_FAQ3_Q", "Accordion question.", "Diagnosis & Staging")
add_block("HL_PED_FAQ3_H1", "Sub-heading.", "Lymph Node Biopsy")
add_block("HL_PED_FAQ3_A1", "Paragraph 1.",
    "Hodgkin's lymphoma is diagnosed from Biopsy of lymph gland or other involved tissues. Histopathology reveals five subtypes: Nodular Sclerosis and Mixed cellularity are more common. Detection of cell surface antigens by Immunohistochemistry (IHC) is confirmatory. Expression of CD30 and CD15 are seen in the RS cells, except in NLPHL, where CD20 is expressed.")
add_block("HL_PED_FAQ3_A2", "Paragraph 2.",
    "Blood tests; Imaging tests (X-rays, CT scan, PET); Bone Marrow Biopsy; Additional Tests: Serum LDH, ECG, Echocardiogram and Lung Function Test before starting treatment.")
add_block("HL_PED_FAQ3_H2", "Sub-heading.", "Staging")
add_label("HL_PED_FAQ3_LIST")
add_hint("Staging bullets.")
add_bullets([
    "Stage I: Hodgkin disease is found in only 1 lymph node area or lymphoid organ such as the thymus. The cancer has not spread to other areas of the body.",
    "Stage II: Hodgkin disease is found in 2 or more lymph node areas on the same side of (above or below) the diaphragm.",
    "Stage III: Hodgkin disease is found in lymph node areas on both sides of (above and below) the diaphragm. It may also have spread to an area or organ near the lymph node areas or to the spleen.",
    "Stage IV: Hodgkin disease has spread widely through 1 or more organs outside of the lymph system.",
])

add_h3("FAQ — Treatment & Role of BMT")
add_block("HL_PED_FAQ4_Q", "Accordion question.", "Treatment & Role of BMT")
add_block("HL_PED_FAQ4_H1", "Sub-heading.", "Treatment Options")
add_label("HL_PED_FAQ4_LIST1")
add_hint("Treatment options bullets.")
add_bullets(["Chemotherapy", "Radiation therapy", "High-dose chemotherapy and Bone Marrow transplant"])
add_block("HL_PED_FAQ4_A1", "Paragraph.",
    "Favourable, Limited Stage Disease: Treated with short course chemotherapy and involved field radiotherapy; >95% cure rate.")
add_block("HL_PED_FAQ4_A2", "Paragraph.",
    "Unfavourable, Limited Stage Disease: Similar treatment but more intensive; > 90% cure rate.")
add_block("HL_PED_FAQ4_A3", "Paragraph.", "Advanced Disease: ABVD/BEACOPP >60%-80% cure rates.")
add_block("HL_PED_FAQ4_A4", "Paragraph.",
    "However, use of radiotherapy is not favoured in small children unless absolutely necessary.")
add_block("HL_PED_FAQ4_H2", "Sub-heading.", "BMT for Hodgkin's Lymphoma")
add_block("HL_PED_FAQ4_A5", "Paragraph.",
    "About 30% patients with advanced disease and 10% patients with Limited disease relapse. Once Hodgkin's lymphoma relapses, the treatment is well-defined:")
add_label("HL_PED_FAQ4_LIST2")
add_hint("Treatment sequence.")
add_bullets([
    "Salvage Chemotherapy",
    "Autologous Peripheral Blood Stem Cell Collection",
    "High Dose Chemotherapy and Autologous BMT",
])
add_block("HL_PED_FAQ4_A6", "Paragraph.",
    "If the patient is PET negative before BMT, 80% of the patients are cured. If the patient is PET positive, the cure rate is only 20%.")
add_block("HL_PED_FAQ4_A7", "Paragraph.",
    "If a patient achieves a complete remission, even in advanced disease, there is no need to perform a BMT upfront.")
add_block("HL_PED_FAQ4_H3", "Sub-heading.", "When is Allogeneic BMT needed?")
add_label("HL_PED_FAQ4_LIST3")
add_hint("Allogeneic BMT indications.")
add_bullets([
    "When the disease recurs after Autologous BMT",
    "When PBSC cannot be mobilised for Autologous BMT",
])
add_block("HL_PED_FAQ4_A8", "Paragraph.",
    "There is a strong Graft versus Tumour effect against HL in the setting of Allogeneic BMT. In the setting of High Risk and Relapsed Hodgkin's lymphoma, best results are obtained with HAPLOIDENTICAL DONOR and REDUCED INTENSITY CONDITIONING. 70% patients are cured with this procedure.")

add_h2("Adult Tab")
add_block("HL_AD_OVERLINE", "Overline.", "Adult BMT — Dr. Suparno Chakrabarti")
add_block("HL_AD_TITLE", "Tab title.", "Hodgkin's Lymphoma in Adults")

add_h3("FAQ — What is Hodgkin's Lymphoma? (Adult)")
add_block("HL_AD_FAQ1_Q", "Accordion question.", "What is Hodgkin's Lymphoma? Types, Causes & Symptoms")
add_block("HL_AD_FAQ1_A1", "Paragraph.",
    "Hodgkin's lymphoma is a distinct cancer of lymphoid tissue derived from germinal B cells of lymph glands, defined by the presence of very characteristic cells called REED-STERNBERG cells or its variants with characteristic immunophenotype in a background of different types of cells.")
add_block("HL_AD_FAQ1_A2", "Paragraph.",
    "Types: Nodular sclerosing HL; Mixed-cellularity subtype; Lymphocyte-rich or Lymphocytic predominance; Lymphocyte depleted; Nodular lymphocyte predominant Hodgkin's lymphoma (NLPHL).")
add_block("HL_AD_FAQ1_A3", "Paragraph.",
    "The exact cause of Hodgkin disease is not known. Most Hodgkin's lymphomas occur when an infection-fighting cell called a B cell develops a mutation in its DNA. The mutation tells the cell to divide rapidly and to continue living when a healthy cell would die.")
add_block("HL_AD_FAQ1_A4", "Paragraph.",
    "Symptoms: Painless swelling of lymph nodes in neck, armpits, groin; Fatigue; Fever and chills (38°C/100.4F); Unexplained Itching; Loss of appetite; Soaking night sweats; Unexplained Weight loss; Coughing/chest pains/breathing problems; Excessive sweating; Pain or fullness below ribs; Pain in lymph nodes after drinking alcohol; Skin blushing or flushing.")

add_h3("FAQ — Diagnosis, Staging, Treatment & BMT (Adult)")
add_block("HL_AD_FAQ2_Q", "Accordion question.", "Diagnosis, Staging, Treatment & BMT")
add_block("HL_AD_FAQ2_A1", "Paragraph.",
    "Diagnosis by lymph node biopsy; Histopathology reveals subtypes; IHC is confirmatory (CD30 and CD15 in RS cells, CD20 in NLPHL); Blood tests; Imaging (X-rays, CT, PET); Bone Marrow Biopsy; Additional tests: Serum LDH, ECG, Echocardiogram, Lung Function Test.")
add_block("HL_AD_FAQ2_H1", "Sub-heading.", "Staging: I–IV")
add_block("HL_AD_FAQ2_A2", "Paragraph.",
    "Stage I: 1 lymph node area; Stage II: 2+ areas same side of diaphragm; Stage III: both sides; Stage IV: spread widely to organs outside lymph system.")
add_block("HL_AD_FAQ2_H2", "Sub-heading.", "Treatment")
add_block("HL_AD_FAQ2_A3", "Paragraph.",
    "Favourable Limited: >95% cure rate. Unfavourable Limited: >90% cure rate. Advanced: ABVD/BEACOPP 60–80%. BMT for relapsed cases: Salvage Chemo → Autologous PBSC collection → High Dose Chemo + Autologous BMT.")
add_block("HL_AD_FAQ2_A4", "Paragraph.",
    "PET negative before BMT: 80% cured. PET positive: 20% cure rate. Allogeneic BMT when disease recurs after Autologous BMT or PBSC cannot be mobilised. Haploidentical Donor with Reduced Intensity Conditioning: 70% cure rate.")

condition_cta("HL",
    "When to see us for Hodgkin's Lymphoma",
    "Relapsed or refractory Hodgkin's Lymphoma requires expert BMT evaluation.")

add_page_break()


# ============================================================
# CONDITION: NHL
# ============================================================
add_h1("Condition — Non-Hodgkin's Lymphoma (NHL)")
add_hint("Sub-page under Conditions. Single-tab page.")

add_block("NHL_OVERLINE", "Overline.", "Pediatric BMT — Dr. Mahak Agarwal")
add_block("NHL_TITLE", "Page title.", "Non-Hodgkin's Lymphoma (NHL)")

add_h3("FAQ — What is Non-Hodgkin's Lymphoma?")
add_block("NHL_FAQ1_Q", "Accordion question.", "What is Non-Hodgkin's Lymphoma (NHL)?")
add_block("NHL_FAQ1_A1", "Paragraph 1.",
    "Lymphoma begins when B cells, T cells, or NK cells in the lymphatic system change and grow uncontrollably, which sometimes may form a tumor. Hodgkin lymphoma is a specific type of lymphoma that is covered in another section of this website. Non-Hodgkin lymphoma (NHL) is a term that refers to the many other types of cancer of the lymphatic system, which can have different symptoms and signs, physical findings, and treatments.")
add_block("NHL_FAQ1_A2", "Paragraph 2.",
    "Because lymphatic tissue is found in most parts of the body, NHL can start almost anywhere and can spread to almost any organ. It most often begins in the lymph nodes, liver, spleen, or bone marrow, but it can also involve the stomach, intestines, skin, thyroid gland, brain, or any other part of the body.")
add_block("NHL_FAQ1_A3", "Paragraph 3.",
    "It is very important to know which type and subtype has been diagnosed because this information helps doctors determine the best treatment and a patient's prognosis (chance of recovery).")
add_block("NHL_FAQ1_A4", "Paragraph 4.", "Non-Hodgkin's Lymphoma is more common in adults than children.")

add_h3("FAQ — Causes & Risk Factors")
add_block("NHL_FAQ2_Q", "Accordion question.", "Causes & Risk Factors")
add_block("NHL_FAQ2_INTRO", "Intro line.",
    "Certain types of Non-Hodgkin's Lymphoma are caused or triggered by Infections. Most common causes are:")
add_label("NHL_FAQ2_LIST1")
add_hint("Causes bullets.")
add_bullets([
    "EBV and Burkitts lymphoma",
    "Hepatitis C",
    "HTLV1",
    "Helicobacter Pylori causing lymphoma of the stomach.",
    "HIV infection increases the incidence of high grade lymphomas",
    "Long term intake of drugs suppressing the immune system as after solid organ transplants",
    "Primary immunodeficiency and autoimmune diseases",
    "Chemicals: polychlorinated biphenyls (PCBs), diphenylhydantoin, dioxin, and phenoxy herbicides.",
    "Medical treatments: like radiation therapy and chemotherapy",
    "Genetic diseases, like Klinefelter's syndrome, Chédiak-Higashi syndrome, ataxia telangiectasia syndrome",
    "Autoimmune diseases, like Sjögren's syndrome, celiac sprue, rheumatoid arthritis, and systemic lupus erythematosus.",
])
add_block("NHL_FAQ2_H1", "Sub-heading.", "Risk Factors")
add_label("NHL_FAQ2_LIST2")
add_hint("Risk factor bullets.")
add_bullets([
    "Being male: Non-Hodgkin's Lymphoma is more common in men than in women.",
    "Age: Likelihood of getting Non-Hodgkin's Lymphoma increases with age.",
    "Impaired immune system: NHL is most common among those who have an impaired immune system, an autoimmune disease, or HIV or other viral infections.",
    "Viral infections: A viral infection, such as Epstein-Barr virus, increases the risk of developing NHL.",
    "Bacterial infections: Infection with Helicobacter pylori increases the risk of lymphoma involving the stomach.",
    "Environmental exposure: Exposure to agricultural pesticides or fertilizers, solvents, and other chemicals may increase the risk of getting NHL.",
])

add_h3("FAQ — Symptoms")
add_block("NHL_FAQ3_Q", "Accordion question.", "Symptoms")
add_block("NHL_FAQ3_INTRO", "Intro paragraph.",
    "The symptoms are protean. They range from painless lymph node enlargement to fits and paralysis. Most commonly, patients present with 'B symptoms' defined by:")
add_label("NHL_FAQ3_LIST")
add_hint("Symptoms bullets.")
add_bullets([
    "A painless swelling of the lymph nodes in the neck, underarms, or groin. This is the most common symptom of NHL.",
    "Fever above 38°C (100.4 F)",
    "Drenching Night sweats.",
    "Feeling very tired.",
    "Weight loss > 10% of baseline body weight",
    "Itchy skin.",
    "Reddened patches on the skin.",
    "A cough or shortness of breath.",
    "Pain in the belly or back.",
])

add_h3("FAQ — Diagnosis")
add_block("NHL_FAQ4_Q", "Accordion question.", "Diagnosis")
add_block("NHL_FAQ4_H1", "Sub-heading.", "Biopsy and Histopathology")
add_block("NHL_FAQ4_A1", "Paragraph.",
    "NHL is diagnosed from Biopsy of lymph gland or other involved tissues. Fine Needle aspiration Cytology (FNAC) is not enough for diagnosis of Lymphoma.")
add_block("NHL_FAQ4_A2", "Paragraph.",
    "Histopathological diagnosis reveals two major subtypes: T cell Non-Hodgkin's Lymphoma and B cell Non-Hodgkin's Lymphoma. B cell Non-Hodgkin's Lymphoma is more common. This is further classified as HIGH GRADE and LOW GRADE. Detection of cell surface antigens by Immunohistochemistry (IHC) is confirmatory.")
add_block("NHL_FAQ4_A3", "Paragraph.",
    "Complete Blood Count: Anemia and Thrombocytopenia may be seen. Eosinophil count is often raised.")
add_block("NHL_FAQ4_A4", "Paragraph.",
    "Bone Marrow Biopsy: This is necessary to ascertain the extent of the disease. Rarely, the presentation is with Bone Marrow and Blood involvement.")
add_block("NHL_FAQ4_A5", "Paragraph.",
    "PET-CT Scan: This has revolutionised the diagnosis and management of NHL. In the early days, surgery was performed to stage the disease.")
add_block("NHL_FAQ4_A6", "Paragraph.",
    "Additional Tests: Serum LDH, ECG, Echocardiogram and Lung Function Tests are done before starting treatment.")

add_h3("FAQ — Treatment")
add_block("NHL_FAQ5_Q", "Accordion question.", "Treatment")
add_block("NHL_FAQ5_A1", "Paragraph.",
    "Treatment plans are designed to meet the unique needs of each person with lymphoma. Treatment decisions are made depending on: the type of NHL; the stage; how quickly the NHL is growing (grade); prognostic factors; person's age; person's overall health status; previous treatment, if any.")
add_block("NHL_FAQ5_H1", "Sub-heading.", "General Principles")
add_block("NHL_FAQ5_A2", "Paragraph.",
    "HIGH GRADE B CELL NHL: Chemotherapy called R-CHOP regimen remains the gold standard for primary treatment. However, based on GENE EXPRESSION PROFILE, Diffuse large B cell Lymphoma (DLBCL) is classified in three groups, GCB, ABC and PMBCL. The latter two groups have a poorer outcome and require upfront BMT. However, we do not favour use of radiotherapy in small children unless absolutely necessary.")
add_block("NHL_FAQ5_A3", "Paragraph.",
    "LOW GRADE B CELL NHL: These lymphomas occur exclusively in older individuals and are treated in the same way as CLL.")
add_block("NHL_FAQ5_A4", "Paragraph.",
    "T CELL NHL: They are mostly high grade and some of them are triggered by HTLV-1 virus. The response to CHOP is less encouraging. This group of patients require early BMT.")
add_block("NHL_FAQ5_H2", "Sub-heading.", "Treatment Options Other Than BMT")
add_label("NHL_FAQ5_LIST")
add_hint("Non-BMT treatment bullets.")
add_bullets([
    "Chemotherapy: Single chemotherapy drugs or combinations of drugs can be given.",
    "Biological therapy: May be used on its own or in combination with chemotherapy.",
    "Radiation therapy: External beam radiation therapy may be used on its own to treat localized areas of early stage lymphoma.",
    "Surgery: Surgery is mainly used to remove all or part of a lymph node (biopsy) to diagnose lymphoma.",
    "Follow-up: It is important to have regular follow-up visits, especially during the first 2 years after treatment.",
])

add_h3("FAQ — Role of BMT in NHL")
add_block("NHL_FAQ6_Q", "Accordion question.", "Role of BMT in NHL")
add_block("NHL_FAQ6_H1", "Sub-heading.", "What to do when Non-Hodgkin's Lymphoma recurs?")
add_block("NHL_FAQ6_A1", "Paragraph.",
    "About 30% patients with advanced disease and 10% patients with Limited disease relapse. Once Non-Hodgkin's Lymphoma relapses, the treatment is well-defined:")
add_label("NHL_FAQ6_LIST1")
add_hint("Relapse treatment steps.")
add_bullets([
    "Salvage Chemotherapy",
    "Autologous Peripheral Blood Stem Cell Collection",
    "High Dose Chemotherapy and Autologous BMT",
])
add_block("NHL_FAQ6_H2", "Sub-heading.", "Cure Rate with Autologous BMT")
add_block("NHL_FAQ6_A2", "Paragraph.",
    "If the patient is PET negative before BMT, 80% of the patients are cured. If the patient is PET positive, the cure rate is only 20%.")
add_block("NHL_FAQ6_H3", "Sub-heading.", "When is Upfront BMT Required?")
add_label("NHL_FAQ6_LIST2")
add_hint("Upfront BMT indications.")
add_bullets(["Patients with more advanced disease", "T cell NHL", "Mantle Cell Lymphoma"])
add_block("NHL_FAQ6_H4", "Sub-heading.", "When is Allogeneic BMT Needed for Non-Hodgkin's Lymphoma?")
add_label("NHL_FAQ6_LIST3")
add_hint("Allogeneic BMT indications.")
add_bullets([
    "When the disease recurs after Autologous BMT",
    "When Peripheral Blood Stem Cells cannot be mobilised for Autologous BMT",
    "T cell NHL",
    "Relapsed Low Grade NHL",
    "Relapsed Mantle Cell Lymphoma",
])
add_block("NHL_FAQ6_A3", "Paragraph.",
    "There is a strong Graft versus Tumour effect against NHL in the setting of Allogeneic BMT. In the setting of High Risk and Relapsed NHL, best results are obtained with HAPLOIDENTICAL DONOR and reduced toxicity transplantation. 70% patients are cured with this procedure. The children have more aggressive varieties of NHL. The cure rates in children are much higher.")

condition_cta("NHL",
    "When to see us for NHL",
    "Relapsed NHL, T cell NHL, and Mantle Cell Lymphoma require BMT specialist evaluation.")

add_page_break()


# ============================================================
# CONDITION: MDS
# ============================================================
add_h1("Condition — Myelodysplastic Syndromes (MDS)")
add_hint("Sub-page under Conditions. Has Pediatric and Adult tabs.")

add_h2("Pediatric Tab")
add_block("MDS_PED_OVERLINE", "Overline.", "Pediatric BMT — Dr. Suparno Chakrabarti")
add_block("MDS_PED_TITLE", "Tab title.", "Myelodysplastic Syndrome (MDS) in Children")

add_h3("FAQ — What is MDS of Childhood?")
add_block("MDS_PED_FAQ1_Q", "Accordion question.", "What is MDS of Childhood? Symptoms & Classification")
add_block("MDS_PED_FAQ1_A1", "Paragraph 1.",
    "Myelodysplastic syndromes (MDS) are a heterogeneous group of disorders that can occur when the blood-forming cells in the bone marrow are damaged. This damage leads to low numbers of one or more type of blood cells. They represent 5–10% of all myeloid malignancies in children. MDS is characterized by ineffective haematopoiesis and increased cell death.")
add_block("MDS_PED_FAQ1_A2", "Paragraph 2.",
    "MDS can affect children as well, although the incidence is less than 5%. Anemia is the commonest manifestation of MDS in adults, but low white cell or platelet count is commoner in children.")
add_block("MDS_PED_FAQ1_A3", "Paragraph 3.",
    "The bone marrow is more often less cellular like Aplastic anemia and this entity is called Refractory Cytopenia of Childhood (RCC). This condition has to be differentiated from Aplastic Anemia.")
add_block("MDS_PED_FAQ1_H1", "Sub-heading.", "Symptoms")
add_label("MDS_PED_FAQ1_LIST1")
add_hint("Symptom bullets.")
add_bullets([
    "Shortness of breath",
    "Weakness or feeling tired.",
    "Having skin that is paler than usual",
    "Easy bruising or bleeding",
    "Petechiae (flat, pinpoint spots under the skin caused by bleeding)",
    "Fever or frequent infections",
])
add_block("MDS_PED_FAQ1_H2", "Sub-heading.", "Classification")
add_block("MDS_PED_FAQ1_A4", "Paragraph.",
    "This is based on number of cell lines affected (i.e. 1, 2 or 3) and the number of blasts in the bone marrow.")
add_label("MDS_PED_FAQ1_LIST2")
add_hint("Classification bullets.")
add_bullets([
    "Refractory Anemia",
    "Refractory Cytopenia with Multilineage Dysplasia",
    "Refractory Anemia with Ring Sideroblasts",
    "Refractory Anemia with Excess Blasts",
])
add_block("MDS_PED_FAQ1_A5", "Paragraph.",
    "The three most important parameters determining the outcome: 1. Number of Cytopenias, 2. %age of blasts in bone marrow, 3. Abnormalities in the chromosomes. Based on these three parameters, an International Prognostic Scoring System has been developed.")

add_h3("FAQ — Diagnosis & Treatment with BMT")
add_block("MDS_PED_FAQ2_Q", "Accordion question.", "Diagnosis & Treatment with BMT")
add_block("MDS_PED_FAQ2_A1", "Paragraph.",
    "The diagnosis is made by careful and diligent examination of Blood and Bone Marrow samples by an experienced Hematologist. Along with that study of chromosomes from the bone marrow cells (Cytogenetics) is needed.")
add_block("MDS_PED_FAQ2_A2", "Paragraph.",
    "The abnormalities of chromosomes are divided as GOOD or BAD: Good: normal, -Y, del5q, del20q. Bad: Monosomy i.e. deletion of one of the pair of chromosomes, mostly chromosomes 5 and 7 and other complex abnormalities.")
add_block("MDS_PED_FAQ2_H1", "Sub-heading.", "Treatment")
add_block("MDS_PED_FAQ2_A3", "Paragraph.",
    "The Only Curative Treatment of MYELODYSPLASTIC SYNDROME (MDS) Is an Allogeneic BMT.")
add_block("MDS_PED_FAQ2_A4", "Paragraph.",
    "As the disease is mostly seen in older patients, a Reduced Intensity Conditioning is preferred.")
add_block("MDS_PED_FAQ2_A5", "Paragraph.",
    "The BMT should be done early, before the onset of life threatening infections or multiple transfusions which may lead to iron overload. Some of the patients with higher blast count or abnormal chromosomes might benefit from a short course of chemotherapy before the BMT.")
add_block("MDS_PED_FAQ2_H2", "Sub-heading.", "Who can be a donor?")
add_block("MDS_PED_FAQ2_A6", "Paragraph.",
    "Although we prefer a matched family donor, a Half matched (Haploidentical) family donor or an unrelated donor provide excellent survival. However a HAPLOIDENTICAL DONOR, who has Natural Killer Cell mismatch with the patient, provides the best survival through its graft versus leukemia effect.")
add_block("MDS_PED_FAQ2_H3", "Sub-heading.", "Other Treatments (for older patients not fit for BMT)")
add_label("MDS_PED_FAQ2_LIST")
add_hint("Other treatment bullets.")
add_bullets([
    "Red Cell and Platelet Transfusions as needed",
    "Erythropoetin injections to reduce blood transfusion requirements",
    "Hypomethylating agents such as Azacytidine or Decitabine which are milder forms of chemotherapy.",
])

add_h2("Adult Tab")
add_block("MDS_AD_OVERLINE", "Overline.", "Adult BMT — Dr. Suparno Chakrabarti")
add_block("MDS_AD_TITLE", "Tab title.", "Myelodysplastic Syndromes (MDS) in Adults")

add_h3("FAQ — What is MDS? (Adult)")
add_block("MDS_AD_FAQ1_Q", "Accordion question.", "What is MDS? Causes, Symptoms & Classification")
add_block("MDS_AD_FAQ1_A1", "Paragraph.",
    "A myelodysplastic syndrome is a type of cancer in which the bone marrow does not make enough healthy blood cells and there are abnormal (blast) cells in the blood and/or bone marrow.")
add_block("MDS_AD_FAQ1_A2", "Paragraph.",
    "Myelodysplastic syndromes (MDS) are conditions that can occur when the blood-forming cells in the bone marrow are damaged. This damage leads to low numbers of one or more type of blood cells.")
add_block("MDS_AD_FAQ1_A3", "Paragraph.",
    "MDS is a spectrum of diseases which are characterised by anemia, thrombocytopenia and low white cell counts along with dysplasia (meaning abnormal looking) blood cells in both peripheral blood and bone marrow. They are early stages of leukemia. Many develop acute leukemia eventually, but most die because of bleeding or infections before development of leukemia.")
add_block("MDS_AD_FAQ1_A4", "Paragraph.",
    "This is commonest disease developing as a side-effect of treatment of other cancers with chemotherapy and radiotherapy and is called Therapy-related MDS or t-MDS.")
add_block("MDS_AD_FAQ1_H1", "Sub-heading.", "Symptoms")
add_block("MDS_AD_FAQ1_A5", "Paragraph.",
    "Shortness of breath; Weakness or feeling tired; Skin paler than usual; Easy bruising or bleeding; Petechiae; Fever or frequent infections.")
add_block("MDS_AD_FAQ1_H2", "Sub-heading.", "Classification")
add_block("MDS_AD_FAQ1_A6", "Paragraph.",
    "Refractory Anemia; Refractory Cytopenia with Multilineage Dysplasia; Refractory Anemia with Ring Sideroblasts; Refractory Anemia with Excess Blasts. Three key parameters: Number of Cytopenias, % of blasts, Chromosomal abnormalities (International Prognostic Scoring System).")

add_h3("FAQ — Diagnosis, Treatment & BMT (Adult)")
add_block("MDS_AD_FAQ2_Q", "Accordion question.", "Diagnosis, Treatment & BMT")
add_block("MDS_AD_FAQ2_A1", "Paragraph.",
    "Diagnosis by careful examination of Blood and Bone Marrow by an experienced Hematologist. Cytogenetics essential. Good chromosomes: normal, -Y, del5q, del20q. Bad: Monosomy (chromosomes 5 and 7) and complex abnormalities.")
add_block("MDS_AD_FAQ2_H1", "Sub-heading.", "Treatment")
add_block("MDS_AD_FAQ2_A2", "Paragraph.",
    "The Only Curative Treatment of MYELODYSPLASTIC SYNDROME (MDS) Is an Allogeneic BMT. Reduced Intensity Conditioning preferred. BMT should be done early, before life-threatening infections or iron overload.")
add_block("MDS_AD_FAQ2_A3", "Paragraph.",
    "A HAPLOIDENTICAL DONOR, who has Natural Killer Cell mismatch with the patient, provides the best cure for MDS via Graft versus Leukemia (GVL) effect.")
add_block("MDS_AD_FAQ2_H2", "Sub-heading.", "Other Treatments (for those not fit for BMT)")
add_label("MDS_AD_FAQ2_LIST")
add_hint("Other treatment bullets.")
add_bullets([
    "Red Cell and Platelet Transfusions as needed",
    "Erythropoetin injections to reduce blood transfusion requirements",
    "Hypomethylating agents such as Azacytidine or Decitabine which are milder forms of chemotherapy.",
])

# MDS CTA (no phone button on this one on the site)
add_h3("Bottom CTA Bar")
add_block("MDS_CTA_TITLE", "CTA heading.", "When to see us for MDS")
add_block("MDS_CTA_SUBTITLE", "CTA subtitle.",
    "Early BMT before iron overload or serious infections is critical for best outcomes.")
add_block("MDS_CTA_BUTTON", "CTA button.", "Book Consultation")

add_page_break()


# ============================================================
# CONDITION: AUTOIMMUNE
# ============================================================
add_h1("Condition — Autoimmune Diseases (AID)")
add_hint("Sub-page under Conditions. Single-tab page.")

add_block("AID_OVERLINE", "Overline.", "Adult BMT — Dr. Suparno Chakrabarti")
add_block("AID_TITLE", "Page title.", "Autoimmune Diseases (AID)")

add_h3("FAQ — What are Autoimmune Diseases?")
add_block("AID_FAQ1_Q", "Accordion question.", "What are Autoimmune Diseases?")
add_block("AID_FAQ1_A1", "Paragraph.",
    "Autoimmune diseases are a group of diverse diseases caused by an abnormal immune response of the body against substances and tissues normally present in the body (autoimmunity). They usually occur in the early adulthood to middle age and are more common in females.")
add_block("AID_FAQ1_H1", "Sub-heading.", "Examples of Autoimmune Diseases (AID)")
add_label("AID_FAQ1_LIST")
add_hint("Examples bullets.")
add_bullets([
    "Rheumatoid Arthritis",
    "Systemic Lupus Erythematosus",
    "Scleroderma",
    "Juvenile Rheumatoid Arthritis",
    "Vasculitis",
    "Multiple Sclerosis",
    "Autoimmune Hemolytic Anemia",
    "Autoimmune Thrombocytopenia",
    "Crohn's Disease etc.",
])
add_block("AID_FAQ1_H2", "Sub-heading.", "What Causes Autoimmune Disease?")
add_block("AID_FAQ1_A2", "Paragraph.",
    "The cause of autoimmune disease is unknown. If one has a family member with an autoimmune disease, one may be more susceptible to developing the disease. There are many theories about what triggers autoimmune diseases, including: Bacteria / viruses; Drugs; Chemical irritants; Environmental irritants.")
add_block("AID_FAQ1_H3", "Sub-heading.", "Treatment")
add_block("AID_FAQ1_A3", "Paragraph.",
    "They are usually treated with drugs that suppress the immune system such as steroids and other drugs.")

add_h3("FAQ — Role of BMT in Autoimmune Disease")
add_block("AID_FAQ2_Q", "Accordion question.", "Role of BMT in Autoimmune Disease")
add_block("AID_FAQ2_H1", "Sub-heading.", "How does Bone Marrow Transplantation (BMT) work in Autoimmune Disease?")
add_block("AID_FAQ2_A1", "Paragraph.", "This is possible in two different ways:")
add_block("AID_FAQ2_A2", "Paragraph.",
    "Autologous Bone Marrow Transplantation (BMT): The bone marrow and the immune system of the patient is killed with high dose chemotherapy. The stem cells of the patient, which are collected before the procedure, are reinfused. When the stem cells are reinfused, the immune system starts from scratch and if the trigger for autoimmune disease has been removed, the disease is cured. This is called 'Re-booting' of the immune system.")
add_block("AID_FAQ2_A3", "Paragraph.",
    "Allogeneic Bone Marrow Transplantation (BMT): This is a more definitive way of replacing the defective or diseased immune system with a new or healthy immune system from a donor.")
add_block("AID_FAQ2_H2", "Sub-heading.", "What type of BMT is used?")
add_block("AID_FAQ2_A4", "Paragraph.",
    "This depends on the severity and type of autoimmune disease. The decision is taken by the rheumatologist, neurologist and BMT physician together.")
add_block("AID_FAQ2_H3", "Sub-heading.", "Which diseases benefit most from BMT?")
add_block("AID_FAQ2_A5", "Paragraph.",
    "Autologous BMT shows most promising results for Scleroderma and Multiple Sclerosis. This should be done before permanent organ damage has occurred.")
add_block("AID_FAQ2_A6", "Paragraph.",
    "It can be useful for all autoimmune diseases, when the conventional treatments fail.")
add_block("AID_FAQ2_H4", "Sub-heading.", "Why was this procedure not being carried out earlier in India?")
add_label("AID_FAQ2_LIST")
add_hint("Reasons bullets.")
add_bullets([
    "Lack of awareness",
    "Lack of collaboration between rheumatologists, neurologists and BMT Physicians",
    "Lack of the right expertise and experience in carrying out BMT for this condition.",
    "Lack of benefit when the patients are referred late.",
])
add_block("AID_FAQ2_A7", "Paragraph.", "This procedure is now available at our BMT centre.")

# CTA (no phone button)
add_h3("Bottom CTA Bar")
add_block("AID_CTA_TITLE", "CTA heading.", "When to see us for Autoimmune Disease")
add_block("AID_CTA_SUBTITLE", "CTA subtitle.",
    "Refractory autoimmune disease, especially Scleroderma or MS, may benefit from BMT before permanent organ damage.")
add_block("AID_CTA_BUTTON", "CTA button.", "Book Consultation")

add_page_break()


# ============================================================
# CONDITION: PRIMARY IMMUNODEFICIENCY
# ============================================================
add_h1("Condition — Primary Immunodeficiency Disorders (PID)")
add_hint("Sub-page under Conditions. Single-tab page.")

add_block("PID_OVERLINE", "Overline.", "Pediatric BMT — Dr. Suparno Chakrabarti")
add_block("PID_TITLE", "Page title.", "Primary Immunodeficiency Disorders (PID)")

add_h3("FAQ — What is Primary Immune Deficiency Disease?")
add_block("PID_FAQ1_Q", "Accordion question.", "What is Primary Immune Deficiency Disease?")
add_block("PID_FAQ1_A1", "Paragraph 1.",
    "Primary cellular immunodeficiencies (PID) are a group of inherited disorders characterized by severe impairment of the immune systems in the body, which generally leads to early death from infectious complications. Mainly T lymphocytes, B lymphocytes and Natural Killer cells play a pivotal role for immune system, whereas neutrophils and other phagocytes play a role in primary defence.")
add_block("PID_FAQ1_A2", "Paragraph 2.",
    "To be considered a primary immunodeficiency, the cause of the immune deficiency must not be secondary in nature (i.e., caused by other disease, drug treatment, or environmental exposure to toxins). Most primary immunodeficiencies are genetic disorders; the majority are diagnosed in children under the age of one, although milder forms may not be recognized until adulthood.")
add_block("PID_FAQ1_H1", "Sub-heading.", "What Causes Primary Immune Deficiency Disease?")
add_block("PID_FAQ1_A3", "Paragraph.",
    "While not contagious, these diseases are caused by hereditary or genetic defects, and, although most disorders present at birth or in early childhood, the disorders can affect anyone, regardless of age or gender. However, certain varieties affect the male child only. Some affect a single part of the immune system; others may affect one or more components of the system.")

add_h3("FAQ — Symptoms & Classification")
add_block("PID_FAQ2_Q", "Accordion question.", "Symptoms & Classification")
add_block("PID_FAQ2_A1", "Intro.",
    "While the diseases may differ depending on type and person to person, they all share one common feature: each results from a defect in one of the functions of the body's normal immune system and results in severe and recurrent infections. The symptoms are usually in the form of:")
add_label("PID_FAQ2_LIST")
add_hint("Symptom bullets.")
add_bullets([
    "Frequent and recurrent pneumonia",
    "Bronchitis",
    "Sinus infections",
    "Ear infections",
    "Meningitis",
    "Skin infections and blood infections",
])
add_block("PID_FAQ2_A2", "Paragraph.",
    "In addition to frequent infections, other problems that may occur include: Abnormalities in nails, skin and bones; Inflammation and infection of internal organs; Blood disorders, such as low platelet counts or anemia; Digestive problems; Delayed growth and development; Autoimmune disorders.")
add_block("PID_FAQ2_H1", "Sub-heading.", "Classification of Primary Immune Deficiency Disease")
add_block("PID_FAQ2_A3", "Paragraph.",
    "Unlike cancers, diagnostic criteria and conditioning for PID diseases is often more complex. Common diseases include:")
add_block("PID_FAQ2_A4", "Paragraph.",
    "SEVERE COMBINED IMMUNE DEFICIENCY (SCID) is caused by a group of genetic disorders with a shared phenotype of deficient or absent T lymphocyte function. Unless the immune system is restored by Allogeneic BMT, children with SCID generally die of infections during the first year of life.")
add_block("PID_FAQ2_A5", "Paragraph.",
    "COMMON VARIABLE IMMUNE DEFICIENCY (CVID) is a disorder characterised by abnormality in B cells resulting in decreased or absent levels of immunoglobulins (IgG, IgA, IgM). Patients are highly susceptible to infections of the ears, sinuses, and lungs.")
add_block("PID_FAQ2_A6", "Paragraph.",
    "Other diseases include syndromes with T-cell defects such as Wiskott–Aldrich Syndrome (WAS) and Hyper IGM1 Syndrome, and inherited predispositions to Hemophagocytic Lymphohistiocytosis (HLH).")

add_h3("FAQ — Diagnosis")
add_block("PID_FAQ3_Q", "Accordion question.", "Diagnosis")
add_block("PID_FAQ3_A1", "Paragraph.",
    "Complete Blood Count: A careful examination of routine blood count often gives an indication by reduced lymphocyte or neutrophil counts.")
add_block("PID_FAQ3_A2", "Paragraph.",
    "Flow Cytometry: This is the key test to determine the presence or absence or reduction in various immune cells. This also looks for specific marker proteins like WASP, Perforin, SAP, XIAP which are absent in particular syndromes.")
add_block("PID_FAQ3_A3", "Paragraph.",
    "Quantitative Immunoglobulins: A measure of various immunoglobulin levels. This is done to detect abnormalities in the B lymphocytes.")
add_block("PID_FAQ3_A4", "Paragraph.", "Cytogenetics: Done to find out the abnormalities in the chromosomes.")
add_block("PID_FAQ3_A5", "Paragraph.",
    "Antenatal Screening: For mothers who have already had a child with a PID, prenatal testing can be done through Chorionic Villus Sampling or Amniocentesis.")
add_block("PID_FAQ3_H1", "Sub-heading.", "Complications")
add_label("PID_FAQ3_LIST")
add_hint("Complications bullets.")
add_bullets([
    "Recurrent infections",
    "Autoimmune disorders",
    "Damage to heart, lungs, nervous system or digestive tract",
    "Slowed growth",
    "Increased risk of cancer",
    "Death from serious infection",
])

add_h3("FAQ — Role of BMT in Primary Immunodeficiency")
add_block("PID_FAQ4_Q", "Accordion question.", "Role of BMT in Primary Immunodeficiency")
add_block("PID_FAQ4_A1", "Paragraph.",
    "Definitive cure is generally only achieved by ALLOGENEIC BMT, though recent advances in gene therapy hold significant promise.")
add_block("PID_FAQ4_A2", "Paragraph.",
    "When is BMT needed? As soon as it is diagnosed — several studies have demonstrated that infants transplanted at less than 3.5 months of age have improved survival. This is likely due to development of pulmonary infections prior to transplant.")
add_block("PID_FAQ4_A3", "Paragraph.",
    "Conditioning: Few cases don't need conditioning. Otherwise Reduced Intensity Conditioning with low doses of chemotherapy or radiation is the preferred option.")
add_block("PID_FAQ4_A4", "Paragraph.",
    "Donor: Although we prefer a matched family donor, a Half Matched (Haploidentical) family donor or a Matched Unrelated donor provide excellent survival.")
add_block("PID_FAQ4_HIGHLIGHT", "Highlighted box.",
    "Our Results: BMT is the only cure for patients who suffer from PID. Our own experience with T cell depleted HAPLOIDENTICAL BMT has shown nearly 100% cure rate in this group of children.")

# CTA (no phone button)
add_h3("Bottom CTA Bar")
add_block("PID_CTA_TITLE", "CTA heading.", "When to Refer for Primary Immunodeficiency")
add_block("PID_CTA_SUBTITLE", "CTA subtitle.",
    "Early diagnosis and BMT significantly improves survival. Refer immediately on diagnosis.")
add_block("PID_CTA_BUTTON", "CTA button.", "Book Consultation")

add_page_break()


# ============================================================
# CONDITION: INHERITED METABOLIC DISORDERS
# ============================================================
add_h1("Condition — Inherited Metabolic Disorders (IMD)")
add_hint("Sub-page under Conditions. Single-tab page.")

add_block("IMD_OVERLINE", "Overline.", "Pediatric BMT — Dr. Suparno Chakrabarti")
add_block("IMD_TITLE", "Page title.", "Inherited Metabolic Disorders (IMD)")

add_h3("FAQ — What are Inherited Metabolic Disorders?")
add_block("IMD_FAQ1_Q", "Accordion question.", "What are Inherited Metabolic Disorders?")
add_block("IMD_FAQ1_A1", "Paragraph 1.",
    "Inherited Metabolic Disorders refer to medical conditions which are caused by genetic defects mostly inherited from both parents that interfere with the body's metabolism. Metabolism is the complex set of chemical reactions that our body uses to maintain life, including energy production. Special enzymes break down food or certain chemicals so our body can use them right away for fuel or store them. Also, certain chemical processes break down substances that our body no longer needs, or make those it lacks.")
add_block("IMD_FAQ1_A2", "Paragraph 2.",
    "IMD are a complex group of disorders, due to an inborn deficiency of a particular enzyme, functions of one or more organs are affected. The onset and severity depends on the degree of deficiency of the particular enzyme and the organ of affection. Most diseases manifest in the childhood.")

add_h3("FAQ — Symptoms")
add_block("IMD_FAQ2_Q", "Accordion question.", "Symptoms")
add_block("IMD_FAQ2_INTRO", "Intro paragraph.",
    "In Children, few early symptoms include: Apnea, Lethargy, Poor Feeding, Tachypnoea, Vomiting. There are enormous numbers of diseases with wide range of systems affected — nearly every organ can be affected:")
add_label("IMD_FAQ2_LIST")
add_hint("Symptoms bullets.")
add_bullets([
    "Growth failure, failure to thrive, weight loss",
    "Developmental delay, seizures, dementia, encephalopathy, stroke",
    "Deafness, blindness, pain agnosia",
    "Skin rash, abnormal pigmentation, lack of pigmentation, excessive hair growth",
    "Dental abnormalities",
    "Immunodeficiency, thrombocytopenia, anemia, enlarged spleen, enlarged lymph nodes",
    "Many forms of cancer",
    "Recurrent vomiting, diarrhea, abdominal pain",
    "Excessive urination, renal failure, dehydration, edema",
    "Hypotension, heart failure, enlarged heart, hypertension, myocardial infarction",
    "Hepatomegaly, jaundice, liver failure",
    "Unusual facial features, congenital malformations",
    "Abnormal behavior, depression, psychosis",
    "Joint pain, muscle weakness, cramps",
    "Hypothyroidism, adrenal insufficiency, hypogonadism, diabetes mellitus",
])

add_h3("FAQ — Diagnosis")
add_block("IMD_FAQ3_Q", "Accordion question.", "Diagnosis")
add_block("IMD_FAQ3_A1", "Paragraph.",
    "They need a high index of suspicion to be diagnosed by an experienced child specialist. The confirmation is through biochemical tests and genetic tests.")
add_block("IMD_FAQ3_H1", "Sub-heading.", "Common Screening Tests")
add_label("IMD_FAQ3_LIST1")
add_hint("Screening tests bullets.")
add_bullets([
    "Ferric chloride test (turned colors in reaction to various abnormal metabolites in urine)",
    "Ninhydrin paper chromatography (detected abnormal amino acid patterns)",
    "Guthrie bacterial inhibition assay (detected a few amino acids in excessive amounts in blood)",
])
add_block("IMD_FAQ3_H2", "Sub-heading.", "Modern Diagnostic Tools")
add_label("IMD_FAQ3_LIST2")
add_hint("Modern tools bullets.")
add_bullets([
    "Quantitative measurement of amino acids in plasma and urine",
    "Urine organic acid analysis by Gas chromatography-mass spectrometry",
    "Plasma acylcarnitines analysis by mass spectrometry",
    "Urine purines and pyrimidines analysis by Gas chromatography-mass spectrometry",
])
add_block("IMD_FAQ3_H3", "Sub-heading.", "Specific Diagnostic Tests")
add_label("IMD_FAQ3_LIST3")
add_hint("Specific tests bullets.")
add_bullets([
    "Tissue biopsy or necropsy: liver, muscle, brain, bone marrow",
    "Skin biopsy and fibroblast cultivation for specific enzyme testing",
    "Specific DNA testing",
])

add_h3("FAQ — Role of BMT")
add_block("IMD_FAQ4_Q", "Accordion question.", "Role of BMT")
add_block("IMD_FAQ4_A1", "Paragraph.",
    "How is BMT curative? Certain cells produced in the bone marrow produce the particular enzymes deficient in a particular IMD. When healthy donor cells are transplanted, they populate the body and the brain and start producing the necessary enzyme. This leads to a gradual improvement in the organ functions and stabilizes the brain.")
add_block("IMD_FAQ4_H1", "Sub-heading.", "Diseases Curable with BMT")
add_label("IMD_FAQ4_LIST")
add_hint("Diseases bullets.")
add_bullets([
    "Hurlers Syndrome",
    "Morteaux-Lamy Syndrome",
    "Childhood onset cerebral X-linked Leukodystrophy",
    "Globoid Cell Leukodystrophy",
    "Metachromatic leukodystrophy",
    "α-Mannosidosis",
    "Osteopetrosis etc.",
])
add_block("IMD_FAQ4_A2", "Paragraph.",
    "Who is the best donor? As there are rarely matched unaffected siblings, alternate donors are the best options. In certain conditions like Hurler's syndrome, Cord Blood Transplants have given excellent results. However, HAPLOIDENTICAL TRANSPLANTS are now providing nearly 100% chance of finding a donor.")

add_h3("Bottom CTA Bar")
add_block("IMD_CTA_TITLE", "CTA heading.", "When to see us for Inherited Metabolic Disorders")
add_block("IMD_CTA_SUBTITLE", "CTA subtitle.",
    "Early BMT before significant neurological damage is critical for best outcomes.")
add_block("IMD_CTA_BUTTON", "CTA button.", "Book Consultation")

add_page_break()


# ============================================================
# CONDITION: INHERITED BONE MARROW FAILURE (IBMFS)
# ============================================================
add_h1("Condition — Inherited Bone Marrow Failure Syndromes (IBMFS)")
add_hint("Sub-page under Conditions. Single-tab page.")

add_block("IBMFS_OVERLINE", "Overline.", "Pediatric BMT — Dr. Suparno Chakrabarti")
add_block("IBMFS_TITLE", "Page title.", "Inherited Bone Marrow Failure Syndromes (IBMFS)")

add_h3("FAQ — What are Inherited Bone Marrow Failure Syndromes?")
add_block("IBMFS_FAQ1_Q", "Accordion question.", "What are Inherited Bone Marrow Failure Syndromes?")
add_block("IBMFS_FAQ1_A1", "Paragraph 1.",
    "The inherited bone marrow failure syndromes are a heterogeneous group of disorders characterized by bone marrow failure usually in association with one or more physical abnormalities. It is often present in childhood but may not present till adulthood in some patients.")
add_block("IBMFS_FAQ1_A2", "Paragraph 2.",
    "Most of the disorders are Autosomal Recessive (needs defective gene from both parents) but few are Autosomal Dominant (needs defective gene from any one parent) or X-Linked (only the male child will develop the disease).")
add_block("IBMFS_FAQ1_H1", "Sub-heading.", "Signs and Symptoms")
add_block("IBMFS_FAQ1_A3", "Paragraph.",
    "The symptoms are usually in the form of unilineage (only anemia or Neutropenia or thrombocytopenia) or multilineage cytopenias and infections. Many patients have physical abnormalities like short stature, darkening of skin, bone abnormalities, heart defects.")
add_label("IBMFS_FAQ1_LIST")
add_hint("Sign and symptom bullets.")
add_bullets([
    "Bleeding and Bruising",
    "Blood in gums, nose or the skin, tending to last longer than normal",
    "Blood in urine or stools",
    "Tooth loss or tooth decay",
    "Feeling tired most of the time",
    "Shortness of breath",
    "Recurring cold",
])

add_h3("FAQ — Classification of IBMFS")
add_block("IBMFS_FAQ2_Q", "Accordion question.", "Classification of IBMFS")
add_block("IBMFS_FAQ2_FA", "Paragraph.",
    "FANCONI ANEMIA (FA): Children with FA have characteristic physical abnormalities. FA is characterised by the defect in DNA repair leading to extensive chromosomal breakage. Affected individuals have a very high predisposition to develop Leukemia and Myelodysplastic Syndrome.")
add_block("IBMFS_FAQ2_DKC", "Paragraph.",
    "DYSKERATOSIS CONGENITA (DKC): DKC is characterised by the triad of abnormal nails, reticular skin pigmentation and oral leukoplakia. These patients have very short telomeres and have a high risk of developing lung fibrosis and liver disease.")
add_block("IBMFS_FAQ2_SDS", "Paragraph.",
    "SHWACHMAN-DIAMOND SYNDROME (SDS): It includes bone marrow failure with exocrine pancreatic insufficiency, short stature, metaphyseal dysostosis, and a predisposition to leukemia.")
add_block("IBMFS_FAQ2_SCN", "Paragraph.",
    "SEVERE CONGENITAL NEUTROPENIA (SCN) INCLUDING KOSTMANN SYNDROME: SCN is characterized by profound peripheral neutropenia. Patients usually present with severe bacterial infections.")
add_block("IBMFS_FAQ2_CAMT", "Paragraph.",
    "CONGENITAL AMEGAKARYOCYTIC THROMBOCYTOPENIA (CAMT): CAMT usually presents in infancy and is characterized by isolated thrombocytopenia and a reduction or absence of megakaryocytes in the bone marrow.")
add_block("IBMFS_FAQ2_DBA", "Paragraph.",
    "DIAMOND-BLACKFAN ANEMIA (DBA): DBA usually presents in early infancy with features of anemia. The hallmark is a profound reduction in erythroid precursors in the bone marrow.")

add_h3("FAQ — Diagnosis")
add_block("IBMFS_FAQ3_Q", "Accordion question.", "Diagnosis")
add_block("IBMFS_FAQ3_A1", "Paragraph.",
    "Complete Blood Count: The characteristic findings are unilineage or multilineage cytopenia which varies with the type of disease.")
add_block("IBMFS_FAQ3_A2", "Paragraph.",
    "Bone Marrow Study: Mainly marrow biopsy is done to confirm the hypocellularity of the marrow.")
add_block("IBMFS_FAQ3_A3", "Paragraph.",
    "Flow Cytometry: Often performed to detect early changes of leukemia and Myelodysplastic syndrome.")
add_block("IBMFS_FAQ3_A4", "Paragraph.",
    "Cytogenetics: Testing for abnormalities in the chromosomes. Certain genes are mutated which leads to complete absence of marrow cells.")
add_block("IBMFS_FAQ3_A5", "Paragraph.",
    "Antenatal Screening: For mothers who have already had a child with IBMFS, prenatal testing can be done through Chorionic Villus Sampling or Amniocentesis.")

add_h3("FAQ — Treatment & Role of BMT")
add_block("IBMFS_FAQ4_Q", "Accordion question.", "Treatment & Role of BMT")
add_block("IBMFS_FAQ4_A1", "Paragraph.",
    "Definitive cure is generally only achieved by ALLOGENEIC BMT; some disease may respond with steroids, but the disease reappears once they are withdrawn.")
add_block("IBMFS_FAQ4_A2", "Paragraph.",
    "When is BMT needed? Best results are obtained when BMT is carried out at the earliest, before the onset of leukemia, Myelodysplastic syndrome or too many blood transfusions.")
add_block("IBMFS_FAQ4_A3", "Paragraph.",
    "Conditioning: Reduced Intensity Conditioning with low doses of chemotherapy or radiation is the preferred option. Otherwise these patients develop severe organ damage during the BMT procedure.")
add_block("IBMFS_FAQ4_A4", "Paragraph.",
    "Donor: Although we prefer a matched family donor, a Half Matched (Haploidentical) family donor or a Matched Unrelated donor provide excellent survival. As BMT is required as an emergency procedure in this condition, Haploidentical Family Donor is often the quickest way of finding a donor.")
add_block("IBMFS_FAQ4_HIGHLIGHT", "Highlighted box.",
    "Results: With improved results with alternate donors and safer conditioning regimens, cure rates of 70-80% can be obtained if ALLOGENEIC BMT is performed in time.")

add_h3("Bottom CTA Bar")
add_block("IBMFS_CTA_TITLE", "CTA heading.", "When to see us for IBMFS")
add_block("IBMFS_CTA_SUBTITLE", "CTA subtitle.",
    "Early BMT before onset of leukemia or MDS offers 70-80% cure rates.")
add_block("IBMFS_CTA_BUTTON", "CTA button.", "Book Consultation")

add_page_break()


# ============================================================
# CONDITION OVERVIEW: NON-MALIGNANT
# ============================================================
add_h1("Conditions Overview — Non-Malignant")
add_hint("Sub-page under Conditions. Overview page covering non-malignant conditions.")

add_block("NONMAL_OVERLINE", "Overline.", "Pediatric — Dr. Suparno Chakrabarti")
add_block("NONMAL_TITLE", "Page title.", "Non-Malignant Conditions")

add_block("NONMAL_INTRO", "Intro paragraph.",
    "BMT (Bone marrow transplant) is also an effective therapeutic option for non-malignant conditions as well. Non-malignant disorders are such conditions which generally include metabolic disorders, immunodeficiency, autoimmune diseases and hemoglobinopathies. Given below is a brief information about various types of non-malignant conditions.")

add_h3("Sickle Cell Anemia and Thalassemia")
add_block("NONMAL_SCD_HEAD", "Sub-heading.", "Sickle Cell Anemia and Thalassemia")
add_block("NONMAL_SCD_P1", "Paragraph 1.",
    "Sickle Cell Anemia is actually an inherited disorder of RBCs (Red Blood cells). It is a genetic blood disease or disorder in which RBCs carrying oxygen in the body tends to develop abnormally. Instead of being flexible and round, the RBCs tend to develop in sickle or crescent shape. These abnormally developed red blood cells leads to clog blood vessel sections and this indeed leads to pain in the body. Frequent pain, frequent infections, vision problems and delayed growth are some of the major symptoms of sickle cell anemia that can have major impact on one's quality of life.")
add_block("NONMAL_SCD_P2", "Paragraph 2.",
    "Thalassemia is yet another blood disease or disorder in which a person's body tends to make abnormal hemoglobin form. This blood disorder results in extreme destruction of RBCs that indeed leads to severe anemia. Bone deformities in face, growth failure, fatigue, jaundice and breathing difficulty are some major symptoms of Thalassemia.")

add_h3("Inherited Disorders of Bone Marrow Failure")
add_block("NONMAL_IBMF_HEAD", "Sub-heading.", "Inherited Disorders of Bone Marrow Failure")
add_block("NONMAL_IBMF_P1", "Paragraph.",
    "Genetic bone marrow failure diseases are some other non-malignant conditions that should be cured by choosing effective treatment options.")
add_block("NONMAL_IBMF_FA", "Paragraph.",
    "Fanconi Anemia — a recessive genetic disorder which results in decreased production of all types of blood cells. Birth defect, growth problems, blasts (abnormal WBCs) and failure of bone marrow are the symptoms.")
add_block("NONMAL_IBMF_DKC", "Paragraph.",
    "Dyskeratosis Congenita — sometimes also called Zinsser-Engman-Cole syndrome, is a progressive and rare bone marrow disorder. Discoloration of skin, abnormal nails and white patches inside the mouth are some of the common symptoms.")
add_block("NONMAL_IBMF_BDS", "Paragraph.",
    "Blackfan-Diamond Syndrome — also known as diamond-blackfan anemia, is actually a congenital erythroid aplasia generally diagnosed in children. In this disorder, bone marrow does not produce enough RBCs.")
add_block("NONMAL_IBMF_SDS", "Paragraph.",
    "Swachman-Diamond Syndrome — SDS is a very rare congenital disease characterized by bone marrow failure, exocrine pancreatic insufficiency and skeletal abnormalities.")
add_block("NONMAL_IBMF_AMT", "Paragraph.",
    "Amegakaryocytic Thrombocytopenia — a disorder in which the body suffers from blood platelet deficiency. Easy bruising, impaired clotting and abnormal bleeding are its symptoms.")

add_h3("Primary Immunodeficiency Disorders")
add_block("NONMAL_PID_HEAD", "Sub-heading.", "Primary Immunodeficiency Disorders")
add_block("NONMAL_PID_P1", "Paragraph.",
    "Primary-immunodeficiency disorders are those diseases in which body's immunity system parts do not function correctly or are even missing. They usually include recurrent infections, poor growth or development, unexplained weight loss, recurrent pneumonia and recurrent skin abscesses.")

add_h3("Inherited Metabolic Disorders")
add_block("NONMAL_IMD_HEAD", "Sub-heading.", "Inherited Metabolic Disorders")
add_block("NONMAL_IMD_P1", "Paragraph.",
    "Inherited-metabolic disorders are those genetic conditions which result in problems of metabolism. In most of the metabolic disorders, enzymes required for metabolic processes are missing or do not function properly.")

add_h3("Bottom CTA Bar")
add_block("NONMAL_CTA_TITLE", "CTA heading.", "Consult our experts for Non-Malignant Blood Disorders")
add_block("NONMAL_CTA_SUBTITLE", "CTA subtitle.",
    "Many non-malignant conditions are curable with BMT. Early referral improves outcomes significantly.")
add_block("NONMAL_CTA_BUTTON", "CTA button.", "Book Consultation")

add_page_break()


# ============================================================
# CONDITION OVERVIEW: MALIGNANT
# ============================================================
add_h1("Conditions Overview — Malignant")
add_hint("Sub-page under Conditions. Overview page covering malignant conditions.")

add_block("MAL_OVERLINE", "Overline.", "Pediatric — Dr. Suparno Chakrabarti")
add_block("MAL_TITLE", "Page title.", "Malignant Conditions")

add_block("MAL_P1", "Paragraph 1.",
    "Malignancy is a tendency of some medical conditions, mainly tumors, that as time passes, started becoming more and more worse, and probably results in death. This is generally familiar to the cancer. Generally, malignant word is used to define the conditions which are getting inferior. For the cancer, it normally means that the tumor is increasing or growing which affects the nearby situated healthy tissues and which it might metastasize. There are several malignant conditions that can affect one or more systems of the body.")

add_h3("Understanding Malignant Conditions")
add_block("MAL_UNDERSTAND_HEAD", "Sub-heading.", "Understanding Malignant Conditions")
add_block("MAL_UNDERSTAND_P1", "Paragraph.",
    "Before assessing the malignant conditions, all the information must be available regarding the extent of this condition. Such information contains the site of primary and any secondary, results of any kind of surgery and also the effect of chemotherapy or radiotherapy.")
add_block("MAL_UNDERSTAND_P2", "Paragraph.",
    "A malignant tumor is the one which is persistent and it may reach to the other parts of body. Those tumors which remains localized and do not spread out are known as benign. Though the term malignant is frequently used to describe the conditions of cancer, not all the malignant conditions are cancerous.")

add_h3("Blood Cancers Treated at Bloods <em>R</em> Us")
add_block("MAL_CANCERS_HEAD", "Sub-heading.", "Blood Cancers Treated at Bloods <em>R</em> Us")
add_block("MAL_CANCERS_P1", "Paragraph.",
    "Diseases like acute lymphoblastic leukemia is a severe kind of blood cancer. It starts generating in the white blood cells within the bone marrow. Transplantation of Bone Marrow permits doctors to cure cancer with high chemotherapy doses by allowing replacing the cells of bone marrow which get destroyed during the treatment.")
add_block("MAL_CANCERS_P2", "Paragraph.",
    "Diseases that can be and have been successfully treated with the help of BMT include: solid tumors, leukemia (ALL, AML, CML), multiple myeloma, and lymphoma (Hodgkin's and Non-Hodgkin's).")
add_block("MAL_CANCERS_P3", "Paragraph.",
    "Patients suffering with malignant disorders can be cured by autologous or allogeneic bone marrow transplantation.")

add_h3("Bottom CTA Bar")
add_block("MAL_CTA_TITLE", "CTA heading.", "Expert care for Blood Cancers")
add_block("MAL_CTA_SUBTITLE", "CTA subtitle.",
    "Our team has treated 250+ BMT patients with 75% long-term survival in high risk blood cancers.")
add_block("MAL_CTA_BUTTON", "CTA button.", "Book Consultation")

add_page_break()


# ============================================================
# BMT KNOWLEDGE CENTRE PAGE
# ============================================================
add_h1("BMT Centre Page")
add_hint("Dedicated page that explains what Bone Marrow Transplantation is, types, who needs it, and the stages involved.")

add_h2("Page Hero")
add_block("BMT_HERO_OVERLINE", "Small label.", "Education")
add_block("BMT_HERO_TITLE", "Page title.", "BMT Knowledge Centre")
add_block("BMT_HERO_SUBTITLE", "Subtitle paragraph.",
    "Comprehensive information about Bone Marrow Transplantation — what it is, how it works, and who needs it. Written by Dr. Suparno Chakrabarti.")

add_h2("What is Bone Marrow Transplantation? (section)")
add_block("BMT_WHAT_OVERLINE", "Small label.", "Foundation")
add_block("BMT_WHAT_TITLE", "Section title.", "What is Bone Marrow Transplantation?")
add_block("BMT_WHAT_P1", "Paragraph.",
    "Bone Marrow Transplantation is a procedure, where healthy stem cells are transplanted into a patient's body after appropriate treatment. Healthy stem cells can be collected from the patient, when he becomes disease free after treatment. They can also be collected from fully / half HLA matched family donor, unrelated donor and cord blood. Stored stem cells of the patient / donor / cord blood are transfused into patient's blood stream, quite similar to blood transfusion.")
add_block("BMT_WHAT_H1", "Sub-heading.", "What are Stem Cells?")
add_block("BMT_WHAT_P2", "Paragraph.",
    "Stem cells are the most primitive cells which can differentiate into various other dedicated cells, such as nerve cells, bone cells, liver cells, blood cells etc. The most primitive stem cell gets committed to organ-specific stem cells and thus forms either blood stem cells or nerve stem cells.")
add_block("BMT_WHAT_H2", "Sub-heading.", "Sources of Blood Stem Cells")
add_block("BMT_WHAT_SOURCES_BM", "Paragraph.",
    "Bone Marrow: The spongy material inside large bones of adults and all bones of children is called bone marrow, which is a normal reserve for blood stem cells. Depending on body's requirement, stem cells can produce desired number of Red blood cells, white blood cells and Platelets. When a patient needs blood stem cells, they can be collected from healthy donor's hip bones under anesthesia.")
add_block("BMT_WHAT_SOURCES_PB", "Paragraph.",
    "Peripheral Blood: Normally, there is only an occasional blood stem cell in our circulation. However, when a blood growth factor (G-CSF) is injected, it stimulates the stem cells from the bone marrow to spill over in the circulation. If we carry out a procedure called Apheresis (similar to dialysis), stem cells can be collected from the peripheral blood.")
add_block("BMT_WHAT_SOURCES_CORD", "Paragraph.",
    "Umbilical Cord Blood: Placenta along with the umbilical cord are waste products of pregnancy. However, in the mid-1980s, it was discovered that they are a rich source of blood stem cells. These cells can be collected after birth and stored for later use in a patient.")

add_h2("Types of Bone Marrow Transplantation (two cards)")
add_block("BMT_TYPES_OVERLINE", "Small label.", "Types")
add_block("BMT_TYPES_TITLE", "Section title.", "Types of Bone Marrow Transplantation")

add_block("BMT_TYPE1_LABEL", "Card label above title.", "Type 1")
add_block("BMT_TYPE1_TITLE", "Card title.", "Allogeneic BMT")
add_block("BMT_TYPE1_BODY", "Card body paragraph.",
    "The blood stem cells are obtained from the peripheral blood or bone marrow of a donor who is suitably matched to the patient. The diseased or failing bone marrow is replaced by healthy marrow from a normal donor. This process cures the underlying disease.")

add_block("BMT_TYPE2_LABEL", "Card label above title.", "Type 2")
add_block("BMT_TYPE2_TITLE", "Card title.", "Autologous BMT")
add_block("BMT_TYPE2_BODY", "Card body paragraph.",
    "Standard dose of chemotherapy cannot always cure a cancer. High dose is often needed. However, such high doses damage the patient's bone marrow irreversibly. Patient's own blood stem cells are collected before administering high dose chemotherapy and stored. The stored blood stem cells of the patient are transfused back after Chemotherapy / Radiotherapy.")

add_h2("Conditions Treated with BMT")
add_block("BMT_COND_OVERLINE", "Small label.", "Indications")
add_block("BMT_COND_TITLE", "Section title.", "Conditions Treated with BMT")
add_block("BMT_COND_P1", "Paragraph.",
    "Transplantation of Bone Marrow, also termed as Stem Cell Transplant, is a complex process where the damaged or infected bone marrow is replaced by healthy ones. This process is implemented once the patient has been treated with high doses of chemotherapy and radiation treatment.")
add_block("BMT_COND_P2", "Paragraph.",
    "Bone marrow is a soft, spongy tissue within the bones which produces cells together with white and red blood cells along with platelets. When this bone marrow gets injured, it is no longer capable of producing these red and white blood cells. As a result, you may experience some severe consequences like infections, weakness, anemia, excessive bleeding and possibly even death.")
add_block("BMT_COND_H1", "Sub-heading.", "Who Needs Bone Marrow Transplantation?")
add_block("BMT_COND_P3", "Paragraph.",
    "Bone Marrow is an amazing organ spread all through the body, producing millions of cells every moment of our life to keep us well and healthy. Any disease affecting the bone marrow affects our entire body. Replacing the diseased or failing bone marrow with healthy marrow stem cells is the process of Bone Marrow Transplantation.")
add_block("BMT_COND_BLOOD_CANCERS", "Paragraph.",
    "Blood Cancers: Any blood cancer (leukemia) or lymph gland cancer (Lymphoma) which is not completely cured with chemotherapy or recurs after completion of chemotherapy (relapse), can be cured with Allogeneic or Autologous BMT in about half of the patients.")
add_block("BMT_COND_GENETIC", "Paragraph.",
    "Thalassemia and Other Genetic Conditions: In these conditions, the defective bone marrow cells can be killed by chemotherapy and replaced by healthy bone marrow cells from a healthy donor. This can be cured in about 90% of the patients.")
add_block("BMT_COND_APLASTIC", "Paragraph.",
    "Aplastic Anaemia and Related Conditions: In these conditions, the bone marrow does not produce enough stem cells and healthy stem cells can be transplanted from a healthy donor. This is also curative in about 70-90% of patients.")
add_block("BMT_COND_OTHER", "Paragraph.",
    "Other Cancers: Many other cancers which do not arise from the bone marrow can be cured by infusing patient's own stem cells after high dose chemotherapy.")

add_h2("BMT Process — Four Stages")
add_block("BMT_PROCESS_OVERLINE", "Small label.", "The Journey")
add_block("BMT_PROCESS_TITLE", "Section title.", "The BMT Process — Four Stages")
add_block("BMT_PROCESS_BODY", "Paragraph under the title.",
    "Dr. Suparno Chakrabarti explains the four stages involved in Bone Marrow Transplantation.")

add_h3("Stage 1 — Evaluation")
add_block("BMT_STAGE1_TITLE", "Stage heading.", "Evaluation (Work-Up) — Usually 14–30 days before")
add_block("BMT_STAGE1_P1", "Intro paragraph.",
    "One will undergo complete medical check-up to evaluate one's suitability to go through the BMT procedure. This involves the following:")
add_label("BMT_STAGE1_LIST")
add_hint("Work-up bullets.")
add_bullets(["Blood Tests", "Chest X-ray and CT Scans", "Tests to assess the condition of heart and lungs", "Bone Marrow Tests"])
add_block("BMT_STAGE1_P2", "Closing paragraph.",
    "Patients will be counselled in detail about the procedure, the complications, the chances of success, the cost and the possible length of stay in the hospital. Patient will be encouraged to go through the educational material/booklet and discuss any queries or doubts that he/she might have.")

add_h3("Stage 2 — Conditioning")
add_block("BMT_STAGE2_TITLE", "Stage heading.", "Conditioning — Usually 2–10 days")
add_block("BMT_STAGE2_P1", "Paragraph.",
    "High dose of chemotherapy or radiotherapy is given to destroy the diseased marrow or destroy the cancer cells and make space for the new bone marrow cells. It also suppresses the immunity of the patient so that the new bone marrow is not rejected.")
add_block("BMT_STAGE2_P2", "Paragraph — 'The Transplant Procedure' label bold on website.",
    "The Transplant Procedure: The transplant procedure is actually fairly simple — the stem cells or bone marrow cells to be transplanted are given as a blood transfusion through the central line. This usually takes from one to several hours. Just before the infusion, the patient may be given medication to help avoid any reaction. A monitor is used to check breathing, heart rate and blood pressure during the procedure.")

add_h3("Stage 3 — Pre-Engraftment")
add_block("BMT_STAGE3_TITLE", "Stage heading.", "Pre-Engraftment — Usually 2–3 weeks")
add_block("BMT_STAGE3_P1", "Paragraph.",
    "After high dose chemo-radiotherapy the blood stem cells are destroyed and normal blood cells are not produced. The patients need to be kept in a clean room within the BMT unit in strict isolation during this time. They also need a lot of blood and platelet transfusion. Most patients get serious infections during this period and need treatment with antibiotics.")

add_h3("Stage 4 — Post-Engraftment")
add_block("BMT_STAGE4_TITLE", "Stage heading.", "Post-Engraftment — Months to Years")
add_block("BMT_STAGE4_EARLY", "Paragraph — 'Early Phase' label bold.",
    "Early Phase (first 3 months): There are two types of white blood cells: neutrophils and lymphocytes. Once the neutrophil count is above the critical value of 500 cells per microlitre, the patient can come out of critical isolation. This is called engraftment — the first sign that the transplanted blood stem cells are functioning. There is also a risk of graft-versus-host disease (GVHD) at this stage.")
add_block("BMT_STAGE4_LATE", "Paragraph — 'Late Phase' label bold.",
    "Late Phase (3 months–12 months): The immunity against viruses takes a very long time to recover. Even though some immunity is restored, the patient is still at risk of infections with viruses and fungus. This is more so if they are being treated for GVHD, which can become chronic and lingering. If the patient is well, the frequency of check-ups and blood tests reduce over several months.")

add_h2("Major Complications After BMT")
add_block("BMT_COMP_OVERLINE", "Small overline label for the complications section.", "Major Complications")
add_block("BMT_COMP_TITLE", "Section title.", "Major Complications After BMT")
add_block("BMT_COMP_INTRO", "Intro paragraph for the complications list.",
    "While BMT is carried out to cure the underlying cancer or blood disorder, it remains one of the most challenging procedures in modern medicine due to its unpredictability. The physicians and scientists involved in the field of BMT have toiled for the last five decades to understand the intricacies of this apparently simple but immunologically complex procedure. With best of efforts to understand and intervene and/or often preempt, certain complications remain unavoidable after an allogenic bone marrow transplantation which are listed below.")

add_h3("Card 1 — Graft Failure & Dysfunction")
add_block("BMT_COMP_CARD1_OVERLINE", "Small label.", "Graft Outcomes")
add_block("BMT_COMP_CARD1_TITLE", "Card title.", "Graft Failure & Dysfunction")
add_label("BMT_COMP_CARD1_ITEMS")
add_hint("List items in this card. Body copy for each is intentionally blank — doctors can fill in via CMS later.")
add_bullets(["Graft Failure", "Graft Dysfunction", "Relapse"])

add_h3("Card 2 — Graft vs Host Disease")
add_block("BMT_COMP_CARD2_OVERLINE", "Small label.", "Immune Reactions")
add_block("BMT_COMP_CARD2_TITLE", "Card title.", "Graft vs Host Disease (GvHD)")
add_label("BMT_COMP_CARD2_ITEMS")
add_hint("List items in this card.")
add_bullets(["Acute Graft vs Host Disease", "Chronic Graft vs Host Disease"])

add_h3("Card 3 — Post-Transplant Infections")
add_block("BMT_COMP_CARD3_OVERLINE", "Small label.", "Infections")
add_block("BMT_COMP_CARD3_TITLE", "Card title.", "Post-Transplant Infections")
add_label("BMT_COMP_CARD3_ITEMS")
add_hint("Infection types — body copy to be filled per type later.")
add_bullets(["Bacterial", "Viral", "Fungal"])

add_h3("Card 4 — Hemophagocytosis")
add_block("BMT_COMP_CARD4_OVERLINE", "Small label.", "Inflammatory Storms")
add_block("BMT_COMP_CARD4_TITLE", "Card title.", "Hemophagocytosis")
add_label("BMT_COMP_CARD4_ITEMS")
add_hint("List items in this card.")
add_bullets(["Post Transplant Hemophagocytosis", "Post Transplant Hemophagocytic Syndrome"])

add_h3("Card 5 — Organ & Endothelial Toxicity")
add_block("BMT_COMP_CARD5_OVERLINE", "Small label.", "Conditioning Toxicity")
add_block("BMT_COMP_CARD5_TITLE", "Card title.", "Organ & Endothelial")
add_label("BMT_COMP_CARD5_ITEMS")
add_hint("List items in this card. Endothelial Dysfunction has two sub-items: VOD and TA-TMA.")
add_bullets([
    "Conditioning-related Organ Toxicity",
    "Endothelial Dysfunction",
    "    — Veno-Occlusive Disease",
    "    — Transplant-Associated Thrombotic Microangiopathy",
])

add_h3("Card 6 — Late Sequelae")
add_block("BMT_COMP_CARD6_OVERLINE", "Small label.", "Long-Term Effects")
add_block("BMT_COMP_CARD6_TITLE", "Card title.", "Late Sequelae")
add_label("BMT_COMP_CARD6_ITEMS")
add_hint("Long-term complications — body copy per item to be filled later.")
add_bullets(["Secondary Malignancies", "Reproductive Dysfunction and Infertility", "Endocrine Dysfunction"])

add_page_break()


# ============================================================
# PAGE: NOVEL CELLULAR THERAPY (#cellular-therapy)
# ============================================================
add_h1("Novel Cellular Therapy Page")
add_hint("Cellular therapies developed by our group: AbaNI-15, MAVIcel, REGRAcel, ECP. Body copy by Dr. Suparno Chakrabarti & Dr. Mahak Agarwal.")

add_h2("Page Hero")
add_block("CT_HERO_OVERLINE", "Small label above the title.", "Developed by Our Group")
add_block("CT_HERO_TITLE", "Page title.", "Novel Cellular Therapy")
add_block("CT_HERO_SUBTITLE", "Subtitle paragraph.",
    "Original cellular therapies developed by our team — exploring the natural laws of immunity in the fight against cancer, without genetic manipulations.")

add_h2("Programme Intro")
add_block("CT_PROG_OVERLINE", "Small label.", "Our Research Programme")
add_block("CT_PROG_TITLE", "Section title.", "Built from Bench to Bedside")
add_block("CT_PROG_BODY", "Programme intro paragraph.",
    "Our team has been engaged in both clinical and laboratory research to understand how we might explore the natural laws of immunity in the fight against cancer, without genetic manipulations. In this process we have developed several novel forms of cellular therapy employing four distinct cell types — each branded as a proprietary protocol of our group.")

add_h2("Protocol 1 — AbaNI-15 (Adaptive NK Cell Therapy)")
add_block("CT_ABANI_OVERLINE", "Small label.", "Protocol 1 — NK Cell Based Immunotherapy")
add_block("CT_ABANI_TITLE", "Protocol title.", "AbaNI-15 — Adaptive NK Cell Therapy")
add_block("CT_ABANI_P1", "Paragraph — what AbaNI-15 is.",
    "AbaNI-15 is our proprietary Natural Killer (NK) cell therapy protocol developed by our team in 2022. NK cells are a unique population of lymphocytes (a type of strong and potent cell of the immune system) which are capable of defending us against both viral infections as well as cancers. They comprise about 5–10% of all the lymphocytes circulating in the peripheral blood.")
add_block("CT_ABANI_P2", "Paragraph — how NK cells work vs T cells.",
    "Unlike T cells which comprise 60–70% of the lymphocytes, NK cells do not need to recognise each virus or each cancer cell by their respective signatures (antigens). Because of their very strong killing potential, they are kept under constant inhibitory control. Under specific circumstances, the NK cells get activated and kill the target cells which might be cancerous or infected with viruses, with a ferocity barely encountered in any other human cells. However, the NK cells are short lived (survive for about 2 weeks in the circulation) and have to be in a constant state of production to maintain the necessary protection for the human body.")
add_block("CT_ABANI_P3", "Paragraph — the Adaptive NK cell discovery.",
    "Along with a few groups in the US and Europe, our group identified a unique population of NK cells which are strong killer cells but unlike the conventional NK cells survive in the circulation for a very long time (years). These cells are not naturally produced but develop when the body is faced with a particular virus called Cytomegalovirus (CMV). Almost 95% of the Indian population has been exposed to this virus and hence are naturally endowed with the capability of producing these strong, long-lived, killer NK cells (also known as Adaptive NK cells).")
add_block("CT_ABANI_P4", "Paragraph — discovering the Abatacept × NK effect, leading to AbaNI-15.",
    "Our group has been working on these Adaptive NK cells with respect to Cancers and BMT as well as Covid-19 infection. While exploring Abatacept for prevention of GvHD in Haploidentical transplants, we discovered that abatacept not only has a sparing effect on NK cells but in-fact promotes the killer function of NK cells particularly the adaptive type. Following a series of lab experiments carried out over 5 years, exploring the effects of Abatacept and various cytokines on NK cells, we have developed a novel NK cell product called AbaNI-15.")
add_block("CT_ABANI_P5", "Paragraph — 12-patient compassionate use program.",
    "On a compassionate use named patient program, AbaNI-15 has been used in 12 patients with refractory or resistant Acute Leukemia undergoing Haploidentical Transplantation. The early results are very encouraging. Unlike CAR-T cell therapy, this treatment has not been found to be associated with any major side effects.")
add_block("CT_ABANI_P6", "Closing italic note.",
    "This treatment is available at our institution under a compassionate use named patient program.")

add_h2("Protocol 2 — MAVIcel (Mesenchymal Stem Cell Therapy)")
add_block("CT_MAVI_OVERLINE", "Small label.", "Protocol 2 — Mesenchymal Stem Cell Based Therapy")
add_block("CT_MAVI_TITLE", "Protocol title.", "MAVIcel — MSCs for GvHD & Inflammatory Conditions")
add_block("CT_MAVI_P1", "Paragraph — MSC biology and immunomodulation.",
    "Bone marrow stem cells rest within a nest of other supportive cells, the primary component of which is mesenchymal stem cells or mesenchymal stromal cells (MSC). These cells have the unique ability of differentiating into other cell lineages such as osteoblasts (bone), chondrocytes (cartilage) and adipocytes (fat). Under certain conditions, these cells can also differentiate into various other cell lineages and due to this reason, they have been a special cell of interest in regenerative medicine. At the same time, MSCs also possess a unique immunosuppressive and immunomodulating property which have made them an attractive proposition for inflammatory conditions where the immune system is hyperactive.")
add_block("CT_MAVI_P2", "Paragraph — MSC use in steroid-refractory GvHD.",
    "One such condition in relation to BMT is the dreaded complication of severe GvHD. In patients who fail to respond to steroids and other therapies for severe GvHD, MSCs have often been found to be effective in reducing the severity of symptoms and at times result in complete remission of GvHD.")
add_block("CT_MAVI_P3", "Paragraph — MAVIcel manufacturing process.",
    "MSCs can be sourced from many tissues which include bone marrow, umbilical cord blood, placental tissue and adipose tissue. Our group has developed a distinctive process of manufacturing of MSC for treatment of GvHD sourced from bone marrow of multiple healthy donors. This proprietary product is named MAVIcel.")
add_block("CT_MAVI_P4", "Paragraph — 10-patient compassionate use experience.",
    "MAVIcel has been used on a compassionate named patient program in over 10 patients. This has been found to be highly effective in control of acute and chronic GvHD, with little or no side effects.")
add_block("CT_MAVI_P5", "Closing italic note about ongoing clinical study.",
    "This product is currently being explored in an institutionally approved clinical study for patients with GvHD and inflammatory conditions.")

add_h2("Protocol 3 — REGRAcel (Regulatory T Cell Therapy)")
add_block("CT_REGRA_OVERLINE", "Small label.", "Protocol 3 — Regulatory T Cell Therapy")
add_block("CT_REGRA_TITLE", "Protocol title.", "REGRAcel — Stable Tregs for Aggressive GvHD")
add_block("CT_REGRA_P1", "Paragraph — T cell biology and Tregs.",
    "T lymphocytes are the dominant immune cells in the human body which are extremely sophisticated weapons used against viruses or cancers as and when needed. They are very long lived cells whose production and functions are heavily regulated so that they do not attack normal cells or tissues in the body. This regulation is governed by a specific type of T cells called Regulatory T Cells (Treg) which control excess T cell activity in both health and disease.")
add_block("CT_REGRA_P2", "Paragraph — why GvHD happens and the Treg challenge.",
    "GvHD, the most dreaded complication of Allogenic BMT, happens due to imbalance of the attacking T cells and the Tregs. While all the existing treatment of GvHD is geared towards controlling or eliminating the T cells which are attacking the patient's body, there is no effective therapy to restore the deficit of the Tregs. In addition, under situations of extreme inflammation, Tregs might lose their regulatory capability and turn rogue. This makes reliance on Tregs for controlling situations like GvHD less feasible. Unlike infusion of other cellular components like conventional T cells, NK cells or MSCs, infusion of Tregs has not been established in clinical practice despite its pivotal role in correcting the immune imbalance.")
add_block("CT_REGRA_P3", "Paragraph — REGRAcel: stable Treg subpopulation, ongoing exploration.",
    "Our group has been working on identifying a stable population of Tregs which are less infidel in extremely stressful conditions. This particular subpopulation of Treg might be useful in patients with aggressive GvHD. Based on these findings, we have developed a cellular product named REGRAcel. Clinical exploration of this product is ongoing. We hope to initiate a clinical study in the near future.")

add_h2("Protocol 4 — ECP (Extracorporeal Photopheresis)")
add_block("CT_ECP_OVERLINE", "Small label.", "Protocol 4 — Extracorporeal Photopheresis")
add_block("CT_ECP_TITLE", "Protocol title.", "ECP — Light-Based Lymphocyte Modulation")
add_block("CT_ECP_P1", "Paragraph — historical origin (ancient India/Egypt + vitiligo + Ammi majus / 8-MOP).",
    "ECP (Extracorporeal Photopheresis) dates back to ancient India and Egypt, where people with vitiligo ingested a plant (Ammi majus) found on the banks of the river, bathed in the sun, and noticed recovery in melanin production. Psoralen (8-methoxypsoralen [8-MOP]) is a photoreactive substance isolated from these plants.")
add_block("CT_ECP_P2", "Paragraph — what ECP does and the safety note.",
    "ECP is defined as a technique of manipulating white blood cells (lymphocytes) outside the body in a way that, when they are re-infused to the patient, cause downregulation of lymphocytes (majorly T-lymphocyte activity) in patients. This process involves collection of mononuclear cells from the patient's blood, followed by addition of Psoralen and irradiation with UVA before reinfusion. As the irradiation is done outside the body, the risk of phototoxicity is negligible — however, it is recommended to use sunscreen and photoprotective sunglasses while on treatment.")
add_block("CT_ECP_P3", "Paragraph — ECP for GvHD, both acute and chronic.",
    "This technique is employed in many conditions where the lymphocytes in the blood of an individual become hyperactive and cause excess inflammation. One such condition related to BMT is GvHD. ECP has been found to be successful in dampening the severity of GvHD in situations where it is non-responsive to standard drug-based therapies. This applies to both acute and chronic GvHD.")
add_block("CT_ECP_INDICATIONS_TITLE", "Sub-heading for indications list.", "Approved Indications")
add_label("CT_ECP_INDICATIONS")
add_hint("Bullet list of approved indications.")
add_bullets(["GvHD", "Cutaneous T Cell Lymphoma", "Solid organ transplant rejection", "Certain autoimmune diseases"])
add_block("CT_ECP_HOW_TITLE", "Sub-heading for mechanism hypotheses.", "How ECP Works — Current Hypotheses")
add_block("CT_ECP_HOW_INTRO", "Intro paragraph above the mechanism bullets.",
    "How ECP tames the lymphocytes is not fully understood. There are several hypotheses which include:")
add_label("CT_ECP_HOW_BULLETS")
add_hint("Bullet list of proposed mechanisms.")
add_bullets([
    "Inducing death in the hyperactive lymphocytes by upregulating proteins responsible for cell death.",
    "Transforming the partner cells of T cells (antigen presenting cells) to a more mellowed and non-reactive state.",
    "Increasing the number and function of Regulatory T Cells which are inherently poised to control hyperactive T cells.",
])
add_block("CT_ECP_PROC_TITLE", "Sub-heading for procedure description.", "The Procedure")
add_block("CT_ECP_PROC_BODY", "Paragraph — what the procedure looks like and side-effect profile.",
    "ECP is not a one-off treatment for GvHD. Several such procedures (1–2 every week) are needed for several weeks or months. The collection of lymphocytes in the patient involves putting in a central line (particularly in children). The collected cells are then subjected to brief exposure to Psoralen and UVA therapy under sterile conditions and returned back to the patient as an IV infusion. The procedure is generally safe and no major side effects have been reported to date — apart from the inconvenience and side effects related to the central line and/or lymphopheresis procedure.")
add_block("CT_ECP_HIGHLIGHT", "Highlight callout — keep verbatim.",
    "We are one of the few centers in the country to provide this service.")
add_block("CT_ECP_CLOSING", "Closing paragraph about novel ECP combinations under exploration.",
    "In addition, we are exploring novel approaches of combining other forms of cellular therapies with ECP in treatment of GvHD and other autoimmune disorders.")

add_page_break()


# ============================================================
# PAGE: NOVEL AUTOLOGOUS BMT — IBAHCT (#autologous-bmt)
# ============================================================
add_h1("Novel Autologous BMT (IBAHCT) Page")
add_hint("Immune-Boosted Autologous Hematopoietic Cell Transplantation. Sits under 'Transplant Innovations' alongside the Haploidentical BMT page.")

add_h2("Page Hero")
add_block("IBAHCT_HERO_OVERLINE", "Small label.", "Transplant Innovations")
add_block("IBAHCT_HERO_TITLE", "Page title.", "Immune-Boosted Autologous Hematopoietic Cell Transplantation (IBAHCT)")
add_block("IBAHCT_HERO_SUBTITLE", "Subtitle paragraph.",
    "A novel protocol developed by our group — bringing the curative power of donor-driven immunity into the autologous BMT setting for advanced Lymphoma and Myeloma.")

add_h2("The Background — High-dose chemo with stem cell rescue")
add_block("IBAHCT_BG_OVERLINE", "Small label.", "The Background")
add_block("IBAHCT_BG_TITLE", "Section title.", "High-dose chemotherapy with stem cell rescue")
add_block("IBAHCT_BG_P1", "Paragraph — what autologous BMT/HCT is.",
    "Autologous BMT/HCT is also known as high dose chemotherapy with stem cell rescue. What that means is to be able to give a higher dose of chemotherapy with the idea and the intent that chemotherapy would eliminate the residual cancer. However, with such high doses, the bone marrow is the first to take a hit and is unable to rejuvenate on its own. To sustain life, then the patient requires a \"Stem Cell Rescue\" to repopulate the marrow. Hence, hematopoietic stem cells have to be collected before the chemotherapy is administered and frozen to preserve, and reinfused following the chemotherapy.")
add_block("IBAHCT_BG_P2", "Paragraph — the GvL/GvHD context vs autologous BMT.",
    "While this process is effective in curing a proportion of patients in Lymphoma and prolonging survival in patients with Myeloma, an allogenic transplantation from a compatible donor under similar circumstances has been found to cure 2–3 times more patients. This is primarily because the donor immune cells effectively eradicate the last cancer cell (GvL) in the patient's body. In a same way, these cells attack the normal cells and tissues in the patient's body (GvHD). This phenomenon is rarely been described in an Autologous BMT. Hence, the GvL effect is not appreciably noted in this process.")

add_h2("The Discovery — A GvHD-like phenomenon in Autologous BMT")
add_block("IBAHCT_DISC_OVERLINE", "Small label.", "The Discovery")
add_block("IBAHCT_DISC_TITLE", "Section title.", "A GvHD-like phenomenon in Autologous BMT")
add_block("IBAHCT_DISC_P1", "Paragraph — observation of GvHD-like phenomenon and the trigger agent investigation.",
    "However, we noticed a similar phenomenon akin to GvHD as seen following Allogenic BMT in certain patients of Lymphoma undergoing Autologous BMT. We explored the treatment process that these patients underwent prior to the BMT, and discovered a common agent that they had received. Studying the various mechanisms of action of this agent, we realised there were certain pathways by which this agent was boosting the immune system of the patient and possibly leading to the GvHD-like phenomenon that we noted.")

add_h2("The Protocol — Engineering an immune-boost into Autologous BMT")
add_block("IBAHCT_PROT_OVERLINE", "Small label.", "The Protocol")
add_block("IBAHCT_PROT_TITLE", "Section title.", "Engineering an immune-boost into Autologous BMT")
add_block("IBAHCT_PROT_P1", "Paragraph — the new protocol design and early results.",
    "Following this observation, we designed a protocol where low doses of this agent were employed before the collection of the stem cells and after Autologous BMT at certain fixed time points, in advanced patients with Lymphoma and Myeloma. The initial results are extremely promising.")
add_block("IBAHCT_PROT_HIGHLIGHT", "Highlight callout — keep verbatim.",
    "This therapy is being offered to selected patients in an ethical committee approved clinical study.")

add_page_break()


# ============================================================
# PAGE: FOR DOCTORS (#physicians)
# ============================================================
add_h1("For Doctors Page")
add_hint("Top-level page aimed at referring physicians / haematologists / oncologists. Contains the BMT referral table, publications, and referral contact block.")

add_h2("Page Hero")
add_block("PHYS_HERO_OVERLINE", "Small label above the page title.", "Referral Resource")
add_block("PHYS_HERO_TITLE", "Page title.", "For Doctors")
add_block("PHYS_HERO_SUBTITLE", "Subtitle paragraph.",
    "Clinical reference guide for haematologists, oncologists, and general physicians considering BMT referral.")

add_h2("When to Refer for BMT — Section Header")
add_block("PHYS_REFER_OVERLINE", "Small label above the section title.", "Clinical Reference")
add_block("PHYS_REFER_TITLE", "Section title.", "When to Refer for BMT")
add_block("PHYS_REFER_BODY", "Body paragraph under the section title.",
    "A consolidated reference of BMT indications organised by condition, extracted from Dr. Suparno Chakrabarti's clinical guidelines.")

add_h3("Referral Table — Column Headings")
add_block("PHYS_TABLE_COL_1", "Table column heading.", "Condition")
add_block("PHYS_TABLE_COL_2", "Table column heading.", "BMT Indication")
add_block("PHYS_TABLE_COL_3", "Table column heading.", "Type of BMT")
add_block("PHYS_TABLE_COL_4", "Table column heading.", "Timing")

add_h3("Referral Table — Row 1: Thalassemia Major")
add_block("PHYS_ROW_THAL_COND", "Condition cell.", "Thalassemia Major")
add_block("PHYS_ROW_THAL_IND", "Indication cell.", "Only curative treatment")
add_block("PHYS_ROW_THAL_TYPE", "Type of BMT cell.", "Allogeneic")
add_block("PHYS_ROW_THAL_TIMING", "Timing cell.", "Best age 2–5 years; earlier is better")

add_h3("Referral Table — Row 2: Sickle Cell Disease")
add_block("PHYS_ROW_SCD_COND", "Condition cell.", "Sickle Cell Disease")
add_block("PHYS_ROW_SCD_IND", "Indication cell.", "Severe disease; 90% cure with matched donor")
add_block("PHYS_ROW_SCD_TYPE", "Type of BMT cell.", "Allogeneic (incl. Haploidentical)")
add_block("PHYS_ROW_SCD_TIMING", "Timing cell.", "Before end-organ damage")

add_h3("Referral Table — Row 3: Aplastic Anemia")
add_block("PHYS_ROW_AA_COND", "Condition cell.", "Aplastic Anemia (Severe/Very Severe)")
add_block("PHYS_ROW_AA_IND", "Indication cell.", "Treatment of choice for young patients with matched sibling donor")
add_block("PHYS_ROW_AA_TYPE", "Type of BMT cell.", "Allogeneic")
add_block("PHYS_ROW_AA_TIMING", "Timing cell.", "Early; avoid multiple transfusions")

add_h3("Referral Table — Row 4: ALL")
add_block("PHYS_ROW_ALL_COND", "Condition cell.", "ALL — High Risk / MRD+")
add_block("PHYS_ROW_ALL_IND", "Indication cell.", "High Risk ALL; positive MRD after chemotherapy; relapsed ALL")
add_block("PHYS_ROW_ALL_TYPE", "Type of BMT cell.", "Allogeneic")
add_block("PHYS_ROW_ALL_TIMING", "Timing cell.", "First complete remission")

add_h3("Referral Table — Row 5: AML")
add_block("PHYS_ROW_AML_COND", "Condition cell.", "AML (non-APML)")
add_block("PHYS_ROW_AML_IND", "Indication cell.", "All AML except good risk or APML; relapsed AML")
add_block("PHYS_ROW_AML_TYPE", "Type of BMT cell.", "Allogeneic")
add_block("PHYS_ROW_AML_TIMING", "Timing cell.", "First complete remission")

add_h3("Referral Table — Row 6: CML")
add_block("PHYS_ROW_CML_COND", "Condition cell.", "CML — TKI-resistant / Accelerated / Blast Crisis")
add_block("PHYS_ROW_CML_IND", "Indication cell.", "Failed TKI therapy; accelerated phase; blast crisis")
add_block("PHYS_ROW_CML_TYPE", "Type of BMT cell.", "Allogeneic")
add_block("PHYS_ROW_CML_TIMING", "Timing cell.", "As soon as failure of TKI confirmed")

add_h3("Referral Table — Row 7: Hodgkin's Lymphoma")
add_block("PHYS_ROW_HL_COND", "Condition cell.", "Hodgkin's Lymphoma (Relapsed)")
add_block("PHYS_ROW_HL_IND", "Indication cell.", "Relapsed/refractory HL; after autologous BMT failure")
add_block("PHYS_ROW_HL_TYPE", "Type of BMT cell.", "Autologous → Allogeneic (if relapsed after auto)")
add_block("PHYS_ROW_HL_TIMING", "Timing cell.", "After salvage chemotherapy")

add_h3("Referral Table — Row 8: NHL")
add_block("PHYS_ROW_NHL_COND", "Condition cell.", "NHL (Relapsed/High Risk)")
add_block("PHYS_ROW_NHL_IND", "Indication cell.", "Relapsed NHL; T cell NHL; Mantle Cell; relapsed low grade")
add_block("PHYS_ROW_NHL_TYPE", "Type of BMT cell.", "Autologous → Allogeneic (if indicated)")
add_block("PHYS_ROW_NHL_TIMING", "Timing cell.", "PET-negative before BMT = 80% cure")

add_h3("Referral Table — Row 9: MDS")
add_block("PHYS_ROW_MDS_COND", "Condition cell.", "MDS")
add_block("PHYS_ROW_MDS_IND", "Indication cell.", "Only curative treatment; early before infections/iron overload")
add_block("PHYS_ROW_MDS_TYPE", "Type of BMT cell.", "Allogeneic")
add_block("PHYS_ROW_MDS_TIMING", "Timing cell.", "Before life-threatening complications")

add_h3("Referral Table — Row 10: Primary Immunodeficiency")
add_block("PHYS_ROW_PID_COND", "Condition cell.", "Primary Immunodeficiency (SCID etc.)")
add_block("PHYS_ROW_PID_IND", "Indication cell.", "Definitive cure; urgently at diagnosis for SCID")
add_block("PHYS_ROW_PID_TYPE", "Type of BMT cell.", "Allogeneic")
add_block("PHYS_ROW_PID_TIMING", "Timing cell.", "Ideally <3.5 months for SCID")

add_h3("Referral Table — Row 11: Inherited Bone Marrow Failure")
add_block("PHYS_ROW_IBMF_COND", "Condition cell.", "Inherited Bone Marrow Failure")
add_block("PHYS_ROW_IBMF_IND", "Indication cell.", "Before onset of leukemia/MDS or excessive transfusions")
add_block("PHYS_ROW_IBMF_TYPE", "Type of BMT cell.", "Allogeneic")
add_block("PHYS_ROW_IBMF_TIMING", "Timing cell.", "Early, as an elective procedure")

add_h3("Referral Table — Row 12: Inherited Metabolic Disorders")
add_block("PHYS_ROW_IMD_COND", "Condition cell.", "Inherited Metabolic Disorders")
add_block("PHYS_ROW_IMD_IND", "Indication cell.", "Before significant neurological damage")
add_block("PHYS_ROW_IMD_TYPE", "Type of BMT cell.", "Allogeneic")
add_block("PHYS_ROW_IMD_TIMING", "Timing cell.", "Early childhood")

add_h3("Referral Table — Row 13: Autoimmune Disease")
add_block("PHYS_ROW_AID_COND", "Condition cell.", "Autoimmune Disease (Severe)")
add_block("PHYS_ROW_AID_IND", "Indication cell.", "Scleroderma, MS — when conventional treatments fail")
add_block("PHYS_ROW_AID_TYPE", "Type of BMT cell.", "Autologous or Allogeneic")
add_block("PHYS_ROW_AID_TIMING", "Timing cell.", "Before permanent organ damage")

add_h2("Research & Credentials — Section Header")
add_block("PHYS_RESEARCH_OVERLINE", "Small label above the section title.", "Our Expertise")
add_block("PHYS_RESEARCH_TITLE", "Section title.", "Research & Credentials")
add_block("PHYS_RESEARCH_BODY", "Body text with an inline link ('Meet our full team').",
    "Led by internationally trained specialists. Meet our full team →")

add_h3("Research Output Card")
add_block("PHYS_RESEARCH_OUT_OVERLINE", "Small label at top of the highlighted card.", "Research Output")
add_block("PHYS_RESEARCH_OUT_TITLE", "Card heading with numbers.", "125+ Peer-Reviewed Publications · 4,700+ Citations")
add_block("PHYS_RESEARCH_OUT_P1", "Paragraph 1.",
    "Our team has published extensively on Haploidentical BMT, post-transplant viral infections, and immune reconstitution. We were the first in the world to identify specific risks in Haploidentical BMT for children and have developed the most effective approach for transplanting children from a parent or haploidentical donor.")
add_block("PHYS_RESEARCH_OUT_P2", "Paragraph 2.",
    "The only Indian centre with international standing in Haploidentical BMT research.")
add_block("PHYS_RESEARCH_BTN_RG", "Button linking out to ResearchGate profile.", "View on ResearchGate")
add_block("PHYS_RESEARCH_BTN_PUBMED", "Button linking out to a PubMed search.", "Search PubMed")

add_h3("Landmark Publications — Card Header")
add_block("PHYS_LANDMARK_OVERLINE", "Small label at top of the publications card.", "Landmark Publications")
add_block("PHYS_LANDMARK_TITLE", "Card heading.", "Selected Key Papers")

add_h3("Publication 1")
add_block("PHYS_PUB1_TITLE", "Paper title.",
    "Adenovirus infections following allogeneic stem cell transplantation: incidence and outcome in relation to graft manipulation, immunosuppression, and immune recovery")
add_block("PHYS_PUB1_CITATION", "Journal · year · authors line.",
    "Blood · 2002 · Chakrabarti S, Mautner V, Osman H, et al.")

add_h3("Publication 2")
add_block("PHYS_PUB2_TITLE", "Paper title.",
    "High incidence of cytomegalovirus infection after nonmyeloablative stem cell transplantation: potential role of Campath-1H in delaying immune reconstitution")
add_block("PHYS_PUB2_CITATION", "Journal · year · authors line.",
    "Blood · 2002 · Chakrabarti S, Mackinnon S, Chopra R, et al.")

add_h3("Publication 3")
add_block("PHYS_PUB3_TITLE", "Paper title.",
    "Respiratory virus infections in transplant recipients after reduced-intensity conditioning with Campath-1H: high incidence but low mortality")
add_block("PHYS_PUB3_CITATION", "Journal · year · authors line.",
    "British Journal of Haematology · 2002 · Chakrabarti S, Avivi I, Mackinnon S, et al.")

add_h3("Publication 4")
add_block("PHYS_PUB4_TITLE", "Paper title.",
    "Haploidentical Family Donor Transplantation: At the Crossroads of a Changing Paradigm")
add_block("PHYS_PUB4_CITATION", "Journal · year · authors line.",
    "Advances in Hematology · 2016 · Chakrabarti S")

add_h3("Publication 5")
add_block("PHYS_PUB5_TITLE", "Paper title.",
    "Pre-emptive oral ribavirin therapy of paramyxovirus infections after haematopoietic stem cell transplantation: a pilot study")
add_block("PHYS_PUB5_CITATION", "Journal · year · authors line.",
    "Bone Marrow Transplantation · 2001 · Chakrabarti S, Collingham K, Holder K, et al.")

add_h3("Publication 6")
add_block("PHYS_PUB6_TITLE", "Paper title.",
    "Cidofovir as primary pre-emptive therapy for post-transplant cytomegalovirus infections")
add_block("PHYS_PUB6_CITATION", "Journal · year · authors line.",
    "Bone Marrow Transplantation · 2001 · Chakrabarti S, Collingham K, Osman H, et al.")

add_h3("Publication 7")
add_block("PHYS_PUB7_TITLE", "Paper title.",
    "Fulminant adenovirus hepatitis following unrelated bone marrow transplantation: failure of intravenous ribavirin therapy")
add_block("PHYS_PUB7_CITATION", "Journal · year · authors line.",
    "Bone Marrow Transplantation · 1999 · Chakrabarti S, Collingham KE, Fegan C, Milligan DW")

add_h3("Publication 8")
add_block("PHYS_PUB8_TITLE", "Paper title.",
    "Post-transplant lymphoproliferative disorders following reduced intensity conditioning with in vivo T cell depletion")
add_block("PHYS_PUB8_CITATION", "Journal · year · authors line.",
    "Bone Marrow Transplantation · 2003 · Chakrabarti S, et al.")

add_h3("Publication 9")
add_block("PHYS_PUB9_TITLE", "Paper title.",
    "EBV-related disease following haematopoietic stem cell transplantation with reduced intensity conditioning")
add_block("PHYS_PUB9_CITATION", "Journal · year · authors line.",
    "Leukemia & Lymphoma · 2007 · Cohen JM, Cooper N, Chakrabarti S, et al.")

add_h3("Publication 10")
add_block("PHYS_PUB10_TITLE", "Paper title.",
    "Polyoma viruria following T-cell-depleted allogeneic transplants using Campath-1H: incidence and outcome")
add_block("PHYS_PUB10_CITATION", "Journal · year · authors line.",
    "Bone Marrow Transplantation · 2003 · Chakrabarti S, Osman H, Collingham K, et al.")

add_block("PHYS_PUB_FOOTER", "Footer line under the publications list (with inline links).",
    "Showing 10 of 125+ publications. View the complete list on ResearchGate or search PubMed.")

add_h2("Refer a Patient — Section Header")
add_block("PHYS_REFER_PATIENT_OVERLINE", "Small label above the section title.", "Get Started")
add_block("PHYS_REFER_PATIENT_TITLE", "Section title.", "Refer a Patient")

add_h3("Send a Referral Card")
add_block("PHYS_REFERRAL_CARD_TITLE", "Card heading.", "Send a Referral")
add_block("PHYS_REFERRAL_PHONE_LABEL", "Label above the referral phone number.", "Referral Phone")
add_block("PHYS_REFERRAL_PHONE_VALUE", "Referral phone number.", "+91-8287078906")
add_block("PHYS_REFERRAL_WA_LABEL", "Label above the WhatsApp number.", "WhatsApp")
add_block("PHYS_REFERRAL_WA_VALUE", "WhatsApp number.", "+91-9871264073")
add_block("PHYS_REFERRAL_LOC_LABEL", "Label above the address.", "Location")
add_block("PHYS_REFERRAL_LOC_VALUE", "Address shown on the referral card.",
    "Action Cancer Hospital, A-4, Paschim Vihar, New Delhi – 110063")

add_h3("Teleconsultation Card")
add_block("PHYS_TELE_TITLE", "Card heading.", "Teleconsultation Available")
add_block("PHYS_TELE_BODY", "Card body paragraph.",
    "We offer teleconsultation for remote physician-to-physician consultation on complex cases and BMT suitability assessment.")
add_block("PHYS_TELE_MODES", "Small line listing available modes of teleconsultation.",
    "Modes: Phone call · Text/WhatsApp · Video call")
add_block("PHYS_TELE_CTA", "Button at the bottom of the card.", "Contact for Referral")

add_page_break()


# ============================================================
# PAGE: FOR PATIENTS / RESOURCES (#resources)
# ============================================================
add_h1("For Patients Page (Patient Resources)")
add_hint("Tabbed page for patients and families. Tabs: Your BMT Journey · Clean Room Guide · Nutrition · Donor Guide · Patient Stories.")

add_h2("Page Hero")
add_block("RES_HERO_OVERLINE", "Small label above the page title.", "Information for Patients & Families")
add_block("RES_HERO_TITLE", "Page title.", "For Patients")
add_block("RES_HERO_SUBTITLE", "Subtitle paragraph.",
    "Guides written by Dr. Suparno Chakrabarti and Dr. Mahak Agarwal to help patients and families prepare for Bone Marrow Transplantation.")

add_h2("Tab Labels")
add_block("RES_TAB_1", "Tab button label.", "Your BMT Journey")
add_block("RES_TAB_2", "Tab button label.", "Clean Room Guide")
add_block("RES_TAB_3", "Tab button label.", "Nutrition")
add_block("RES_TAB_4", "Tab button label.", "Donor Guide")
add_block("RES_TAB_5", "Tab button label.", "Patient Stories")

# ---------- TAB 1: BMT JOURNEY ----------
add_h2("Tab 1 — Your BMT Journey")
add_block("RES_JOURNEY_OVERLINE", "Small label above the tab title.", "For Patients")
add_block("RES_JOURNEY_TITLE", "Tab section title.", "Your BMT Journey")
add_block("RES_JOURNEY_BODY", "Body paragraph under the title.",
    "What to expect before, during, and after your Bone Marrow Transplant.")

add_h3("Accordion 1 — Step-Down Isolation")
add_block("RES_JOURNEY_STEPDOWN_TITLE", "Accordion title.", "Step-Down Isolation")
add_block("RES_JOURNEY_STEPDOWN_BODY", "Accordion body paragraph.",
    "Once transplanted bone marrow starts to grow, or engraft, the white blood cell count will rise. Once the white blood cell count reaches safe levels, the patient can be moved from the high-level isolation (clean room) to the low-level isolation. In low-level isolation, standard precautions are taken like hand washing and not visiting if sick. This period is also known as Step down isolation.")

add_h3("Accordion 2 — Central Venous Lines (CVL)")
add_block("RES_JOURNEY_CVL_TITLE", "Accordion title.", "Central Venous Lines (CVL)")
add_block("RES_JOURNEY_CVL_P1", "Accordion body, paragraph 1.",
    "If you do not have a double lumen central venous catheter, you need one for the transplant period. The transplant doctor or an anaesthetist will place the line in the operating room, before the conditioning starts.")
add_block("RES_JOURNEY_CVL_P2", "Accordion body, paragraph 2.",
    "This CVL remains in place for the entire transplant period. It may be used to give intravenous (IV) fluids, medications, blood products, and for the transplant itself. It is also used to draw most blood samples. The BMT nurses will flush the line with a medication (heparin) to keep it from clotting.")
add_block("RES_JOURNEY_CVL_P3", "Accordion body, paragraph 3.",
    "When the transplant doctor decides that the line is no longer needed, it will be removed in the operating room or at the bedside.")

# ---------- TAB 2: CLEAN ROOM ----------
add_h2("Tab 2 — Clean Room Guide")
add_block("RES_CR_OVERLINE", "Small label above the tab title.", "For Patients")
add_block("RES_CR_TITLE", "Tab section title.", "Clean Room Guide")
add_block("RES_CR_BODY", "Body paragraph under the title.",
    "Protective isolation details, precautions, and equipment in the BMT Unit.")

add_h3("Accordion 1 — Protective Isolation")
add_block("RES_CR_ISO_TITLE", "Accordion title.", "Protective Isolation — What Are Clean Rooms?")
add_block("RES_CR_ISO_P1", "Accordion body, intro paragraph.",
    "Transplant patients have very low immunity to fight bacterial, viral and fungal infections and therefore need to be isolated into CLEAN ROOMS which will protect them from infections.")
add_block("RES_CR_ISO_AHU_H", "Sub-heading.", "Air Handling Unit (AHU)")
add_block("RES_CR_ISO_AHU_P", "Paragraph.",
    "Each room of the BMT Unit has its own dedicated Air Handling Unit (AHU) to provide 10–15 Hepafiltered fresh air changes per hour. This means the fresh air entering the patient's room is first treated through special filters. Treated fresh air then passes through 0.3 Micron High Efficiency Particulate Air Filter (HEPA). HEPA removes all the bacteria, viruses and fungus. This hepafiltered air passes through a laminar floor to reach the patient's room at the desired humidity and temperature.")
add_block("RES_CR_ISO_ACS_H", "Sub-heading.", "Automatic and Selective Control System")
add_block("RES_CR_ISO_ACS_P", "Paragraph.",
    "Automatic and selective control system provides positive air pressure in the BMT room compared to ante room and the BMT corridor. This has been done to ensure that, on opening the door of the BMT Room or the ante room, no outside air from BMT corridor or the ante room enters the BMT room.")
add_block("RES_CR_ISO_ANTE_H", "Sub-heading.", "Anteroom")
add_block("RES_CR_ISO_ANTE_P", "Paragraph.",
    "Ante room is a small room between the corridor and the BMT room for maintaining positive air pressure in the BMT room.")
add_block("RES_CR_ISO_SURFACES_H", "Sub-heading.", "Stainless Steel Doors, Vinyl Flooring and Cladding of Walls")
add_block("RES_CR_ISO_SURFACES_P", "Paragraph.",
    "Stainless steel doors and vinyl surfaces are most practical to clean with disinfectants to maintain high standards of hygiene and infection control.")

add_h3("Accordion 2 — Precautions Before Entering the BMT Unit")
add_block("RES_CR_BEFORE_TITLE", "Accordion title.", "Precautions Before Entering the BMT Unit")
add_label("RES_CR_BEFORE_LIST")
add_hint("Numbered/bulleted precautions list (10 items).")
add_bullets([
    "Do not carry any valuables e.g. jewellery, cash etc. to the BMT room",
    "Ring the bell outside the BMT unit, so that the nurse can decide whether to let you in or not.",
    "If allowed, keep your cell phone, toys, games, pictures etc. in the pass box fitted with ultraviolet light for 20 minutes before taking them to the BMT room.",
    "Enter the change room",
    "Change shoes",
    "Change into sterile clothes, wear cap and mask.",
    "Wash your hands with soap and water by following 6 step hand wash instructions, written on the wall above the sink.",
    "Use all the three disinfectants (placed above the sink) one by one. Dry your hands",
    "Enter the BMT corridor.",
    "Enter the Anteroom, scrub again as mentioned above, wear a sterile gown and then enter the BMT room.",
])

add_h3("Accordion 3 — Precautions Inside the BMT Room")
add_block("RES_CR_INSIDE_TITLE", "Accordion title.", "Precautions Inside the BMT Room")
add_block("RES_CR_INSIDE_ATT_H", "Sub-heading.", "Attendants")
add_block("RES_CR_INSIDE_ATT_P", "Paragraph.",
    "No attendants are allowed except with children and very sick patients with the doctor's permission. Only one attendant is allowed in the BMT unit at any time.")
add_block("RES_CR_INSIDE_TOIL_H", "Sub-heading.", "Toilets")
add_block("RES_CR_INSIDE_TOIL_P", "Paragraph.",
    "Attached toilet is exclusively meant for the patient. Only treated water should be used for patient's bath and mouth wash.")
add_block("RES_CR_INSIDE_FOOD_H", "Sub-heading.", "Food")
add_block("RES_CR_INSIDE_FOOD_P1", "Paragraph.",
    "No outside food is allowed inside the BMT room. Patients will be served prescribed, pressure cooked and microwaved food inside the BMT room. Attendants will be served food in the BMT pantry as per their order.")
add_block("RES_CR_INSIDE_FOOD_NOTE", "Italic closing note.",
    "Note: BMT Nurse will explain all the above in detail to the patient and the family for their full cooperation.")

add_h3("Accordion 4 — Equipment & Facilities")
add_block("RES_CR_EQUIP_TITLE", "Accordion title.", "Equipment & Facilities in the BMT Unit")
add_block("RES_CR_EQUIP_EQ_H", "Sub-heading.", "Equipment")
add_block("RES_CR_EQUIP_EQ_P", "Paragraph.",
    "BMT Unit is equipped with a dedicated X-ray machine, ultrasound machine, dialysis machine and a ventilator to handle any medical emergency.")
add_block("RES_CR_EQUIP_FURN_H", "Sub-heading.", "Furniture and Furnishing of BMT Room")
add_label("RES_CR_EQUIP_FURN_LIST")
add_hint("Bulleted list of furnishings in the BMT room.")
add_bullets([
    "Centralised oxygen, suction system with double outlets",
    "State-of-art, six parameter monitors for monitoring heart rate, breathing rate, blood pressure, oxygen saturation and ECG 24x7.",
    "Infusion pumps",
    "Syringe pumps",
    "C.C. TV monitoring system for 24x7 vigilance of the patient",
    "Television",
    "Telephone",
    "Patient bed, bedside locker and desk",
    "Attendant bed",
    "Two trolleys for medicines/dressing etc.",
])

# ---------- TAB 3: NUTRITION ----------
add_h2("Tab 3 — Nutrition Guidelines")
add_block("RES_NUT_OVERLINE", "Small label above the tab title.", "For Patients")
add_block("RES_NUT_TITLE", "Tab section title.", "Nutrition Guidelines")
add_block("RES_NUT_BODY", "Body paragraph under the title.",
    "Diet guidelines, restrictions, and nutritional support during BMT.")

add_block("RES_NUT_P1", "Paragraph 1.",
    "During bone marrow transplantation, the bone marrow is destroyed by high doses of chemotherapy. This causes a decrease in white blood cells, which help to fight infections. Because of this, it is very important to eat foods that are less likely to contain high levels of bacteria.")
add_block("RES_NUT_P2", "Paragraph 2.",
    "When Patients are admitted in the isolation room, they are placed on a low bacteria diet. Most food items served in the hospital are pressure cooked and microwaved before serving to patients.")
add_block("RES_NUT_P3", "Paragraph 3.",
    "Home prepared foods are NOT allowed. Commercially packaged items are ONLY allowed after they have been approved by the dietician or nurse. Food should NOT be left at room temperature for longer than one hour.")
add_block("RES_NUT_P4", "Paragraph 4.",
    "Eating well is very important during transplant and recovery. Patients may have low appetite, change in taste, dry mouth and/or nausea. There may be times when Patients may not feel well enough to eat. To maintain the nutrition, TPN (Total Parental Nutrition) or tube feeding is started.")

add_block("RES_NUT_LIST_HEAD", "Sub-heading above the guidelines list.", "Low Bacteria Diet Guidelines")
add_label("RES_NUT_LIST")
add_hint("Seven-point bulleted list of diet dos and don'ts.")
add_bullets([
    "No restaurant food, take out, cafeteria food or vendor food is allowed.",
    "All foods must be cooked thoroughly. Avoid rare to medium cooked meats and fish.",
    "Herbs, spices and pepper should not be added to food after it is cooked, but are allowed when cooked with the food.",
    "Avoid raw fruits and vegetables including salads, garnishes, stir-fried vegetables, egg rolls and any fruit garnish on a dessert.",
    "Avoid foil-sealed plastic cups of juices because they do not have best before date.",
    "Avoid food containing raw eggs including soft cooked eggs.",
    "Dried fruits, nuts and seeds are not allowed unless cooked in a food item.",
])

# ---------- TAB 4: DONOR GUIDE ----------
add_h2("Tab 4 — Donor Guide")
add_block("RES_DONOR_OVERLINE", "Small label above the tab title.", "For Donors")
add_block("RES_DONOR_TITLE", "Tab section title.", "Donor Guide")

add_h3("Accordion 1 — Who Can Be a Donor?")
add_block("RES_DONOR_WHO_TITLE", "Accordion title.", "Who Can Be a Donor for BMT?")
add_block("RES_DONOR_WHO_INTRO", "Intro paragraph.",
    "Donor for BMT has to be matched with the patient in their 'tissue type'. This is confirmed by typing their HLA antigens.")
add_block("RES_DONOR_WHO_FAM_H", "Sub-heading.", "Family Donor — Fully HLA Matched")
add_block("RES_DONOR_WHO_FAM_P", "Paragraph.",
    "Within a family, there is about 25–30% chance of finding such a match in a brother or a sister. If there is no match within the close family, the chances of finding a fully HLA matched donor in distant relatives is remote.")
add_block("RES_DONOR_WHO_HAP_H", "Sub-heading.", "Half-Matched / Haploidentical")
add_block("RES_DONOR_WHO_HAP_P", "Paragraph.",
    "We inherit two sets of HLA ANTIGENS; one from each parent. Thus, parents are always half matched with us. In addition, even if the brothers and sisters are not fully matched with the patient, there is 90% chance that they shall be half-matched. BMT from a half matched or HAPLOIDENTICAL donor is feasible in centres with adequate infrastructure and expertise.")
add_block("RES_DONOR_WHO_VOL_H", "Sub-heading.", "Volunteer Unrelated Donors")
add_block("RES_DONOR_WHO_VOL_P", "Paragraph.",
    "To find a match with a random person is less than one in a billion. However, if we screen million people of similar ethnic background, we might find a close match. Based on this concept, volunteer unrelated donor registries have been set up in all countries. In India also, we have many volunteer unrelated donor registries. It takes a few months to search for an HLA matched donor in these registries.")

add_h3("Accordion 2 — Donating Bone Marrow")
add_block("RES_DONOR_BM_TITLE", "Accordion title.", "Donating Bone Marrow — Is It Safe?")
add_block("RES_DONOR_BM_BODY", "Accordion body paragraph.",
    "Hundred milliliter to a litre of bone marrow (depending on the age of the patient and the donor) is removed from the donor's hip bones under general anesthesia. It is an extremely safe procedure and is completed within two to three hours. The donor can be discharged on the same day or next morning and he can go back to his work within two to three days. Bone marrow is naturally regenerated in the body within 2 to 3 weeks.")

add_h3("Accordion 3 — Donating Peripheral Blood Stem Cells")
add_block("RES_DONOR_PBSC_TITLE", "Accordion title.", "Donating Peripheral Blood Stem Cells")
add_block("RES_DONOR_PBSC_BODY", "Accordion body paragraph.",
    "No. The donor receives an injection of a growth-factor for 4 days and on the fifth day the donation of blood stem cells is done. It is similar to a long blood donation through a machine and no operation or general anaesthesia is required. The donor can go back to work on the same day.")

# ---------- TAB 5: PATIENT STORIES ----------
add_h2("Tab 5 — Patient Stories")
add_block("RES_STORIES_OVERLINE", "Small label above the tab title.", "Patient Testimonials")
add_block("RES_STORIES_TITLE", "Tab section title.", "Patient Stories")
add_block("RES_STORIES_BODY", "Body paragraph under the title.",
    "Real stories from patients and families who have been treated at Bloods <em>R</em> Us under Dr. Suparno Chakrabarti and Dr. Mahak Agarwal.")

add_h3("Testimonial 1")
add_block("RES_TESTI_1_QUOTE", "Quote.",
    '"My father was diagnosed with Mantle Cell Lymphoma on 3rd Feb 2021. My father start receiving his treatment from Dr. Suparno and his team and he is fine now with his health."')
add_block("RES_TESTI_1_AUTHOR", "Author name.", "Anonymous")
add_block("RES_TESTI_1_ROLE", "Author role / condition.", "Patient's Family — Mantle Cell Lymphoma")

add_h3("Testimonial 2")
add_block("RES_TESTI_2_QUOTE", "Quote.",
    '"The care and dedication shown by the entire Bloods <em>R</em> Us team is truly remarkable. From diagnosis to treatment and follow-up, every step was handled with professionalism and compassion."')
add_block("RES_TESTI_2_AUTHOR", "Author name.", "Nitin Garg")
add_block("RES_TESTI_2_ROLE", "Author role / condition.", "Patient")

add_h3("Testimonial 3")
add_block("RES_TESTI_3_QUOTE", "Quote.",
    '"I am very satisfied with my course of treatment and the doctors and staff under Dr. Suparno sir are very good and cooperative."')
add_block("RES_TESTI_3_AUTHOR", "Author name.", "Kamal Sarkar")
add_block("RES_TESTI_3_ROLE", "Author role / condition.", "Patient")

add_h3("Testimonial 4")
add_block("RES_TESTI_4_QUOTE", "Quote.",
    '"The treatment was done by expertise doctors."')
add_block("RES_TESTI_4_AUTHOR", "Author name.", "Anonymous")
add_block("RES_TESTI_4_ROLE", "Author role / condition.", "Patient")

add_h3("Testimonial 5")
add_block("RES_TESTI_5_QUOTE", "Quote.",
    '"I have taken treatment of Acute Myeloid Leukemia (AML) for my daughter. My daughter had a successful Bone Marrow Transplant."')
add_block("RES_TESTI_5_AUTHOR", "Author name.", "Aimira Isamidinov")
add_block("RES_TESTI_5_ROLE", "Author role / condition.", "Parent of Patient — AML")

add_h3("Testimonial 6")
add_block("RES_TESTI_6_QUOTE", "Quote.",
    '"Story of Mr. Jagdish Sharma, Lymphoma Patient after his successful Bone Marrow Transplant."')
add_block("RES_TESTI_6_AUTHOR", "Author name.", "Mr. Jagdish Sharma")
add_block("RES_TESTI_6_ROLE", "Author role / condition.", "Patient — Lymphoma")

add_h3("Testimonial 7")
add_block("RES_TESTI_7_QUOTE", "Quote.",
    '"Success Story of Fatima Yahaya Zango from Nigeria. Treated for Sickle Cell Disease."')
add_block("RES_TESTI_7_AUTHOR", "Author name.", "Fatima Yahaya Zango")
add_block("RES_TESTI_7_ROLE", "Author role / condition.", "Patient — Sickle Cell Disease, Nigeria")

add_page_break()


# ============================================================
# PAGE: CONTACT (#contact)
# ============================================================
add_h1("Contact Page")
add_hint("Final page: location, hours, phone numbers, teleconsultation note, and the enquiry form.")

add_h2("Page Hero")
add_block("CONTACT_HERO_OVERLINE", "Small label above the page title.", "Reach Us")
add_block("CONTACT_HERO_TITLE", "Page title.", "Contact Bloods <em>R</em> Us")
add_block("CONTACT_HERO_SUBTITLE", "Subtitle paragraph.",
    "Book an appointment, ask a question, or request a teleconsultation. We are here to help.")

add_h2("Location & Hours Column")
add_block("CONTACT_LOC_OVERLINE", "Small label above the left column.", "Find Us")
add_block("CONTACT_LOC_TITLE", "Left column heading.", "Location & Hours")

add_h3("Address Item")
add_block("CONTACT_ADDR_LABEL", "Label.", "Address")
add_block("CONTACT_ADDR_VALUE", "Full address shown on the contact page.",
    "Action Cancer Hospital, A-4, Paschim Vihar, Near Paschim Vihar East Metro Station, New Delhi – 110063, Delhi, India")

add_h3("Clinic Hours Item")
add_block("CONTACT_HOURS_LABEL", "Label.", "Clinic Hours")
add_block("CONTACT_HOURS_DAYS", "Days line.", "Monday – Saturday")
add_block("CONTACT_HOURS_TIMES", "Times line.",
    "10:00 AM – 1:00 PM & 6:00 PM – 8:00 PM")

add_h3("Primary Contact")
add_block("CONTACT_PRIMARY_LABEL", "Label.", "Primary Contact")
add_block("CONTACT_PRIMARY_VALUE", "Primary phone number.", "+91-8287078906")

add_h3("Appointment Number")
add_block("CONTACT_APPT_LABEL", "Label.", "Appointment")
add_block("CONTACT_APPT_VALUE", "Appointment phone number.", "+91-9911800616")

add_h3("WhatsApp Item")
add_block("CONTACT_WA_LABEL", "Label.", "WhatsApp")
add_block("CONTACT_WA_VALUE", "WhatsApp number.", "+91-9871264073")

add_h3("Teleconsultation Card (under the contact info)")
add_block("CONTACT_TELE_TITLE", "Card title.", "Teleconsultation")
add_block("CONTACT_TELE_P1", "Card paragraph 1.",
    "Your convenience is our priority! Telemedicine allows you to consult with us from the comfort of your home, and without any delays.")
add_block("CONTACT_TELE_MODES_LABEL", "Bold line above the modes.", "Available modes:")
add_block("CONTACT_TELE_MODES_VALUE", "Modes line.",
    "Phone call · Text/WhatsApp chat · Video call")
add_block("CONTACT_TELE_DISCLAIMER", "Small italic disclaimer at the bottom of the card.",
    "Telemedicine is not a substitute for in-person care. For serious medical conditions, please go to the nearest emergency room.")

add_h2("Enquiry Form Column")
add_block("CONTACT_FORM_OVERLINE", "Small label above the right column.", "Get in Touch")
add_block("CONTACT_FORM_TITLE", "Right column heading.", "Send an Enquiry")

add_h3("Form Field Labels & Placeholders")
add_block("CONTACT_FIELD_NAME_LABEL", "Input label (required).", "Full Name *")
add_block("CONTACT_FIELD_NAME_PLACEHOLDER", "Input placeholder.", "Your full name")
add_block("CONTACT_FIELD_NAME_ERROR", "Error message shown if name is blank.", "Please enter your name.")

add_block("CONTACT_FIELD_EMAIL_LABEL", "Input label.", "Email Address")
add_block("CONTACT_FIELD_EMAIL_PLACEHOLDER", "Input placeholder.", "your@email.com")

add_block("CONTACT_FIELD_PHONE_LABEL", "Input label.", "Phone Number")
add_block("CONTACT_FIELD_PHONE_PLACEHOLDER", "Input placeholder.", "+91-XXXXXXXXXX")

add_block("CONTACT_FIELD_FOR_LABEL", "Select label.", "Enquiry For")
add_label("CONTACT_FIELD_FOR_OPTIONS")
add_hint("Dropdown options for 'Enquiry For'.")
add_bullets([
    "Select a topic",
    "Appointment Booking",
    "Thalassemia / Sickle Cell",
    "Aplastic Anemia",
    "Leukemia (ALL/AML/CML)",
    "Hodgkin's / Non-Hodgkin's Lymphoma",
    "MDS",
    "BMT Evaluation",
    "Teleconsultation",
    "Second Opinion",
    "Other",
])

add_block("CONTACT_FIELD_DOCTOR_LABEL", "Select label.", "Preferred Doctor")
add_label("CONTACT_FIELD_DOCTOR_OPTIONS")
add_hint("Dropdown options for 'Preferred Doctor'.")
add_bullets([
    "No preference",
    "Dr. Suparno Chakrabarti",
    "Dr. Mahak Agarwal",
])

add_block("CONTACT_FIELD_MESSAGE_LABEL", "Textarea label.", "Message")
add_block("CONTACT_FIELD_MESSAGE_PLACEHOLDER", "Textarea placeholder.",
    "Briefly describe the patient's condition and your query...")

add_h3("Form Submit & Success")
add_block("CONTACT_FORM_SUBMIT", "Submit button.", "Send Enquiry")
add_block("CONTACT_FORM_RESPONSE_NOTE", "Small note under the submit button.",
    "We aim to respond within 24 hours on working days.")
add_block("CONTACT_FORM_SUCCESS_TITLE", "Green success card — title.", "Enquiry Sent Successfully")
add_block("CONTACT_FORM_SUCCESS_TEXT", "Green success card — body.",
    "Thank you for reaching out. Our team will respond within 24 hours on working days.")
add_block("CONTACT_FORM_SUBMIT_AFTER", "Button text shown immediately after submitting.",
    "Enquiry Sent")

add_page_break()


# ============================================================
# SEO / METADATA
# ============================================================
add_h1("SEO / Metadata")
add_hint("Text that lives in the HTML <head> or JSON structured data. Not visible on the page itself, but shown by Google, social-share previews, and browser tabs. Safe to edit for wording / clarity / SEO.")

add_h2("Browser Tab Title")
add_block("SEO_TITLE", "The <title> tag — appears in the browser tab and as the headline in Google search results.",
    "Bloods R Us — Centre of Excellence for Blood Disorders | Dr. Suparno Chakrabarti")

add_h2("Meta Description")
add_block("SEO_META_DESCRIPTION",
    "The description shown under the title in Google search results. Keep under ~160 characters for best display.",
    "Bloods R Us — Dr. Suparno Chakrabarti's Centre of Excellence for Blood Disorders. Pioneers of Haploidentical BMT. 250+ procedures, 75% long-term survival. Action Cancer Hospital, New Delhi.")

add_h2("Meta Keywords & Author")
add_block("SEO_META_KEYWORDS", "Comma-separated keywords used by some search engines.",
    "blood disorders, bone marrow transplant, BMT, haploidentical, thalassemia, leukemia, lymphoma, haematology, Dr Suparno Chakrabarti, Action Cancer Hospital, New Delhi")
add_block("SEO_META_AUTHOR", "Author meta tag.",
    "Dr. Suparno Chakrabarti, Dr. Mahak Agarwal")

add_h2("Open Graph (Facebook, WhatsApp, LinkedIn share previews)")
add_block("SEO_OG_TITLE", "Large title shown when the site is shared on social media / WhatsApp.",
    "Bloods R Us — Centre of Excellence for Blood Disorders")
add_block("SEO_OG_DESCRIPTION", "Description line shown in social share previews.",
    "Internationally acclaimed haematologists and BMT specialists. Pioneers of Haploidentical Bone Marrow Transplantation in India. 250+ procedures, 75% long-term survival.")
add_block("SEO_OG_SITE_NAME", "Small site-name label in the share preview.", "Bloods R Us")

add_h2("Twitter / X Card")
add_block("SEO_TW_TITLE", "Title shown when the site is shared on Twitter / X.",
    "Bloods R Us — Centre of Excellence for Blood Disorders")
add_block("SEO_TW_DESCRIPTION", "Description line in the Twitter / X card.",
    "Internationally acclaimed haematologists and BMT specialists. Pioneers of Haploidentical BMT in India.")

add_h2("Structured Data (schema.org) — Organization")
add_hint("JSON-LD fields used by Google to build the rich 'Knowledge Panel' on the right-hand side of search results. These are user-visible strings.")
add_block("SEO_SCHEMA_ORG_NAME", "Clinic name in structured data.",
    "Bloods R Us — Centre of Excellence for Blood Disorders")
add_block("SEO_SCHEMA_ORG_DESCRIPTION", "Clinic description in structured data.",
    "Internationally acclaimed haematologists and BMT specialists offering world-class care. Pioneers of Haploidentical Bone Marrow Transplantation in India.")
add_block("SEO_SCHEMA_ORG_STREET", "Structured-data street address.",
    "Action Cancer Hospital, A-4, Paschim Vihar")
add_block("SEO_SCHEMA_ORG_LOCALITY", "Structured-data locality.", "New Delhi")
add_block("SEO_SCHEMA_ORG_REGION", "Structured-data region.", "Delhi")
add_block("SEO_SCHEMA_ORG_POSTAL", "Structured-data postal code.", "110063")
add_block("SEO_SCHEMA_ORG_COUNTRY", "Structured-data country code.", "IN")
add_block("SEO_SCHEMA_ORG_TELEPHONE", "Structured-data telephone.", "+91-8287078906")

add_label("SEO_SCHEMA_ORG_SPECIALTIES")
add_hint("Medical specialties listed in structured data.")
add_bullets([
    "Hematology",
    "Oncology",
    "Bone Marrow Transplantation",
])

add_label("SEO_SCHEMA_ORG_SERVICES")
add_hint("Services listed in structured data.")
add_bullets([
    "Haploidentical Bone Marrow Transplantation",
    "Allogeneic BMT",
    "Autologous BMT",
    "CAR-T Cell Therapy",
])

add_block("SEO_SCHEMA_ORG_HOURS", "Opening hours block in structured data.",
    "Monday–Saturday · 09:00 – 17:00")

add_h2("Structured Data — Dr. Suparno Chakrabarti")
add_block("SEO_SCHEMA_SUPARNO_NAME", "Name.", "Dr. Suparno Chakrabarti")
add_block("SEO_SCHEMA_SUPARNO_SPECIALTY", "Medical specialty.", "Hematology")
add_block("SEO_SCHEMA_SUPARNO_JOBTITLE", "Job title used in structured data.",
    "Principal Director, Action Institute for Blood Diseases, Transplantation & Cellular Therapy (AIBTCT)")

add_label("SEO_SCHEMA_SUPARNO_CREDENTIALS")
add_hint("Credentials listed in structured data.")
add_bullets([
    "MD (Internal Medicine) — PGIMER, Chandigarh",
    "Doctor of Medicine (Hematopoietic Cell Transplantation) — University of Birmingham, UK",
    "FRCPath (Haematology) — Royal College of Pathologists, London",
    "CCT in Haematology — UK",
])

add_label("SEO_SCHEMA_SUPARNO_AFFILIATIONS")
add_hint("Affiliations listed in structured data.")
add_bullets([
    "Action Cancer Hospital, New Delhi",
    "Amity University (Hon Professor)",
    "Jamia Hamdard University (Adjunct Professor)",
    "Manashi Chakrabarti Foundation (Director & Senior Scientist)",
    "CIBMTR Expert Panel (Post-Transplant Infections)",
    "EBMT Working Party on Infectious Diseases",
])

add_label("SEO_SCHEMA_SUPARNO_EXPERTISE")
add_hint("'Knows about' list in structured data — the scientific topics Dr. Chakrabarti is associated with.")
add_bullets([
    "Haploidentical BMT",
    "T Cell Co-stimulation Blockade",
    "Adaptive NK Cell Biology",
    "NK Cell-Based Immunotherapy",
    "Viral Immunology Post-BMT",
    "Reduced Intensity Conditioning",
    "Campath-1H T Cell Depletion",
])

add_h2("Structured Data — Dr. Mahak Agarwal")
add_block("SEO_SCHEMA_MAHAK_NAME", "Name.", "Dr. Mahak Agarwal")
add_block("SEO_SCHEMA_MAHAK_SPECIALTY", "Medical specialty.", "Hematology")
add_block("SEO_SCHEMA_MAHAK_JOBTITLE", "Job title.", "Consultant Haematologist")

add_page_break()


# ============================================================
# SAVE
# ============================================================
# Count how many [LABEL] blocks we produced.
# We count paragraphs whose first run text starts with '[' and ends with ']'.
_block_count = 0
for _p in doc.paragraphs:
    if _p.runs:
        _t = _p.runs[0].text or ""
        if _t.startswith("[") and _t.endswith("]"):
            _block_count += 1

_output_path = '/Users/shivambansal/Desktop/Claude Code/BloodsRUs-Redesigned/Bloods-R-Us_Website_Copy_for_Review.docx'
doc.save(_output_path)
print(f"Saved: {_output_path}")
print(f"Total editable blocks: {_block_count}")
